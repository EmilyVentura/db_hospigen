#Universidad del Valle de Guatemala
#Autora: Emily Ventura
#Ingeniería Biomédica
#Proyecto de graduación "Optimización de gestión de datos dentro de los procesos de donación de órganos en el Hospital General San Juan De Dios"


###################### Importación de módulos #########################
import sys
from PyQt5.QtWidgets import QMessageBox, QMainWindow, QApplication, QPushButton, QVBoxLayout, QMessageBox, QTableWidgetItem, QFileDialog, QInputDialog, QLineEdit,QDialog, QLabel
from PyQt5.QtCore import pyqtSlot, QDate
from PyQt5.QtGui import QIntValidator, QPixmap
from datetime import datetime
import matplotlib.pyplot as plt
from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg as FigureCanvas
import pandas as pd
import matplotlib.dates as mdates
import mplcursors
from PyQt5 import QtWidgets
from docx import Document
import os
import pprint
import bcrypt

# Conexión a hacia archivos propios, sobre el inicio de sesión (Login), interfaz 
from interfaz_ui import Ui_Form
from tesis_conexionSQL import ConectaDB
from tesis_login import LoginWindow


#################### Código principal ############################

# Clase para la ventana principal que hereda de QMainWindow
class MainWindow(QMainWindow):

    #Funciones que manejan la autenticación de usuarios

    def mostrar_login(self):
        #Muestra la ventana de login y conecta las señales
        self.login_window = LoginWindow()
        self.login_window.sesion_exitosa.connect(self.inicio_exitoso)
        self.login_window.show()

    def inicio_exitoso(self, username):
        #Esta es la función que se ejecutará cuando el usuario logre entrar al sistema exitosamente
        self.usuario_actual = username
        print(f"Usuario autenticado: {self.usuario_actual}")
        self.show()
        self.login_window.close()

    def solicitar_contraseña(self):
        dialog = contraseña_dialogo(self)
        return dialog.obtener_contraseña()

    def validar_contraseña(self, password):
        try:
            if not self.usuario_actual:
                raise ValueError("No hay un usuario autenticado.")

            self.db.conecta_base_datos()
            query = "SELECT contrasena FROM administradores WHERE usuario = %s"
            self.db.cursor.execute(query, (self.usuario_actual,))
            result = self.db.cursor.fetchone()

            if result is None:
                print("Usuario no encontrado en la base de datos.")
                return False

            #Se comprueba que si está detectado el hash almacenado en la base de datos
            hash_almacenado = result['contrasena']
            print(f"Contraseña ingresada: {password}")
            print(f"Hash almacenado: {hash_almacenado}")

            return bcrypt.checkpw(password.encode('utf-8'), hash_almacenado.encode('utf-8'))

        except Exception as e:
            print(f"Error al validar la contraseña: {e}")
            return False

        finally:
            self.db.cursor.close()
            self.db.con.close()

    # Función para iniciar ventana principal y configuración de componentes básicos de la interfaz con PyQt5
    def __init__(self):
        super().__init__() # Llama a la clase padre (QMainWindow)

        # Inicializa la interfaz de usuario desde el archivo UI generado por Qt Designer
        self.ui = Ui_Form()
        self.ui.setupUi(self)
        self.id_paciente_actual = None
        self.usuario_actual = None
        self.habilitar_radio_buttons(False)
        self.db =  ConectaDB()
        self.hide()
        self.mostrar_login()

        # Conexión de cada elemento de la UI con variables de la clase para facilitar  el acceso y manipulación de la información de pacientes y fases

        #Datos de paciente : receptor
        self.id_receptor = self.ui.lineEdit_34
        self.id_receptor.setValidator(QIntValidator())  # Restringe a enteros
        self.fecha_registro = self.ui.dateEdit_21
        self.nombre = self.ui.lineEdit_44
        self.edad = self.ui.spinBox_26
        self.etnia = self.ui.comboBox_66
        self.sexo = self.ui.comboBox_68
        self.fecha_dg_erc = self.ui.dateEdit_12
        self.ter_sust_act = self.ui.comboBox_73
        self.inst_provee_hd = self.ui.comboBox_78
        self.vol_residual = self.ui.doubleSpinBox_41
        self.tiempo_anuria = self.ui.dateEdit_15
        self.grupo_sanguineo = self.ui.comboBox_67
        self.procedencia = self.ui.comboBox_74
        self.residencia = self.ui.comboBox_75
        self.ocupacion = self.ui.comboBox_69
        self.etiologia_erc = self.ui.comboBox_70
        self.tiempo_ini_st_renal = self.ui.dateEdit_13
        self.tiempo_dialisis = self.ui.dateEdit_14
        self.riesgo_cmv = self.ui.comboBox_65
        self.yd = self.ui.spinBox_27
        self.yi = self.ui.spinBox_28
        self.fd = self.ui.spinBox_29
        self.fi = self.ui.spinBox_30
        self.est_acc_vasc =self.ui.comboBox_81
        self.obs_acc_vasc = self.ui.plainTextEdit_153

        #Datos de paciente: Donante
        self.id_donante = self.ui.lineEdit_46
        self.id_donante.setValidator(QIntValidator())
        self.fecha_registro_1 = self.ui.dateEdit_22
        self.nombre_2 = self.ui.lineEdit_48
        self.sexo_2 = self.ui.comboBox_76
        self.edad_2 = self.ui.spinBox_33
        self.grupo_sanguineo_2 = self.ui.comboBox_84

        #Datos de fases: Antecedentes
        self.id_antecedente = self.ui.lineEdit_36
        self.id_antecedente.setValidator(QIntValidator()) 
        self.id_paciente_2 = self.ui.lineEdit_35
        self.id_paciente_2.setValidator(QIntValidator()) 
        self.fecha_registro_2 = self.ui.dateEdit_16
        self.ant_medicos = self.ui.plainTextEdit_183
        self.ant_quirurgicos = self.ui.plainTextEdit_184
        self.ant_traumaticos = self.ui.plainTextEdit_185
        self.ant_alergicos = self.ui.plainTextEdit_186
        self.ant_transfusionales = self.ui.plainTextEdit_187
        self.ant_ginecoobstetricos = self.ui.plainTextEdit_188

        #Datos de fases: Fase 1
        self.id_fase1 = self.ui.lineEdit_40
        self.id_fase1.setValidator(QIntValidator())  
        self.id_paciente_3 = self.ui.lineEdit_38
        self.id_paciente_3.setValidator(QIntValidator()) 
        self.fecha_registro_3 = self.ui.dateEdit_17
        self.peso = self.ui.doubleSpinBox_61
        self.talla = self.ui.doubleSpinBox_59
        self.valor_imc = self.ui.doubleSpinBox_329
        self.est_imc = self.ui.comboBox_79
        self.hema_wbc = self.ui.doubleSpinBox_361
        self.hema_neutro = self.ui.doubleSpinBox_363
        self.hema_hgb = self.ui.doubleSpinBox_365
        self.hema_hct = self.ui.doubleSpinBox_362
        self.hema_plt = self.ui.doubleSpinBox_364
        self.creatinina = self.ui.doubleSpinBox_60
        self.tfg = self.ui.doubleSpinBox_65
        self.bun = self.ui.doubleSpinBox_62
        self.gli_pre = self.ui.doubleSpinBox_64
        self.gli_post = self.ui.doubleSpinBox_54
        self.hb_glicolisada = self.ui.doubleSpinBox_55
        self.rd_longitud = self.ui.doubleSpinBox_334
        self.rd_anchura = self.ui.doubleSpinBox_332
        self.rd_grosor = self.ui.doubleSpinBox_333
        self.ri_longitud = self.ui.doubleSpinBox_337
        self.ri_anchura = self.ui.doubleSpinBox_416
        self.ri_grosor = self.ui.doubleSpinBox_335
        self.usg_hep_bil = self.ui.comboBox_129
        self.obs_usg_hep_bil = self.ui.plainTextEdit_190
        self.pra = self.ui.doubleSpinBox_336

        #Datos de fases:Fase 2a
        self.id_fase2a = self.ui.lineEdit_41
        self.id_fase2a.setValidator(QIntValidator()) 
        self.id_paciente_4 = self.ui.lineEdit_42
        self.id_paciente_4.setValidator(QIntValidator()) 
        self.fecha_registro_4 = self.ui.dateEdit_18
        self.vac_cov19 = self.ui.comboBox_208
        self.dos1_cov19 = self.ui.dateEdit_66
        self.tipo_dos1 = self.ui.comboBox_209
        self.dos2_cov19 = self.ui.dateEdit_67
        self.tipo_dos2 = self.ui.comboBox_210
        self.dos3_cov19 = self.ui.dateEdit_68
        self.tipo_dos3 = self.ui.comboBox_211
        self.vac_influen = self.ui.comboBox_212 
        self.dos_influen = self.ui.dateEdit_69
        self.vac_neum = self.ui.comboBox_213 
        self.dos_neum = self.ui.dateEdit_70
        self.cert_dental = self.ui.comboBox_218 
        self.obs_dental = self.ui.plainTextEdit_195
        self.eval_psico = self.ui.comboBox_214 
        self.obs_eval1 = self.ui.plainTextEdit_195
        self.eval_trab_social = self.ui.comboBox_215
        self.obs_eval2 = self.ui.plainTextEdit_192
        self.eval_legal = self.ui.comboBox_217
        self.obs_eval3 = self.ui.plainTextEdit_194
        self.eval_nutri = self.ui.comboBox_216
        self.obs_eval4 = self.ui.plainTextEdit_193
        self.masa_muscular = self.ui.doubleSpinBox_63
        self.grasa = self.ui.doubleSpinBox_90
        self.agua = self.ui.doubleSpinBox_91

        #Datos de fases: Fase 2b
        self.id_fase2b = self.ui.lineEdit_45
        self.id_fase2b.setValidator(QIntValidator()) 
        self.id_paciente_5 = self.ui.lineEdit_43
        self.id_paciente_5.setValidator(QIntValidator())  
        self.fecha_registro_5 = self.ui.dateEdit_20
        self.vih = self.ui.comboBox_253
        self.hepa_b = self.ui.comboBox_254
        self.hepa_c = self.ui.comboBox_255
        self.vdrl = self.ui.comboBox_300
        self.toxo_igg = self.ui.comboBox_301
        self.toxo_igm = self.ui.comboBox_305
        self.rubeo_igg = self.ui.comboBox_302
        self.rubeo_igm = self.ui.comboBox_306
        self.cmv_igg = self.ui.comboBox_303
        self.cmv_igm = self.ui.comboBox_307
        self.herpes_igg = self.ui.comboBox_304
        self.herpes_igm = self.ui.comboBox_308
        self.veb_igg = self.ui.comboBox_310
        self.veb_igm = self.ui.comboBox_309
        self.cuantiferon = self.ui.comboBox_326
        self.obs_cuantiferon = self.ui.plainTextEdit_196
        self.tgo = self.ui.doubleSpinBox_339
        self.tgp = self.ui.doubleSpinBox_338
        self.bt = self.ui.doubleSpinBox_340
        self.bd = self.ui.doubleSpinBox_341
        self.bi = self.ui.doubleSpinBox_342
        self.dhl = self.ui.doubleSpinBox_343
        self.fa = self.ui.doubleSpinBox_344
        self.albumina = self.ui.doubleSpinBox_345
        self.acido_urico = self.ui.doubleSpinBox_346
        self.perfil_hdl = self.ui.doubleSpinBox_355
        self.perfil_ldl = self.ui.doubleSpinBox_354
        self.perfil_ct = self.ui.doubleSpinBox_357
        self.perfil_tg = self.ui.doubleSpinBox_356
        self.perfil_vldl = self.ui.doubleSpinBox_358
        self.sodio = self.ui.doubleSpinBox_347
        self.potasio = self.ui.doubleSpinBox_348
        self.fosforo = self.ui.doubleSpinBox_349
        self.calcio = self.ui.doubleSpinBox_350
        self.magnesio = self.ui.doubleSpinBox_351
        self.pth = self.ui.doubleSpinBox_352
        self.tsh = self.ui.doubleSpinBox_368
        self.t4_libre = self.ui.doubleSpinBox_369
        self.c3 = self.ui.doubleSpinBox_367
        self.c4 = self.ui.doubleSpinBox_370
        self.bnp = self.ui.doubleSpinBox_405
        self.tp = self.ui.doubleSpinBox_408
        self.tpt = self.ui.doubleSpinBox_407
        self.inr = self.ui.doubleSpinBox_406
        self.nivel_ecg = self.ui.comboBox_653
        self.obs_ecg = self.ui.plainTextEdit_229
        self.anti_dna = self.ui.comboBox_654
        self.ana = self.ui.comboBox_655
        self.b2_glicoprot = self.ui.comboBox_656
        self.anticoag_lupico = self.ui.comboBox_657
        self.anticardio_igg_igm = self.ui.comboBox_658
        self.p_anca = self.ui.comboBox_659
        self.c_anca = self.ui.comboBox_660
        self.urocultivo = self.ui.comboBox_732
        self.obs_urocultivo = self.ui.plainTextEdit_261
        self.orocultivo = self.ui.comboBox_730
        self.obs_orocultivo = self.ui.plainTextEdit_262
        self.ex_nasal = self.ui.comboBox_731
        self.obs_nasal = self.ui.plainTextEdit_263

        #Datos de fases: Fase 3
        self.id_fase3 = self.ui.lineEdit_52
        self.id_fase3.setValidator(QIntValidator())  
        self.id_paciente_6 = self.ui.lineEdit_51
        self.id_paciente_6.setValidator(QIntValidator())  
        self.fecha_registro_6 = self.ui.dateEdit_23
        self.est_antig_prost = self.ui.comboBox_733
        self.obs_antig_prost = self.ui.plainTextEdit_208
        self.est_hgc_sbeta = self.ui.comboBox_735
        self.obs_hgc_sbeta = self.ui.plainTextEdit_207
        self.est_pap = self.ui.comboBox_734
        self.obs_pap = self.ui.plainTextEdit_209
        self.est_mamo = self.ui.comboBox_736
        self.obs_mamo = self.ui.plainTextEdit_212
        self.est_guay_hec = self.ui.comboBox_739
        self.obs_guay_hec = self.ui.plainTextEdit_213
        self.est_endo_colon = self.ui.comboBox_745
        self.obs_endo_colon = self.ui.plainTextEdit_214
        self.est_rxt = self.ui.comboBox_749
        self.obs_rxt = self.ui.plainTextEdit_223
        self.est_rx_spn = self.ui.comboBox_750
        self.obs_rx_spn = self.ui.plainTextEdit_215
        self.est_cisto = self.ui.comboBox_746
        self.obs_cisto =self.ui.plainTextEdit_216
        self.est_usg_vesi = self.ui.comboBox_784
        self.obs_usg_vesi = self.ui.plainTextEdit_306
        self.est_eco_trans = self.ui.comboBox_785
        self.obs_eco_trans = self.ui.plainTextEdit_307
        self.est_eco_trans_dm = self.ui.comboBox_786
        self.obs_eco_trans_dm = self.ui.plainTextEdit_308
        self.est_dop_iliac = self.ui.comboBox_787
        self.obs_dop_iliac = self.ui.plainTextEdit_310
        self.est_dop_art = self.ui.comboBox_788
        self.obs_dop_art = self.ui.plainTextEdit_311
        self.est_2donantes = self.ui.comboBox_789
        self.obs_2donantes = self.ui.plainTextEdit_309
        self.est_pielograma = self.ui.comboBox_791
        self.obs_pielograma = self.ui.plainTextEdit_313

        #Datos de fases: Fase 4
        self.id_fase4 = self.ui.lineEdit_58
        self.id_fase4.setValidator(QIntValidator()) 
        self.id_paciente_7 = self.ui.lineEdit_57
        self.id_paciente_7.setValidator(QIntValidator()) 
        self.fecha_registro_7 = self.ui.dateEdit_25
        self.eval_urologia = self.ui.plainTextEdit_210
        self.eval_cardiologia = self.ui.plainTextEdit_211
        self.angiotac_miem_inf =self.ui.plainTextEdit_252
        self.angiotac_ven_art = self.ui.plainTextEdit_251
        self.a1 = self.ui.spinBox_45
        self.a2 = self.ui.spinBox_47
        self.b1 = self.ui.spinBox_46
        self.b2 = self.ui.spinBox_48
        self.dr1 = self.ui.spinBox_49
        self.dr2 = self.ui.spinBox_50
        self.dq1 = self.ui.spinBox_52
        self.dq2 = self.ui.spinBox_51
        self.est_protocolo =self.ui.comboBox_219

        #Datos de funciones extras: Imágenes. Observar imágenes médicas para análisis por parte de los médicos
        self.id_imagen = self.ui.lineEdit_49
        self.id_paciente_8 = self.ui.lineEdit_37
        self.fecha_registro_8 = self.ui.dateEdit_19
        self.tipo_examen = self.ui.comboBox_80
        self.data_imagen = self.ui.labelimagen

        #Datos de funciones extras: Fase de compatibilidad. Unión de potencial donante para receptor
        self.id_receptor_2 = self.ui.lineEdit_39
        self.rango_edad = self.ui.comboBox_77
        self.estado_protocolo = self.ui.lineEdit_60
        self.est_acc_vascul = self.ui.lineEdit_59
        self.grupo_sanguineo_r = self.ui.lineEdit_61
        self.id_fase_final = self.ui.lineEdit_62
        self.id_fase_final.setValidator(QIntValidator()) 
        self.id_receptor_3 = self.ui.lineEdit_50
        self.id_donante_2 = self.ui.lineEdit_53
        self.fecha_registro_9 = self.ui.dateEdit_24
        self.parentesco = self.ui.comboBox_82
        self.tipo_donante = self.ui.comboBox_39
        self.hla_1 = self.ui.lineEdit_54
        self.hla_2 = self.ui.lineEdit_56
        self.pra_hla1_2 = self.ui.lineEdit_55
        self.prueba_cruzada = self.ui.comboBox_83

        #Se dejan estos widgets en modo lectura
        self.estado_protocolo.setReadOnly(True)
        self.est_acc_vascul.setReadOnly(True)
        self.grupo_sanguineo_r.setReadOnly(True)
        self.id_receptor_3.setReadOnly(True)
        self.id_donante_2.setReadOnly(True)
        self.hla_1.setReadOnly(True)
        self.hla_2.setReadOnly(True)

        #Datos de funciones extras: Evaluación periódica. Comparación de parámetros en diferentes fechas
        self.id_paciente_grafico = self.ui.lineEdit_68    

        #Parámetros a graficar de la Fase 1
        self.figure = plt.figure()
        self.canvas = FigureCanvas(self.figure)
        layout = QVBoxLayout(self.ui.widget_3)
        layout.addWidget(self.canvas)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)
        self.figure.patch.set_facecolor((230/255, 230/255, 230/255))  # Color RGB de la interfaz

        #Parámetros a graficar de la Fase 2b
        self.figure2 = plt.figure()
        self.canvas2 = FigureCanvas(self.figure2)

        layout2 = QVBoxLayout(self.ui.widget_4)
        layout2.addWidget(self.canvas2)
        layout2.setContentsMargins(0, 0, 0, 0)
        layout2.setSpacing(0)
        self.figure2.patch.set_facecolor((230/255, 230/255, 230/255))

        # Conección de radio buttons para función extra de evaluación periódica
        #Radiobuttons de Fase 1
        self.ui.radioButton_29.toggled.connect(lambda checked: checked and self.graficar_datos('hema_wbc'))
        self.ui.radioButton_30.toggled.connect(lambda checked: checked and self.graficar_datos('hema_neutro'))
        self.ui.radioButton.toggled.connect(lambda checked: checked and self.graficar_datos('hema_hgb'))
        self.ui.radioButton_14.toggled.connect(lambda checked: checked and self.graficar_datos('hema_hct'))
        self.ui.radioButton_17.toggled.connect(lambda checked: checked and self.graficar_datos('hema_plt'))
        self.ui.radioButton_18.toggled.connect(lambda checked: checked and self.graficar_datos('creatinina'))
        self.ui.radioButton_21.toggled.connect(lambda checked: checked and self.graficar_datos('tfg'))
        self.ui.radioButton_2.toggled.connect(lambda checked: checked and self.graficar_datos('bun'))
        self.ui.radioButton_9.toggled.connect(lambda checked: checked and self.graficar_datos('gli_pre'))
        self.ui.radioButton_27.toggled.connect(lambda checked: checked and self.graficar_datos('gli_post'))
        self.ui.radioButton_22.toggled.connect(lambda checked: checked and self.graficar_datos('hb_glicolisada'))

        ##Radiobuttons de Fase 1
        self.ui.radioButton_44.toggled.connect(lambda checked: checked and self.graficar_datos2('tgo'))
        self.ui.radioButton_43.toggled.connect(lambda checked: checked and self.graficar_datos2('tgp'))
        self.ui.radioButton_23.toggled.connect(lambda checked: checked and self.graficar_datos2('bt'))
        self.ui.radioButton_24.toggled.connect(lambda checked: checked and self.graficar_datos2('bd'))
        self.ui.radioButton_25.toggled.connect(lambda checked: checked and self.graficar_datos2('bi'))
        self.ui.radioButton_26.toggled.connect(lambda checked: checked and self.graficar_datos2('dhl'))
        self.ui.radioButton_31.toggled.connect(lambda checked: checked and self.graficar_datos2('fa'))
        self.ui.radioButton_32.toggled.connect(lambda checked: checked and self.graficar_datos2('albumina'))
        self.ui.radioButton_33.toggled.connect(lambda checked: checked and self.graficar_datos2('acido_urico'))
        self.ui.radioButton_34.toggled.connect(lambda checked: checked and self.graficar_datos2('perfil_hdl'))
        self.ui.radioButton_35.toggled.connect(lambda checked: checked and self.graficar_datos2('perfil_ldl'))
        self.ui.radioButton_36.toggled.connect(lambda checked: checked and self.graficar_datos2('perfil_ct'))
        self.ui.radioButton_37.toggled.connect(lambda checked: checked and self.graficar_datos2('perfil_tg'))
        self.ui.radioButton_38.toggled.connect(lambda checked: checked and self.graficar_datos2('perfil_vldl'))
        self.ui.radioButton_10.toggled.connect(lambda checked: checked and self.graficar_datos2('sodio'))
        self.ui.radioButton_11.toggled.connect(lambda checked: checked and self.graficar_datos2('potasio'))
        self.ui.radioButton_12.toggled.connect(lambda checked: checked and self.graficar_datos2('fosforo'))
        self.ui.radioButton_28.toggled.connect(lambda checked: checked and self.graficar_datos2('calcio'))
        self.ui.radioButton_39.toggled.connect(lambda checked: checked and self.graficar_datos2('magnesio'))
        self.ui.radioButton_40.toggled.connect(lambda checked: checked and self.graficar_datos2('pth'))
        self.ui.radioButton_45.toggled.connect(lambda checked: checked and self.graficar_datos2('t4_libre'))
        self.ui.radioButton_46.toggled.connect(lambda checked: checked and self.graficar_datos2('c3'))
        self.ui.radioButton_53.toggled.connect(lambda checked: checked and self.graficar_datos2('c4'))
        self.ui.radioButton_47.toggled.connect(lambda checked: checked and self.graficar_datos2('bnp'))
        self.ui.radioButton_48.toggled.connect(lambda checked: checked and self.graficar_datos2('tp'))
        self.ui.radioButton_51.toggled.connect(lambda checked: checked and self.graficar_datos2('tpt'))
        self.ui.radioButton_52.toggled.connect(lambda checked: checked and self.graficar_datos2('inr'))

        #Datos de funciones extras: Radiobuttons para seleccionar si es receptor o donante y desactivar widgets dependiendo la opción a elegir

        #Antecedentes
        self.id_receptor_ant = self.ui.radioButton_4
        self.id_donante_ant = self.ui.radioButton_3

        self.id_receptor_ant.toggled.connect(self.toggle_antecedentes_widgets)
        self.id_donante_ant.toggled.connect(self.toggle_antecedentes_widgets)

        #Fase 1
        self.id_receptor_f1= self.ui.radioButton_5
        self.id_donante_f1 = self.ui.radioButton_6

        self.id_receptor_f1.toggled.connect(self.toggle_f1_widgets)
        self.id_donante_f1.toggled.connect(self.toggle_f1_widgets)

        #Fase 2a
        self.id_receptor_f2a= self.ui.radioButton_7
        self.id_donante_f2a = self.ui.radioButton_8

        self.id_receptor_f2a.toggled.connect(self.toggle_f2a_widgets)
        self.id_donante_f2a.toggled.connect(self.toggle_f2a_widgets)

        #Fase 2b
        self.id_receptor_f2b= self.ui.radioButton_19
        self.id_donante_f2b = self.ui.radioButton_20

        self.id_receptor_f2b.toggled.connect(self.toggle_f2b_widgets)
        self.id_donante_f2b.toggled.connect(self.toggle_f2b_widgets)

        #Fase 3
        self.id_receptor_f3= self.ui.radioButton_41
        self.id_donante_f3 = self.ui.radioButton_42

        self.id_receptor_f3.toggled.connect(self.toggle_f3_widgets)
        self.id_donante_f3.toggled.connect(self.toggle_f3_widgets)

        #Imagenes
        self.id_receptor_imagen= self.ui.radioButton_49
        self.id_donante_imagen = self.ui.radioButton_50

        self.id_receptor_imagen.toggled.connect(self.buscar_imagen_condicional)

        #Fase 4
        self.id_receptor_f4= self.ui.radioButton_15
        self.id_donante_f4 = self.ui.radioButton_16

        self.id_receptor_f4.toggled.connect(self.toggle_f4_widgets)
        self.id_receptor_f4.toggled.connect(self.buscar_fase4_condicional)

        ## Mapeo de botones de acción para la gestión de datos, entre ellos: agregar, actualizar, seleccionar, buscar, limpiar y eliminar
        #  para los pacientes y sus fases

        #Botones para paciente: receptor
        self.add_btn = self.ui.agregar
        self.update_btn = self.ui.actualizar
        self.select_btn = self.ui.seleccionar
        self.search_btn = self.ui.buscar_2
        self.clear_btn = self.ui.limpiar
        self.delete_btn = self.ui.eliminar

        #Botones para paciente: donante
        self.agregar_btn_1 = self.ui.agregar_2
        self.actualizar_btn_1 = self.ui.actualizar_2
        self.seleccionar_btn_1 = self.ui.seleccionar_2
        self.buscar_btn_1 = self.ui.buscar_6
        self.limpiar_btn_1 = self.ui.limpiar_2
        self.eliminar_btn_1 = self.ui.eliminar_2

        #Botones para antecedentes
        self.agregar_btn_2 = self.ui.agregar_3
        self.actualizar_btn_2 = self.ui.actualizar_3
        self.seleccionar_btn_2 = self.ui.seleccionar_3
        self.buscar_btn_2 = self.ui.buscar_3
        self.limpiar_btn_2 = self.ui.limpiar_3
        self.eliminar_btn_2 = self.ui.eliminar_3

        #Botones para fase 1
        self.agregar_btn_3 = self.ui.agregar_4
        self.actualizar_btn_3 = self.ui.actualizar_4
        self.seleccionar_btn_3 = self.ui.seleccionar_4
        self.buscar_btn_3 = self.ui.buscar_4
        self.limpiar_btn_3 = self.ui.limpiar_4
        self.eliminar_btn_3 = self.ui.eliminar_4
        self.imc_btn_3 = self.ui.pushButton

        #Botones para fase 2a
        self.agregar_btn_4 = self.ui.agregar_5
        self.actualizar_btn_4 = self.ui.actualizar_5
        self.seleccionar_btn_4 = self.ui.seleccionar_5
        self.buscar_btn_4 = self.ui.buscar_5
        self.limpiar_btn_4 = self.ui.limpiar_5
        self.eliminar_btn_4 = self.ui.eliminar_5

        #Botones para fase 2b
        self.agregar_btn_5 = self.ui.agregar_8
        self.actualizar_btn_5 = self.ui.actualizar_8
        self.seleccionar_btn_5 = self.ui.seleccionar_8
        self.buscar_btn_5 = self.ui.buscar_8
        self.limpiar_btn_5 = self.ui.limpiar_8
        self.eliminar_btn_5 = self.ui.eliminar_8

        #Botones para fase 3
        self.agregar_btn_6 = self.ui.agregar_19
        self.actualizar_btn_6 = self.ui.actualizar_19
        self.seleccionar_btn_6 = self.ui.seleccionar_19
        self.buscar_btn_6 = self.ui.buscar_20
        self.limpiar_btn_6 = self.ui.limpiar_19
        self.eliminar_btn_6 = self.ui.eliminar_19

        #Botones para fase 4
        self.agregar_btn_7 = self.ui.agregar_13
        self.actualizar_btn_7 = self.ui.actualizar_13
        self.seleccionar_btn_7 = self.ui.seleccionar_13
        self.buscar_btn_7 = self.ui.buscar_14
        self.limpiar_btn_7 = self.ui.limpiar_13
        self.eliminar_btn_7 = self.ui.eliminar_13

        #Botones para función extra: imágenes
        self.guardar_imagen_btn = self.ui.agregar_32
        self.cargar_imagen_btn = self.ui.cargar_imagen
        self.actualizar_imagen_btn = self.ui.actualizar_32
        self.seleccionar_imagen_btn = self.ui.seleccionar_32
        self.buscar_imagen_btn = self.ui.buscar_33
        self.limpiar_imagen_btn = self.ui.limpiar_32
        self.eliminar_imagen_btn = self.ui.eliminar_32
        
        #Botón para función extra: evaluación periódica
        self.graficar_btn = self.ui.pushButton_57

        #Botones para función extra: fase de compatibilidad
        self.agregar_btn_8 = self.ui.agregar_17
        self.seleccionar_btn_8 = self.ui.Seleccionar_15
        self.buscar_btn_8 = self.ui.buscar_22
        self.limpiar_btn_8 = self.ui.limpiar_15
        self.seleccionar_registro_8 = self.ui.seleccionar_17
        self.actualizar_registro_8 = self.ui.actualizar_18
        self.eliminar_registro_8 = self.ui.eliminar_16
        #self.exportar_docx_8 = self.ui.pushButton_56

        #Configuración de tablas para visualizar los datos
        #Tabla para receptor
        self.result_table = self.ui.tableWidget_3
        self.result_table.setSortingEnabled(False)
        self.buttons_list = self.ui.function_frame.findChildren(QPushButton)

        #Tabla para donante
        self.result_table1 = self.ui.tableWidget_5
        self.result_table1.setSortingEnabled(False)
        self.buttons_list = self.ui.function_frame1.findChildren(QPushButton)

        #Tabla para antecedentes
        self.result_table2 = self.ui.tableWidget_6
        self.result_table2.setSortingEnabled(False)
        self.buttons_list2 = self.ui.function_frame2.findChildren(QPushButton)
        
        #Tabla para fase 1
        self.result_table3 = self.ui.tableWidget_7
        self.result_table3.setSortingEnabled(False)
        self.buttons_list3 = self.ui.function_frame3.findChildren(QPushButton)

        #Tabla para fase 2a
        self.result_table4 = self.ui.tableWidget_8
        self.result_table4.setSortingEnabled(False)
        self.buttons_list4 = self.ui.function_frame4.findChildren(QPushButton)

        #Tabla para fase 2b
        self.result_table5 = self.ui.tableWidget_9
        self.result_table5.setSortingEnabled(False)
        self.buttons_list5 = self.ui.function_frame5.findChildren(QPushButton)

        #Tabla para fase 3
        self.result_table6 = self.ui.tableWidget_19
        self.result_table6.setSortingEnabled(False)
        self.buttons_list6 = self.ui.function_frame6.findChildren(QPushButton)


        #Tabla para fase 4
        self.result_table7 = self.ui.tableWidget_16
        self.result_table7.setSortingEnabled(False)
        self.buttons_list7 = self.ui.function_frame4_3.findChildren(QPushButton)

        #Tablas para compatibilidad
        self.result_table8 = self.ui.tableWidget
        self.result_table8.setSortingEnabled(False)
        self.buttons_list8 = self.ui.function_frame7.findChildren(QPushButton)

        self.result_table_9 = self.ui.tableWidget_2
        self.result_table_9.setSortingEnabled(False)

        #Tabla para imágenes   
        self.result_table10 = self.ui.tableWidget_34
        self.result_table10.setSortingEnabled(False)

        #Se definen los diccionarios de mapeo para convertir valores descriptivos a códigos numéricos
        self.etnia_map = {
            "Ladino": 1,
            "Maya": 2,
            "Garífuna": 3,
            "Xinca": 4
        }

        self.sexo_map = {
            "Masculino": 1,
            "Femenino": 2
        }
        
        self.ter_sust_act_map = {
            "DPCA": 1,
            "HD": 2
        }
        
        self.inst_provee_hd_map = {
            "UNAERC": 1,
            "IGSS": 2,
            "Privado" : 3,
            "HGSJDD": 4
        }

        self.grupo_sanguineo_map = {
            "A+": 1,
            "B+": 2,
            "O+": 3,
            "AB+": 4,
            "A-": 5,
            "B-": 6,
            "O-": 7,
            "AB-": 8  
        }

        self.ocupacion_map = {
            "Agricultor": 1,
            "Construcción": 2,
            "Oficina": 3,
            "Ama de casa": 4,
            "Desempleado": 5
        }

        self.etiologia_erc_map = {
            "Extrarrenal": 1,
            "Renal": 2,
            "Prerrenal": 3,
            "Desconocido": 4
        }

        self.procedencia_map = {
            "Alta Verapaz": 1,
            "Baja Verapaz": 2,
            "Chimaltenango": 3,
            "Chiquimula": 4,
            "El Progreso": 5,
            "Escuintla": 6,
            "Guatemala": 7,
            "Huehuetenango": 8,
            "Izabal": 9,
            "Jalapa": 10,
            "Jutiapa": 11,
            "Petén": 12,
            "Quetzaltenango": 13,
            "Quiché": 14,
            "Retalhuleu": 15,
            "Sacatepéquez": 16,
            "San Marcos": 17,
            "Santa Rosa": 18,
            "Sololá": 19,
            "Suchitepéquez": 20,
            "Totonicapán": 21,
            "Zacapa": 22
        }

        self.residencia_map = {
            "Alta Verapaz": 1,
            "Baja Verapaz": 2,
            "Chimaltenango": 3,
            "Chiquimula": 4,
            "El Progreso": 5,
            "Escuintla": 6,
            "Guatemala": 7,
            "Huehuetenango": 8,
            "Izabal": 9,
            "Jalapa": 10,
            "Jutiapa": 11,
            "Petén": 12,
            "Quetzaltenango": 13,
            "Quiché": 14,
            "Retalhuleu": 15,
            "Sacatepéquez": 16,
            "San Marcos": 17,
            "Santa Rosa": 18,
            "Sololá": 19,
            "Suchitepéquez": 20,
            "Totonicapán": 21,
            "Zacapa": 22
        }

        self.riesgo_cmv_map = {
            "Alto": 1,
            "Intermedio": 2,
            "Bajo": 3
        }
        
        self.est_acc_vasc_map = {
            "Agotados": 1,
            "No agotados": 2
        }

        #Fase 1
        self.est_imc_map = {
            "Bajo peso": 1,
            "Normal": 2,
            "Sobrepeso": 3,
            "Obesidad": 4,
            "Obesidad mórbida": 5
        }

        self.usg_hep_bil_map = {
            "Normal": 1,
            "Anormal" : 2
        }

        #Fase 2a
        self.estado_vacuna_map = {
            "Vacunado": 1,
            "No vacunado" : 2,
            "Pendiente" : 3
        }

        self.tipo_vacuna_map = {
            "Pfizer": 1,
            "Moderna" : 2,
            "Astrazeneca" : 3,
            "Janssen" : 4
        }

        self.estado_eval_map = {
            "Aprobado": 1,
            "No aprobado" : 2,
            "Al ingreso" : 3,
            "Pendiente" : 4
        }

        #Fase 2b
        self.estado_pruebas = {
            "Positivo": 1,
            "Negativo" : 2,
            "No reactivo" : 3,
            "NHR" : 4,
            "Estéril": 5,
            "S. aureus" : 6,
            "No aplica" : 7,
            "No crecimiento" : 8,
            "Microbiota normal": 9,
            "Anuria" : 10,
            "S. epidermis" : 11
        }

        self.nivel_ecg_map = {
            "Normal": 1,
            "Anormal" : 2
        }

        self.estado_pruebas1 = {
            "+": 1,
            "-" : 2
        }

        #Fase 3
        self.estado_pruebas2 = {
            "Positivo": 1,
            "Negativo" : 2,
            "Normal" : 3,
            "Anormal" : 4,
            "No aplica": 5,
            "Completado" : 6,
            "No ompletado" : 7,
            "Si" : 8,
            "No": 9
        }

        #Fase 4
        self.est_protocolo_map = {
            "Completado": 1,
            "No completado": 2
        }

        #Imagenes
        self.tipo_imagen = {
            "RXT": 1,
            "RX SPN": 2,
            "Cistograma": 3,
            "ECO": 4,
            "Pielograma": 5,
            "AngioTAC": 6
        }

        #Fase de compatibilidad
        self.parentesco_map = {
            "Madre": 1,
            "Padre": 2,
            "Hermano": 3,
            "Hermana": 4,
            "Hijo": 5,
            "Hija": 6,
            "Abuelo": 7,
            "Abuela": 8 ,
            "Amigo": 9,
            "Amiga": 10,
            "Cónyuge": 11,
            "Primo": 12 ,
            "Prima": 13
        }

        self.est_acc_vascul_map = {
            "Agotados": 1,
            "No agotados": 2
        }

        self.prueba_cruzada_map = {
            "Positivo": 1,
            "Negativo": 2
        }

        self.rango_edad_map = {
            "Adulto joven (20 - 39) años": 1,
            "Adulto medio (40-59) años": 2,
            "Adulto mayor (60+) años": 3
        }

        self.tipo_donante_map = {
            "Vivo relacionado": 1,
            "Vivo no relacionado": 2,
            "Cadavérico": 3,
            "Altruista": 4
        }

        # Se inicializa conexiones entre señales y sus funciones correspondientes
        self.init_signal_slot()

        # Se cargar datos iniciales en las tablas
        self.buscar_receptor()
        self.buscar_donante()
        self.buscar_registros()

        # Se conectan los botones de navegación para las distintas secciones de la interfaz
        self.ui.registros_2.clicked.connect(self.ir_a_pag_registros_2)
        self.ui.evaluacion_2.clicked.connect(self.ir_a_pag_evaluacion_2)
        self.ui.compatibilidad_2.clicked.connect(self.ir_a_pag_compatibilidad_2)

    @pyqtSlot()
    #Se definen las funciones que están conectadas a los botones de navegación por la interfaz
    def ir_a_pag_registros_2(self):
        self.ui.stackedWidget.setCurrentWidget(self.ui.pag_registros)
    def ir_a_pag_evaluacion_2(self):
        self.ui.stackedWidget.setCurrentWidget(self.ui.pag_evaluacion)  
    def ir_a_pag_compatibilidad_2(self):
        self.ui.stackedWidget.setCurrentWidget(self.ui.pag_compatibilidad)  

    #Se conecta las señales 'clicked' de cada botón con su correspondiente función (Agregar, buscar, limpiar, seleccionar, actualizar y eliminar)
    def init_signal_slot(self):   

        # Botones para datos personales de receptor
        self.add_btn.clicked.connect(self.agregar_receptor)
        self.search_btn.clicked.connect(self.buscar_receptor)
        self.clear_btn.clicked.connect(self.limpiar_receptor)
        self.select_btn.clicked.connect(self.seleccionar_receptor)
        self.update_btn.clicked.connect(self.actualizar_receptor)
        self.delete_btn.clicked.connect(self.eliminar_receptor)

        # Botones para datos personales de donante
        self.agregar_btn_1.clicked.connect(self.agregar_donante)
        self.buscar_btn_1.clicked.connect(self.buscar_donante)
        self.limpiar_btn_1.clicked.connect(self.limpiar_donante)
        self.seleccionar_btn_1.clicked.connect(self.seleccionar_donante)
        self.actualizar_btn_1.clicked.connect(self.actualizar_donante)
        self.eliminar_btn_1.clicked.connect(self.eliminar_donante)

        #Botones para antecedentes
        self.agregar_btn_2.clicked.connect(self.agregar_antecedentes)
        self.buscar_btn_2.clicked.connect(self.buscar_antecedentes_condicional)
        self.limpiar_btn_2.clicked.connect(self.limpiar_antecedentes)
        self.seleccionar_btn_2.clicked.connect(self.seleccionar_antecedentes)
        self.actualizar_btn_2.clicked.connect(self.actualizar_antecedentes)
        self.eliminar_btn_2.clicked.connect(self.eliminar_antecedentes)

        #Fase 1
        self.agregar_btn_3.clicked.connect(self.agregar_fase1)
        self.buscar_btn_3.clicked.connect(self.buscar_fase1_condicional)
        self.limpiar_btn_3.clicked.connect(self.limpiar_fase1)
        self.seleccionar_btn_3.clicked.connect(self.seleccionar_fase1)
        self.actualizar_btn_3.clicked.connect(self.actualizar_fase1)
        self.eliminar_btn_3.clicked.connect(self.eliminar_fase1)
        self.imc_btn_3.clicked.connect(self.calcular_imc)

        #Fase 2a
        self.agregar_btn_4.clicked.connect(self.agregar_fase2a)
        self.buscar_btn_4.clicked.connect(self.buscar_fase2a_condicional)
        self.limpiar_btn_4.clicked.connect(self.limpiar_fase2a)
        self.seleccionar_btn_4.clicked.connect(self.seleccionar_fase2a)
        self.actualizar_btn_4.clicked.connect(self.actualizar_fase2a)
        self.eliminar_btn_4.clicked.connect(self.eliminar_fase2a)
        
        #Fase 2b
        self.agregar_btn_5.clicked.connect(self.agregar_fase2b)
        self.buscar_btn_5.clicked.connect(self.buscar_fase2b_condicional)
        self.limpiar_btn_5.clicked.connect(self.limpiar_fase2b)
        self.seleccionar_btn_5.clicked.connect(self.seleccionar_fase2b)
        self.actualizar_btn_5.clicked.connect(self.actualizar_fase2b)
        self.eliminar_btn_5.clicked.connect(self.eliminar_fase2b)

        #Fase 3
        self.agregar_btn_6.clicked.connect(self.agregar_fase3)
        self.buscar_btn_6.clicked.connect(self.buscar_fase3_condicional)
        self.limpiar_btn_6.clicked.connect(self.limpiar_fase3)
        self.seleccionar_btn_6.clicked.connect(self.seleccionar_fase3)
        self.actualizar_btn_6.clicked.connect(self.actualizar_fase3)
        self.eliminar_btn_6.clicked.connect(self.eliminar_fase3) 

        #Fase 4
        self.agregar_btn_7.clicked.connect(self.agregar_fase4)
        self.buscar_btn_7.clicked.connect(self.buscar_fase4_condicional)
        self.limpiar_btn_7.clicked.connect(self.limpiar_fase4)
        self.seleccionar_btn_7.clicked.connect(self.seleccionar_fase4)
        self.actualizar_btn_7.clicked.connect(self.actualizar_fase4)
        self.eliminar_btn_7.clicked.connect(self.eliminar_fase4) 

        #Imágenes
        # Conectar botones a funciones
        self.guardar_imagen_btn.clicked.connect(self.guardar_imagen)
        self.cargar_imagen_btn.clicked.connect(self.cargar_imagen)
        self.actualizar_imagen_btn.clicked.connect(self.actualizar_imagen)
        self.seleccionar_imagen_btn.clicked.connect(self.seleccionar_imagen)
        self.buscar_imagen_btn.clicked.connect(self.buscar_imagen_condicional)
        self.limpiar_imagen_btn.clicked.connect(self.limpiar_imagen)
        self.eliminar_imagen_btn.clicked.connect(self.eliminar_imagen)

        #Fase de compatibilidad
        self.agregar_btn_8.clicked.connect(self.agregar_compatibilidad)
        self.buscar_btn_8.clicked.connect(self.combinar_buscar_mostrar)
        self.limpiar_btn_8.clicked.connect(self.limpiar_compatibilidad)
        self.seleccionar_btn_8.clicked.connect(self.seleccionar_compatibilidad)
        self.seleccionar_registro_8.clicked.connect(self.seleccionar_registro)
        self.actualizar_registro_8.clicked.connect(self.actualizar_registro)
        self.eliminar_registro_8.clicked.connect(self.eliminar_registro) 
#        self.exportar_docx_8.clicked.connect(self.exportar_docx)

        #Fase de graficos
        self.graficar_btn.clicked.connect(self.habilitar_graficos)
   
    def combinar_buscar_mostrar(self):
        # Llama a ambas funciones cuando esta clickeado una opción 
        self.buscar_compatibilidad()
        self.mostrar_datos_receptor()

    def desactivar_botones(self):
        # Desactiva todos los botones
        for button in self.buttons_list:
            button.setDisabled(True)

    def activar_botones(self):
        # Habilita todos los botones
        for button in self.buttons_list:
            button.setDisabled(False)

####################################Receptor - Funciones básicas para interactuar con los datos (buscar,agregar,limpiar,eliminar datos)

    def agregar_receptor(self):
        #Agrega un nuevo receptor después de validar la contraseña
        try:
            #Solicita y valida la contraseña
            password = self.solicitar_contraseña()
            if password is None:
                return
                
            if not self.validar_contraseña(password):
                QtWidgets.QMessageBox.warning(self, "Error", "Contraseña incorrecta")
                return
            
            # Recoge datos de los campos de entrada
            id_receptor = self.id_receptor.text()
            fecha_registro = self.fecha_registro.date().toString("yyyy-MM-dd")
            nombre = self.nombre.text()
            edad = self.edad.value()
            fecha_dg_erc = self.fecha_dg_erc.date().toString("yyyy-MM-dd")
            vol_residual = self.vol_residual.value()
            tiempo_anuria = self.tiempo_anuria.date().toString("yyyy-MM-dd")
            tiempo_ini_st_renal = self.tiempo_ini_st_renal.date().toString("yyyy-MM-dd")
            tiempo_dialisis = self.tiempo_dialisis.date().toString("yyyy-MM-dd")
            yd = self.yd.value()
            yi = self.yi.value()
            fd = self.fd.value()
            fi = self.fi.value()
            obs_acc_vasc = self.obs_acc_vasc.toPlainText()

            # Obtiene el valor entero correspondiente a las opciones seleccionadas
            etnia = self.etnia_map.get(self.etnia.currentText(), None)
            sexo = self.sexo_map.get(self.sexo.currentText(), None)
            ter_sust_act = self.ter_sust_act_map.get(self.ter_sust_act.currentText(), None)
            inst_provee_hd = self.inst_provee_hd_map.get(self.inst_provee_hd.currentText(), None)
            grupo_sanguineo = self.grupo_sanguineo_map.get(self.grupo_sanguineo.currentText(), None)
            ocupacion = self.ocupacion_map.get(self.ocupacion.currentText(), None)
            etiologia_erc = self.etiologia_erc_map.get(self.etiologia_erc.currentText(), None)
            procedencia = self.procedencia_map.get(self.procedencia.currentText(), None)
            residencia = self.residencia_map.get(self.residencia.currentText(), None)
            riesgo_cmv = self.riesgo_cmv_map.get(self.riesgo_cmv.currentText(), None)
            est_acc_vasc = self.est_acc_vasc_map.get(self.est_acc_vasc.currentText(), None)

            # Llama al método de la clase de conexión para agregar el paciente
            if self.db.agregar_datos_0(id_receptor, fecha_registro, nombre, edad, etnia, sexo, fecha_dg_erc, ter_sust_act, inst_provee_hd, vol_residual, tiempo_anuria, grupo_sanguineo,
                                    procedencia, residencia, ocupacion, etiologia_erc, tiempo_ini_st_renal, tiempo_dialisis, riesgo_cmv, yd, yi, fd, fi, est_acc_vasc, obs_acc_vasc):
                
                # Actualiza solo id_receptor y fecha_registro en el QTableWidget
                self.result_table.setRowCount(0)  # Limpia la tabla
                row_position = self.result_table.rowCount()
                self.result_table.insertRow(row_position)
                self.result_table.setItem(row_position, 0, QTableWidgetItem(id_receptor))
                self.result_table.setItem(row_position, 1, QTableWidgetItem(fecha_registro))

                QMessageBox.information(self, "Éxito", "Información agregada correctamente.")
                self.limpiar_receptor()  # Limpia el formulario después de agregar
            else:
                QtWidgets.QMessageBox.warning(self, "Error", "No se pudo agregar la información.")
            
        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Error", f"Error al agregar receptor: {str(e)}")

    def buscar_receptor(self):
        results = self.db.buscar_datos_0()  # Llama al método de la clase de conexión

        # Limpia la tabla antes de agregar nuevos datos
        self.result_table.setRowCount(0)

        for row in results:
            row_position = self.result_table.rowCount()
            self.result_table.insertRow(row_position)
            # Se asegura de que los nombres de las columnas coincidan con los de la tabla
            for column, value in enumerate(row.values()):
                self.result_table.setItem(row_position, column, QTableWidgetItem(str(value))) 


    def limpiar_receptor(self):
        # Limpia los campos de texto
        self.id_receptor.clear()
        self.nombre.clear() # Limpia la celda de nombre
        self.edad.setValue(0)  # Restablece el valor de la edad a 0
        self.etnia.setCurrentIndex(0)  # Restablece el índice del comboBox de etnia al primero
        self.sexo.setCurrentIndex(0)  
        fecha_default = QDate(2000, 1, 1)
        self.fecha_registro.setDate(fecha_default) # Restablece a una fecha ya definida
        self.fecha_dg_erc.setDate(fecha_default)
        self.ter_sust_act.setCurrentIndex(0)
        self.inst_provee_hd.setCurrentIndex(0)
        self.vol_residual.setValue(0)
        self.tiempo_anuria.setDate(fecha_default)
        self.grupo_sanguineo.setCurrentIndex(0)
        self.procedencia.setCurrentIndex(0)
        self.residencia.setCurrentIndex(0)
        self.ocupacion.setCurrentIndex(0)
        self.etiologia_erc.setCurrentIndex(0)
        self.tiempo_ini_st_renal.setDate(fecha_default)
        self.tiempo_dialisis.setDate(fecha_default)
        self.riesgo_cmv.setCurrentIndex(0)
        self.yd.setValue(0)
        self.yi.setValue(0)
        self.fd.setValue(0)
        self.fi.setValue(0)
        self.obs_acc_vasc.clear()
        self.est_acc_vasc.setCurrentIndex(0)

    def seleccionar_receptor(self):

        selected_row = self.result_table.currentRow()
        if selected_row >= 0:
            id_receptor = self.result_table.item(selected_row, 0).text()
            
            # Obtiene todos los datos del receptor seleccionado de la base de datos
            receptor_data = self.db.buscar_datos_01(id_receptor)
            
            #Instrucciones para trabajar con datos de la base de datos que se encuentren nulos
            if receptor_data:
                # Función auxiliar para obtener valores float o establecerlos en 0
                def get_float_value(value):
                    try:
                        return float(value) if value not in [None, '', 'None'] else 0.0
                    except ValueError:
                        return 0.0

                # Función auxiliar para obtener valores int o establecerlos en 0
                def get_int_value(value):
                    try:
                        return int(value) if value not in [None, '', 'None'] else 0
                    except ValueError:
                        return 0
                self.limpiar_receptor()

                # Se llenan los widgets con la información del id seleccionado
                self.id_receptor.setText(str(receptor_data['id_receptor']))
                self.fecha_registro.setDate(QDate.fromString(str(receptor_data['fecha_registro']), "yyyy-MM-dd"))
                self.nombre.setText(receptor_data['nombre'])
                self.edad.setValue(get_int_value(receptor_data['edad']))

                etnia_value = get_int_value(receptor_data['etnia'])
                if etnia_value in self.etnia_map.values():
                    self.etnia.setCurrentText(next(key for key, value in self.etnia_map.items() if value == etnia_value))

                sexo_value = get_int_value(receptor_data['sexo'])
                if sexo_value in self.sexo_map.values():
                    self.sexo.setCurrentText(next(key for key, value in self.sexo_map.items() if value == sexo_value))

                self.fecha_dg_erc.setDate(QDate.fromString(str(receptor_data['fecha_dg_erc']), "yyyy-MM-dd"))

                ter_sust_act_value = get_int_value(receptor_data['ter_sust_act'])
                if ter_sust_act_value in self.ter_sust_act_map.values():
                    self.ter_sust_act.setCurrentText(next(key for key, value in self.ter_sust_act_map.items() if value == ter_sust_act_value))

                inst_provee_hd_value = get_int_value(receptor_data['inst_provee_hd'])
                if inst_provee_hd_value in self.inst_provee_hd_map.values():
                    self.inst_provee_hd.setCurrentText(next(key for key, value in self.inst_provee_hd_map.items() if value == inst_provee_hd_value))

                self.vol_residual.setValue(get_float_value(receptor_data['vol_residual']))
                self.tiempo_anuria.setDate(QDate.fromString(str(receptor_data['tiempo_anuria']), "yyyy-MM-dd"))

                grupo_sanguineo_value = get_int_value(receptor_data['grupo_sanguineo'])
                if grupo_sanguineo_value in self.grupo_sanguineo_map.values():
                    self.grupo_sanguineo.setCurrentText(next(key for key, value in self.grupo_sanguineo_map.items() if value == grupo_sanguineo_value))

                procedencia_value = get_int_value(receptor_data['procedencia'])
                if procedencia_value in self.procedencia_map.values():
                    self.procedencia.setCurrentText(next(key for key, value in self.procedencia_map.items() if value == procedencia_value))

                residencia_value = get_int_value(receptor_data['residencia'])
                if residencia_value in self.residencia_map.values():
                    self.residencia.setCurrentText(next(key for key, value in self.residencia_map.items() if value == residencia_value))

                ocupacion_value = get_int_value(receptor_data['ocupacion'])
                if ocupacion_value in self.ocupacion_map.values():
                    self.ocupacion.setCurrentText(next(key for key, value in self.ocupacion_map.items() if value == ocupacion_value))

                etiologia_erc_value = get_int_value(receptor_data['etiologia_erc'])
                if etiologia_erc_value in self.etiologia_erc_map.values():
                    self.etiologia_erc.setCurrentText(next(key for key, value in self.etiologia_erc_map.items() if value == etiologia_erc_value))

                self.tiempo_ini_st_renal.setDate(QDate.fromString(str(receptor_data['tiempo_ini_st_renal']), "yyyy-MM-dd"))
                self.tiempo_dialisis.setDate(QDate.fromString(str(receptor_data['tiempo_dialisis']), "yyyy-MM-dd"))

                riesgo_cmv_value = get_int_value(receptor_data['riesgo_cmv'])
                if riesgo_cmv_value in self.riesgo_cmv_map.values():
                    self.riesgo_cmv.setCurrentText(next(key for key, value in self.riesgo_cmv_map.items() if value == riesgo_cmv_value))

                self.yd.setValue(get_int_value(receptor_data['yd']))
                self.yi.setValue(get_int_value(receptor_data['yi']))
                self.fd.setValue(get_int_value(receptor_data['fd']))
                self.fi.setValue(get_int_value(receptor_data['fi']))

                est_acc_vasc_value = get_int_value(receptor_data['est_acc_vasc'])
                if est_acc_vasc_value in self.est_acc_vasc_map.values():
                    self.est_acc_vasc.setCurrentText(next(key for key, value in self.est_acc_vasc_map.items() if value == est_acc_vasc_value))
                
                self.obs_acc_vasc.setPlainText(str(receptor_data['obs_acc_vasc']))

            else:
                QMessageBox.warning(self, "Error", "No se pudo obtener la información del receptor.")

    def actualizar_receptor(self):
        # Actualiza los datos del receptor después de validar la contraseña
        # Solicita y valida la contraseña
        contraseña = self.solicitar_contraseña()
        if contraseña is None:
            return
            
        if not self.validar_contraseña(contraseña):
            QtWidgets.QMessageBox.warning(self, "Error", "Contraseña incorrecta")
            return
        
        # Si la contraseña es válida, procede con la actualización
        try:
        
            # Recoge datos de los campos de entrada
            id_receptor = self.id_receptor.text() or None
            fecha_registro = self.fecha_registro.date().toString("yyyy-MM-dd")  # Formato de fecha
            nombre = self.nombre.text() or None
            edad = self.edad.value() or 0
            fecha_dg_erc = self.fecha_dg_erc.date().toString("yyyy-MM-dd")  # Formato de fecha
            vol_residual = self.vol_residual.value() or 0.0
            tiempo_anuria = self.tiempo_anuria.date().toString("yyyy-MM-dd")
            tiempo_ini_st_renal = self.tiempo_ini_st_renal.date().toString("yyyy-MM-dd")  # Formato de fecha
            tiempo_dialisis = self.tiempo_dialisis.date().toString("yyyy-MM-dd")
            yd = self.yd.value() or 0
            yi = self.yi.value() or 0
            fd = self.fd.value() or 0
            fi = self.fi.value() or 0
            obs_acc_vasc = self.obs_acc_vasc.toPlainText() or None

            # Obtiene el valor entero correspondiente a la etnia y sexo seleccionados
            etnia = self.etnia_map.get(self.etnia.currentText(), None)
            sexo = self.sexo_map.get(self.sexo.currentText(), None)
            ter_sust_act = self.ter_sust_act_map.get(self.ter_sust_act.currentText(), None)
            inst_provee_hd = self.inst_provee_hd_map.get(self.inst_provee_hd.currentText(), None)
            grupo_sanguineo = self.grupo_sanguineo_map.get(self.grupo_sanguineo.currentText(), None)
            ocupacion = self.ocupacion_map.get(self.ocupacion.currentText(), None)
            etiologia_erc = self.etiologia_erc_map.get(self.etiologia_erc.currentText(), None)
            procedencia = self.procedencia_map.get(self.procedencia.currentText(), None)
            residencia = self.residencia_map.get(self.residencia.currentText(), None)
            riesgo_cmv = self.riesgo_cmv_map.get(self.riesgo_cmv.currentText(), None)
            est_acc_vasc = self.est_acc_vasc_map.get(self.est_acc_vasc.currentText(), None)

            # Llama al método de la clase de conexión para actualizar el paciente
            if self.db.actualizar_datos_0(id_receptor, fecha_registro, nombre, edad, etnia, sexo, fecha_dg_erc,ter_sust_act,inst_provee_hd, vol_residual, tiempo_anuria, grupo_sanguineo,
                                    procedencia, residencia, ocupacion, etiologia_erc, tiempo_ini_st_renal, tiempo_dialisis, riesgo_cmv, yd, yi, fd, fi, est_acc_vasc, obs_acc_vasc):
                QMessageBox.information(self, "Éxito", "Información actualizada correctamente.")
                self.limpiar_receptor()  # Limpia el formulario después de actualizar
                self.buscar_receptor()  # Vuelve a cargar los datos para mostrar todos
            pass
        
        except Exception as e:
            QtWidgets.QMessageBox.warning(self, "Error", f"Error al actualizar receptor: {e}")

    def eliminar_receptor(self):
        #Elimina un receptor después de validar la contraseña
        try:
            # Solicita y valida la contraseña
            contraseña = self.solicitar_contraseña()
            if contraseña is None:
                return
                
            if not self.validar_contraseña(contraseña):
                QtWidgets.QMessageBox.warning(self, "Error", "Contraseña incorrecta")
                return
            
            # Si la contraseña es válida, procede con eliminar
            selected_row = self.result_table.currentRow()
            
            if selected_row >= 0:
                id_receptor = self.result_table.item(selected_row, 0).text()
                
                reply = QtWidgets.QMessageBox.question(
                    self, 
                    'Confirmar eliminación', 
                    '¿Está seguro de que desea eliminar este registro?',
                    QtWidgets.QMessageBox.Yes | QtWidgets.QMessageBox.No, 
                    QtWidgets.QMessageBox.No
                )
                
                if reply == QtWidgets.QMessageBox.Yes:
                    if self.db.eliminar_datos_0(id_receptor):
                        self.result_table.removeRow(selected_row)
                        QtWidgets.QMessageBox.information(self, "Éxito", "Registro eliminado correctamente.")
                        self.limpiar_receptor()
                    else:
                        QtWidgets.QMessageBox.warning(self, "Error", "No se pudo eliminar el registro.")
            else:
                QtWidgets.QMessageBox.warning(self, "Advertencia", "Por favor, seleccione un registro para eliminar.")
                
        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Error", f"Error al eliminar receptor: {str(e)}")
    
 ################## DONANTE - Funciones básicas para interactuar con los datos (buscar,agregar,limpiar,eliminar datos)

    def agregar_donante(self):
        #Agrega un nuevo donante después de validar la contraseña
        try:
            # Solicita y valida la contraseña
            contraseña = self.solicitar_contraseña()
            if contraseña is None:
                return
                
            if not self.validar_contraseña(contraseña):
                QtWidgets.QMessageBox.warning(self, "Error", "Contraseña incorrecta")
                return
            # Recoge datos de los campos de entrada
            id_donante = self.id_donante.text()
            fecha_registro_1 = self.fecha_registro_1.date().toString("yyyy-MM-dd")  # Formato de fecha
            nombre_2 = self.nombre_2.text()
            edad_2 = self.edad_2.value()

            if self.sexo_2.currentText() == "" or self.grupo_sanguineo_2.currentText() == "" :
                QMessageBox.warning(self, "Advertencia", "Debe seleccionar una opción para Sexo, Grupo Sanguíneo y Tipo de Donante.")
                return

            # Obtiene el valor entero correspondiente a las opciones seleccionadas
            sexo_2 = self.sexo_map[self.sexo_2.currentText()]
            grupo_sanguineo_2 = self.grupo_sanguineo_map[self.grupo_sanguineo_2.currentText()]

            # Valida que todos los campos estén llenos
            if not all([id_donante, fecha_registro_1, nombre_2, sexo_2, edad_2, grupo_sanguineo_2]):
                QMessageBox.warning(self, "Advertencia", "Todos los campos son obligatorios.")
                return

            # Llama al método de la clase de conexión para agregar el donante
            if self.db.agregar_datos_1(id_donante, fecha_registro_1, nombre_2, sexo_2, edad_2, grupo_sanguineo_2):
                # Limpia la tabla antes de agregar el nuevo registro
                self.result_table1.setRowCount(0)  # Limpia la tabla

                # Actualiza solo la fila correspondiente en el QTableWidget con las dos primeras variables
                row_position = self.result_table1.rowCount()
                self.result_table1.insertRow(row_position)
                self.result_table1.setItem(row_position, 0, QTableWidgetItem(id_donante))
                self.result_table1.setItem(row_position, 1, QTableWidgetItem(fecha_registro_1))

                QMessageBox.information(self, "Éxito", "Información del donante agregada correctamente.")
                self.limpiar_donante()  # Limpia el formulario después de agregar

            else:
                QtWidgets.QMessageBox.warning(self, "Error", "No se pudo agregar la información.")
            
        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Error", f"Error al agregar donante: {str(e)}")

    def buscar_donante(self):
            results = self.db.buscar_datos_1()  # Llama al método de la clase de conexión

            # Limpia la tabla antes de agregar nuevos datos
            self.result_table1.setRowCount(0)

            for row in results:
                row_position = self.result_table1.rowCount()
                self.result_table1.insertRow(row_position)
                # Se asegura de que los nombres de las columnas coincidan con los de la tabla
                for column, value in enumerate(row.values()):
                    self.result_table1.setItem(row_position, column, QTableWidgetItem(str(value))) 


    def limpiar_donante(self):
        # Limpia los campos de texto
        self.id_donante.clear()
        fecha_default = QDate(2000, 1, 1)
        self.fecha_registro_1.setDate(fecha_default)
        self.nombre_2.clear()
        self.edad_2.setValue(0)  # Restablece el valor de la edad a 0
        self.sexo_2.setCurrentIndex(0)  # Restablece el índice del comboBox de sexo al primero
        self.grupo_sanguineo_2.setCurrentIndex(0)

    def seleccionar_donante(self):
        selected_row = self.result_table1.currentRow()
        if selected_row >= 0:
            id_donante = self.result_table1.item(selected_row, 0).text()
            
            #Obtiene todos los datos del donante seleccionado de la base de datos
            donante_data = self.db.buscar_datos_1_2(id_donante)
            
            if donante_data:
                # Función auxiliar para obtener valores int o establecerlos en 0
                def get_int_value(value):
                    try:
                        return int(value) if value not in [None, '', 'None'] else 0
                    except ValueError:
                        return 0
                self.limpiar_donante()

                # Llena los widgets con la información
                self.id_donante.setText(str(donante_data['id_donante']))
                self.fecha_registro_1.setDate(QDate.fromString(str(donante_data['fecha_registro']), "yyyy-MM-dd"))
                self.nombre_2.setText(donante_data['nombre'])

                sexo_2_value = get_int_value(donante_data['sexo'])
                if sexo_2_value in self.sexo_map.values():
                    self.sexo_2.setCurrentText(next(key for key, value in self.sexo_map.items() if value == sexo_2_value))

                self.edad_2.setValue(get_int_value(donante_data['edad']))

                grupo_sanguineo_2_value = get_int_value(donante_data['grupo_sanguineo'])
                if grupo_sanguineo_2_value in self.grupo_sanguineo_map.values():
                    self.grupo_sanguineo_2.setCurrentText(next(key for key, value in self.grupo_sanguineo_map.items() if value == grupo_sanguineo_2_value))

            else:
                QMessageBox.warning(self, "Error", "No se pudo obtener la información del donante.")

    def actualizar_donante(self):
        #Agrega un nuevo donante después de validar la contraseña
        try:
            # Solicita y valida la contraseña
            contraseña = self.solicitar_contraseña()
            if contraseña is None:
                return
                
            if not self.validar_contraseña(contraseña):
                QtWidgets.QMessageBox.warning(self, "Error", "Contraseña incorrecta")
                return

            # Recoger datos de los campos de entrada
            id_donante = self.id_donante.text()
            fecha_registro_1 = self.fecha_registro_1.date().toString("yyyy-MM-dd")  # Formato de fecha
            nombre_2 = self.nombre_2.text()
            edad_2 = self.edad_2.value()

            # Obtiene el valor entero correspondiente a la etnia y sexo seleccionados
            sexo_2 = self.sexo_map.get(self.sexo_2.currentText())
            grupo_sanguineo_2 = self.grupo_sanguineo_map.get(self.grupo_sanguineo_2.currentText())

            # Llama al método de la clase de conexión para actualizar el paciente
            if self.db.actualizar_datos_1(id_donante, fecha_registro_1, nombre_2, sexo_2, edad_2, grupo_sanguineo_2):
                QMessageBox.information(self, "Éxito", "Información actualizada correctamente.")
                self.limpiar_donante()  # Limpia el formulario después de actualizar
                self.buscar_donante()  # Vuelve cargar los datos para mostrar todos
        
        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Error", f"Error al actualizar donante: {str(e)}")

    def eliminar_donante(self):
        #Elimina un donante después de validar la contraseña
        try:
            #Solicita y valida la contraseña
            contraseña = self.solicitar_contraseña()
            if contraseña is None:
                return
                
            if not self.validar_contraseña(contraseña):
                QtWidgets.QMessageBox.warning(self, "Error", "Contraseña incorrecta")
                return
            
            # Si la contraseña es válida, procede con eliminar
            selected_row = self.result_table1.currentRow()
            
            if selected_row >= 0:
                id_donante = self.result_table1.item(selected_row, 0).text()
                
                reply = QtWidgets.QMessageBox.question(
                    self, 
                    'Confirmar eliminación', 
                    '¿Está seguro de que desea eliminar este registro?',
                    QtWidgets.QMessageBox.Yes | QtWidgets.QMessageBox.No, 
                    QtWidgets.QMessageBox.No
                )
                
                if reply == QtWidgets.QMessageBox.Yes:
                    if self.db.eliminar_datos_1(id_donante):
                        self.result_table1.removeRow(selected_row)
                        QtWidgets.QMessageBox.information(self, "Éxito", "Registro eliminado correctamente.")
                        self.limpiar_donante()
                    else:
                        QtWidgets.QMessageBox.warning(self, "Error", "No se pudo eliminar el registro.")
            else:
                QtWidgets.QMessageBox.warning(self, "Advertencia", "Por favor, seleccione un registro para eliminar.")
                
        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Error", f"Error al eliminar donante: {str(e)}")


    ################## ANTECEDENTES - Funciones básicas para interactuar con los datos (buscar,agregar,limpiar,eliminar datos)
    def agregar_antecedentes(self):
        try:
            # Solicita y valida la contraseña
            contraseña = self.solicitar_contraseña()
            if contraseña is None:
                return
                
            if not self.validar_contraseña(contraseña):
                QtWidgets.QMessageBox.warning(self, "Error", "Contraseña incorrecta")
                return
            
            # Recoge datos de antecedentes
            id_antecedente = self.id_antecedente.text()
            id_paciente_2 = self.id_paciente_2.text()
            fecha_registro_2 = self.fecha_registro_2.date().toString("yyyy-MM-dd")

            ant_medicos = self.ant_medicos.toPlainText()
            ant_quirurgicos = self.ant_quirurgicos.toPlainText()
            ant_traumaticos = self.ant_traumaticos.toPlainText()
            ant_alergicos = self.ant_alergicos.toPlainText()
            ant_transfusionales = self.ant_transfusionales.toPlainText()
            ant_ginecoobstetricos = self.ant_ginecoobstetricos.toPlainText()

            # Convierte campos vacíos a None para permitir valores nulos en la base de datos
            def convert_to_none_if_empty(value):
                return value if value else None

            id_antecedente = convert_to_none_if_empty(id_antecedente)
            id_paciente_2 = convert_to_none_if_empty(id_paciente_2)
            ant_medicos = convert_to_none_if_empty(ant_medicos)
            ant_quirurgicos = convert_to_none_if_empty(ant_quirurgicos)
            ant_traumaticos = convert_to_none_if_empty(ant_traumaticos)
            ant_alergicos = convert_to_none_if_empty(ant_alergicos)
            ant_transfusionales = convert_to_none_if_empty(ant_transfusionales)
            ant_ginecoobstetricos = convert_to_none_if_empty(ant_ginecoobstetricos)

            # Valida que los campos obligatorios estén llenos
            if not all([id_antecedente, id_paciente_2, fecha_registro_2]):
                QMessageBox.warning(self, "Advertencia", "Los campos ID Antecedente, ID Paciente y Fecha de Registro son obligatorios.")
                return

            # Verifica si el id_paciente_3 existe en las tablas receptor o donante
            if not (self.db.existe_paciente_en_receptor_o_donante(id_paciente_2)):
                QMessageBox.warning(self, "Error", f"El ID del paciente {id_paciente_2} no se encuentra en las tablas receptor o donante.")
                return
            
            # Llama al método de la clase de conexión para agregar antecedentes
            if self.db.agregar_datos_2(id_antecedente, id_paciente_2, fecha_registro_2, ant_medicos, ant_quirurgicos, 
                                    ant_traumaticos, ant_alergicos, ant_transfusionales, ant_ginecoobstetricos):
                # Limpia la tabla antes de agregar el nuevo registro
                self.result_table2.setRowCount(0)  # Limpia la tabla

                # Actualiza solo la fila correspondiente en el QTableWidget para antecedentes
                row_position = self.result_table2.rowCount()
                self.result_table2.insertRow(row_position)
                self.result_table2.setItem(row_position, 0, QTableWidgetItem(id_antecedente))
                self.result_table2.setItem(row_position, 1, QTableWidgetItem(id_paciente_2))
                self.result_table2.setItem(row_position, 2, QTableWidgetItem(fecha_registro_2))

                QMessageBox.information(self, "Éxito", "Antecedentes agregados correctamente.")
                self.limpiar_antecedentes()  # Limpia el formulario de antecedentes después de agregar
            else:
                QMessageBox.warning(self, "Error", "No se pudieron agregar los antecedentes.")

        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Error", f"Error al agregar antecedentes: {str(e)}")

    #Es un condicional para verificar que RadioButton esta seleccionado y así restringir widgets dependiendo del tipo de paciente (receptor o donante)
    def buscar_antecedentes_condicional(self):
        if self.id_receptor_ant.isChecked():
            self.mostrar_antecedentes_receptor()
        elif self.id_donante_ant.isChecked():
            self.mostrar_antecedentes_donante()

    def mostrar_antecedentes_receptor(self):
        results = self.db.buscar_antecedentes_receptor()  # Llama al método del objeto de conexión
        self.limpiar_antecedentes()

        # Limpia la tabla antes de agregar nuevos datos
        self.result_table2.setRowCount(0)

        for row in results:
            row_position = self.result_table2.rowCount()
            self.result_table2.insertRow(row_position)
            
            for column, value in enumerate(row.values()):
                self.result_table2.setItem(row_position, column, QTableWidgetItem(str(value)))

    def mostrar_antecedentes_donante(self):
        if self.id_donante_ant.isChecked():
            try:
                self.limpiar_antecedentes()
                # Conecta a la base de datos
                self.db.conecta_base_datos()

                #Consulta para obtener datos de antecedentes de donante
                query = """
                SELECT a.* FROM a_antecedentes a
                JOIN donante d ON a.id_paciente_2 = d.id_donante
                """
                self.db.cursor.execute(query)
                results = self.db.cursor.fetchall()

                # Limpia las filas existentes en la tabla
                self.result_table2.setRowCount(0)

                # Llena la tabla con la información filtrada
                for row in results:
                    row_position = self.result_table2.rowCount()
                    self.result_table2.insertRow(row_position)
                    
                    for column, value in enumerate(row.values()):
                        self.result_table2.setItem(row_position, column, QTableWidgetItem(str(value)))

            except Exception as e:
                QMessageBox.critical(self, "Error", f"Error al mostrar datos de donantes: {str(e)}")
            finally:
                if self.db.cursor:
                    self.db.cursor.close()
                if self.db.con:
                    self.db.con.close()

    def limpiar_antecedentes(self):
    # Limpia los campos de texto de antecedentes
        fecha_default = QDate(2000, 1, 1)
        self.id_antecedente.clear() 
        self.id_paciente_2.clear()  
        self.fecha_registro_2.setDate(fecha_default)

        self.ant_medicos.clear()
        self.ant_quirurgicos.clear()
        self.ant_traumaticos.clear()
        self.ant_alergicos.clear()
        self.ant_transfusionales.clear()
        self.ant_ginecoobstetricos.clear()

    def seleccionar_antecedentes(self):
        selected_row = self.result_table2.currentRow()
        if selected_row >= 0:
            id_antecedente = self.result_table2.item(selected_row, 0).text()
            
            # Obtiene todos los datos del antecedente seleccionado de la base de datos
            antecedente_data = self.db.obtener_datos_antecedente(id_antecedente)
            self.limpiar_antecedentes()

            if antecedente_data:
                # Llena los widgets con la información
                self.id_antecedente.setText(str(antecedente_data['id_antecedente']))
                self.id_paciente_2.setText(str(antecedente_data['id_paciente_2']))
                self.fecha_registro_2.setDate(QDate.fromString(str(antecedente_data['fecha_registro_2']), "yyyy-MM-dd"))
                
                self.ant_medicos.setPlainText(antecedente_data['ant_medicos'] or "")
                self.ant_quirurgicos.setPlainText(antecedente_data['ant_quirurgicos'] or "")
                self.ant_traumaticos.setPlainText(antecedente_data['ant_traumaticos'] or "")
                self.ant_alergicos.setPlainText(antecedente_data['ant_alergicos'] or "")
                self.ant_transfusionales.setPlainText(antecedente_data['ant_transfusionales'] or "")
                self.ant_ginecoobstetricos.setPlainText(antecedente_data['ant_ginecoobstetricos'] or "")
            else:
                QMessageBox.warning(self, "Error", "No se pudo obtener la información del antecedente.")

    def actualizar_antecedentes(self):
        try:
            # Solicita y valida la contraseña
            contraseña = self.solicitar_contraseña()
            if contraseña is None:
                return
                
            if not self.validar_contraseña(contraseña):
                QtWidgets.QMessageBox.warning(self, "Error", "Contraseña incorrecta")
                return
            # Recoger datos de antecedentes
            id_antecedente = self.id_antecedente.text()
            id_paciente_2 = self.id_paciente_2.text()
            fecha_registro_2 = self.fecha_registro_2.date().toString("yyyy-MM-dd")

            ant_medicos = self.ant_medicos.toPlainText()
            ant_quirurgicos = self.ant_quirurgicos.toPlainText()
            ant_traumaticos = self.ant_traumaticos.toPlainText()
            ant_alergicos = self.ant_alergicos.toPlainText()
            ant_transfusionales = self.ant_transfusionales.toPlainText()
            ant_ginecoobstetricos = self.ant_ginecoobstetricos.toPlainText()

            # Validar que todos los campos estén llenos
            if not all([id_antecedente, id_paciente_2, fecha_registro_2, ant_medicos, ant_quirurgicos, ant_traumaticos, ant_alergicos, ant_transfusionales, ant_ginecoobstetricos]):
                QMessageBox.warning(self, "Advertencia", "Todos los campos de antecedentes son obligatorios.")
                return

            # Obtener el ID del registro que se desea actualizar
            selected_row = self.result_table2.currentRow()
            if selected_row < 0:
                QMessageBox.warning(self, "Advertencia", "Por favor, seleccione un registro para actualizar.")
                return

            # Llamar al método de la clase de conexión para actualizar los antecedentes
            if self.db.actualizar_datos_2(id_antecedente, id_paciente_2, fecha_registro_2, ant_medicos, ant_quirurgicos, ant_traumaticos, ant_alergicos, ant_transfusionales, ant_ginecoobstetricos):
                QMessageBox.information(self, "Éxito", "Antecedentes actualizados correctamente.")
                self.limpiar_antecedentes()  # Limpiar el formulario después de actualizar
                self.buscar_antecedentes_condicional()  # Volver a cargar los datos para mostrar todos
        
        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Error", f"Error al actualizar antecedentes: {str(e)}")

    def eliminar_antecedentes(self):
        try:
            # Solicita y valida la contraseña
            contraseña = self.solicitar_contraseña()
            if contraseña is None:
                return
                
            if not self.validar_contraseña(contraseña):
                QtWidgets.QMessageBox.warning(self, "Error", "Contraseña incorrecta")
                return  

            if not self.verificar_contraseña_admin():
                QMessageBox.warning(self, "Error", "Contraseña incorrecta.")
                return
            selected_row = self.result_table2.currentRow()
            
            if selected_row >= 0:
                # Suponiendo que el ID del receptor está en la primera columna
                id_receptor = self.result_table2.item(selected_row, 0).text()
                
                reply = QMessageBox.question(self, 'Confirmar eliminación', 
                                            '¿Está seguro de que desea eliminar este registro de antecedentes?',
                                            QMessageBox.Yes | QMessageBox.No, QMessageBox.No)
                
                if reply == QMessageBox.Yes:
                    if self.db.eliminar_datos_2(id_receptor):
                        self.result_table2.removeRow(selected_row)
                        QMessageBox.information(self, "Éxito", "Registro de antecedentes eliminado correctamente.")
                        self.limpiar_antecedentes()
                    else:
                        QMessageBox.warning(self, "Error", "No se pudo eliminar el registro de antecedentes.")
            else:
                QMessageBox.warning(self, "Advertencia", "Por favor, seleccione un registro de antecedentes para eliminar.")

        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Error", f"Error al eliminar donante: {str(e)}")

   ################## FASE 1 - Funciones básicas para interactuar con los datos (buscar,agregar,limpiar,eliminar datos)
    def agregar_fase1(self):
        try:
            # Solicita y valida la contraseña
            contraseña = self.solicitar_contraseña()
            if contraseña is None:
                return
                
            if not self.validar_contraseña(contraseña):
                QtWidgets.QMessageBox.warning(self, "Error", "Contraseña incorrecta")
                return
            # Recoge datos de los campos de entrada para la fase 1
            id_fase1 = self.id_fase1.text()
            id_paciente_3 = self.id_paciente_3.text()
            fecha_registro_3 = self.fecha_registro_3.date().toString("yyyy-MM-dd")
            peso = self.peso.value()
            talla = self.talla.value()
            valor_imc = self.valor_imc.value()
            est_imc = self.est_imc_map.get(self.est_imc.currentText(), None)
            hema_wbc = self.hema_wbc.value()
            hema_neutro = self.hema_neutro.value()
            hema_hgb = self.hema_hgb.value()
            hema_hct = self.hema_hct.value()
            hema_plt = self.hema_plt.value()
            creatinina = self.creatinina.value()
            tfg = self.tfg.value()
            bun = self.bun.value()
            gli_pre = self.gli_pre.value()
            gli_post = self.gli_post.value()
            hb_glicolisada = self.hb_glicolisada.value()
            rd_longitud = self.rd_longitud.value()
            rd_anchura = self.rd_anchura.value()
            rd_grosor = self.rd_grosor.value()
            ri_longitud = self.ri_longitud.value()
            ri_anchura = self.ri_anchura.value()
            ri_grosor = self.ri_grosor.value()

            usg_hep_bil = self.usg_hep_bil_map.get(self.usg_hep_bil.currentText(), None)
            obs_usg_hep_bil = self.obs_usg_hep_bil.toPlainText()
            pra = self.pra.value()

            # Valida que todos los campos obligatorios estén llenos
            if not all([id_fase1, id_paciente_3, fecha_registro_3]):
                QMessageBox.warning(self, "Advertencia", "Los campos ID Fase 1, ID Paciente y Fecha de Registro son obligatorios.")
                return

            # Verifica si el id_paciente existe en las tablas receptor o donante
            if not (self.db.existe_paciente_en_receptor_o_donante(id_paciente_3)):
                QMessageBox.warning(self, "Error", f"El ID del paciente {id_paciente_3} no se encuentra en las tablas receptor o donante.")
                return

            # Llama al método de la clase de conexión para agregar los datos de la fase 1
            if self.db.agregar_datos_3(id_fase1, id_paciente_3, fecha_registro_3, peso, talla, valor_imc, est_imc,
                                hema_wbc, hema_neutro, hema_hgb, hema_hct,
                                hema_plt, creatinina, tfg, bun,
                                gli_pre, gli_post, hb_glicolisada,
                                rd_longitud, rd_anchura,
                                rd_grosor, ri_longitud,
                                ri_anchura, ri_grosor,
                                usg_hep_bil, obs_usg_hep_bil, pra):
                # Limpia la tabla antes de agregar el nuevo registro
                self.result_table3.setRowCount(0)  # Limpiar la tabla

                # Actualiza solo las primeras tres columnas en el QTableWidget para la fase 1
                row_position = self.result_table3.rowCount()
                self.result_table3.insertRow(row_position)

                items_to_insert = [id_fase1, id_paciente_3, fecha_registro_3]
                
                for column_index, item in enumerate(items_to_insert):
                    self.result_table3.setItem(row_position, column_index, QTableWidgetItem(str(item)))

                QMessageBox.information(self, "Éxito", "Información de Fase 1 agregada correctamente.")
                self.limpiar_fase1()  # Limpia el formulario después de agregar
            else:
                QMessageBox.warning(self, "Error", "No se pudo agregar la información de Fase 1.")

        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Error", f"Error al agregar fase 1: {str(e)}")

     #Es un condicional para verificar que RadioButton esta seleccionado y así restringir widgets dependiendo del tipo de paciente (receptor o donante)
    def buscar_fase1_condicional(self):
        if self.id_receptor_f1.isChecked():
            self.mostrar_f1_receptor()
        elif self.id_donante_f1.isChecked():
            self.mostrar_f1_donante()

    def mostrar_f1_receptor(self):
        self.limpiar_fase1()
        results = self.db.buscar_f1_receptor()  # Llama al método del objeto de conexión

        # Limpia la tabla antes de agregar nuevos datos
        self.result_table3.setRowCount(0)

        for row in results:
            row_position = self.result_table3.rowCount()
            self.result_table3.insertRow(row_position)
            
            for column, value in enumerate(row.values()):
                self.result_table3.setItem(row_position, column, QTableWidgetItem(str(value)))

    def mostrar_f1_donante(self):
        if self.id_donante_f1.isChecked():
            try:
                # Conecta a la base de datos
                self.db.conecta_base_datos()
                self.limpiar_fase1()
                query = """
                SELECT f1.* FROM fase_1 f1
                JOIN donante d ON f1.id_paciente_3 = d.id_donante
                """
                self.db.cursor.execute(query)
                results = self.db.cursor.fetchall()

                # Limpia las filas existentes de la tabla
                self.result_table3.setRowCount(0)

                # Llena la data con la data filtrada
                for row in results:
                    row_position = self.result_table3.rowCount()
                    self.result_table3.insertRow(row_position)
                    
                    for column, value in enumerate(row.values()):
                        self.result_table3.setItem(row_position, column, QTableWidgetItem(str(value)))

            except Exception as e:
                QMessageBox.critical(self, "Error", f"Error al mostrar datos de donantes: {str(e)}")
            finally:
                if self.db.cursor:
                    self.db.cursor.close()
                if self.db.con:
                    self.db.con.close()

    def limpiar_fase1(self):
        # Limpia los campos de texto de la fase 1
        self.id_fase1.clear()  
        self.id_paciente_3.clear()  
        fecha_default = QDate(2000, 1, 1)
        self.fecha_registro_3.setDate(fecha_default)
        self.peso.setValue(0.0)  
        self.talla.setValue(0.0)  
        self.valor_imc.setValue(0.0)  
        self.est_imc.setCurrentIndex(0)
        self.hema_wbc.setValue(0.0)
        self.hema_neutro.setValue(0.0)
        self.hema_hgb.setValue(0.0)
        self.hema_hct.setValue(0.0)
        self.hema_plt.setValue(0.0)
        self.creatinina.setValue(0.0)
        self.tfg.setValue(0.0)
        self.bun.setValue(0.0)
        self.gli_pre.setValue(0.0)
        self.gli_post.setValue(0.0)
        self.hb_glicolisada.setValue(0.0)
        self.rd_longitud.setValue(0.0)
        self.rd_anchura.setValue(0.0)
        self.rd_grosor.setValue(0.0)
        self.ri_longitud.setValue(0.0)
        self.ri_anchura.setValue(0.0)
        self.ri_grosor.setValue(0.0)
        self.usg_hep_bil.setCurrentIndex(0)
        self.obs_usg_hep_bil.clear()
        self.pra.setValue(0.0)          

    def seleccionar_fase1(self):
        selected_row = self.result_table3.currentRow()
        if selected_row >= 0:
            id_fase1 = self.result_table3.item(selected_row, 0).text()
            
            # Obtiene todos los datos de la fase 1 seleccionada de la base de datos
            fase1_data = self.db.obtener_datos_fase1(id_fase1)
            
            if fase1_data:
                # Función auxiliar para obtener valores float o establecerlos en 0
                def get_float_value(value):
                    try:
                        return float(value) if value not in [None, '', 'None'] else 0.0
                    except ValueError:
                        return 0.0

                # Función auxiliar para obtener valores int o establecerlos en 0
                def get_int_value(value):
                    try:
                        return int(value) if value not in [None, '', 'None'] else 0
                    except ValueError:
                        return 0
                    
                self.limpiar_fase1()

                # Llena los widgets con la información
                self.id_fase1.setText(str(fase1_data['id_fase1']))
                self.id_paciente_3.setText(str(fase1_data['id_paciente_3']))
                self.fecha_registro_3.setDate(QDate.fromString(str(fase1_data['fecha_registro_3']), "yyyy-MM-dd"))
                self.peso.setValue(get_float_value(fase1_data['peso']))
                self.talla.setValue(get_float_value(fase1_data['talla']))
                self.valor_imc.setValue(get_float_value(fase1_data['valor_imc']))

                est_imc_value = get_int_value(fase1_data['est_imc'])
                if est_imc_value in self.est_imc_map.values():
                    self.est_imc.setCurrentText(next(key for key, value in self.est_imc_map.items() if value == est_imc_value))

                self.hema_wbc.setValue(get_float_value(fase1_data['hema_wbc']))
                self.hema_neutro.setValue(get_float_value(fase1_data['hema_neutro']))
                self.hema_hgb.setValue(get_float_value(fase1_data['hema_hgb']))
                self.hema_hct.setValue(get_float_value(fase1_data['hema_hct']))
                self.hema_plt.setValue(get_float_value(fase1_data['hema_plt']))
                self.creatinina.setValue(get_float_value(fase1_data['creatinina']))
                self.tfg.setValue(get_float_value(fase1_data['tfg']))
                self.bun.setValue(get_float_value(fase1_data['bun']))
                self.gli_pre.setValue(get_float_value(fase1_data['gli_pre']))
                self.gli_post.setValue(get_float_value(fase1_data['gli_post']))
                self.hb_glicolisada.setValue(get_float_value(fase1_data['hb_glicolisada']))
                self.rd_longitud.setValue(get_float_value(fase1_data['rd_longitud']))
                self.rd_anchura.setValue(get_float_value(fase1_data['rd_anchura']))
                self.rd_grosor.setValue(get_float_value(fase1_data['rd_grosor']))
                self.ri_longitud.setValue(get_float_value(fase1_data['ri_longitud']))
                self.ri_anchura.setValue(get_float_value(fase1_data['ri_anchura']))
                self.ri_grosor.setValue(get_float_value(fase1_data['ri_grosor']))

                usg_hep_bil_value = get_int_value(fase1_data['usg_hep_bil'])
                if usg_hep_bil_value in self.usg_hep_bil_map.values():
                    self.usg_hep_bil.setCurrentText(next(key for key, value in self.usg_hep_bil_map.items() if value == usg_hep_bil_value))

                self.obs_usg_hep_bil.setPlainText(str(fase1_data['obs_usg_hep_bil']) if fase1_data['obs_usg_hep_bil'] else "")

                self.pra.setValue(get_float_value(fase1_data['pra']))
            else:
                QMessageBox.warning(self, "Error", "No se pudo obtener la información de la Fase 1.")
    
    def actualizar_fase1(self):
        try:
            # Solicita y valida la contraseña
            contraseña = self.solicitar_contraseña()
            if contraseña is None:
                return
                
            if not self.validar_contraseña(contraseña):
                QtWidgets.QMessageBox.warning(self, "Error", "Contraseña incorrecta")
                return
        
            # Recoge datos de los campos de entrada para la fase 1
            id_fase1 = self.id_fase1.text()
            id_paciente_3 = self.id_paciente_3.text()
            fecha_registro_3 = self.fecha_registro_3.date().toString("yyyy-MM-dd")
            peso = self.peso.value()
            talla = self.talla.value()
            valor_imc = self.valor_imc.value()

            # Obtiene el valor entero correspondiente al estado IMC seleccionado
            est_imc_text = self.est_imc.currentText()
            est_imc = self.est_imc_map.get(est_imc_text, None)  # Permite None si está vacío

            hema_wbc = self.hema_wbc.value()
            hema_neutro = self.hema_neutro.value()
            hema_hgb = self.hema_hgb.value()
            hema_hct = self.hema_hct.value()
            hema_plt = self.hema_plt.value()
            creatinina = self.creatinina.value()
            tfg = self.tfg.value()
            bun = self.bun.value()
            gli_pre = self.gli_pre.value()
            gli_post = self.gli_post.value()
            hb_glicolisada = self.hb_glicolisada.value()
            rd_longitud = self.rd_longitud.value()
            rd_anchura = self.rd_anchura.value()
            rd_grosor = self.rd_grosor.value()
            ri_longitud = self.ri_longitud.value()
            ri_anchura = self.ri_anchura.value()
            ri_grosor = self.ri_grosor.value()

            # Obtiene el valor entero correspondiente al USG Hepático Biliar seleccionado
            usg_hep_bil_text = self.usg_hep_bil.currentText()
            usg_hep_bil = self.usg_hep_bil_map.get(usg_hep_bil_text, None)  # Permitir None si está vacío

            obs_usg_hep_bil = self.obs_usg_hep_bil.toPlainText() or None  # Permitir None si está vacío
            pra = self.pra.value()

            # Llama al método de la clase de conexión para actualizar los datos de la fase 1
            if self.db.actualizar_datos_3(
                id_fase1, id_paciente_3, fecha_registro_3, peso, talla, valor_imc, est_imc,
                hema_wbc, hema_neutro, hema_hgb, hema_hct,
                hema_plt, creatinina, tfg, bun,
                gli_pre, gli_post, hb_glicolisada,
                rd_longitud, rd_anchura,
                rd_grosor, ri_longitud,
                ri_anchura, ri_grosor,
                usg_hep_bil, obs_usg_hep_bil,
                pra):
                
                QMessageBox.information(self, "Éxito", "Información actualizada correctamente.")
                self.limpiar_fase1()  # Limpiar el formulario después de actualizar
                self.buscar_fase1_condicional()  # Volver a cargar los datos para mostrar todos
        
        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Error", f"Error al actualizar fase 1: {str(e)}")


    def eliminar_fase1(self):
        #Elimina un registro de fase 1 después de validar la contraseña
        try:
            # Solicita y valida la contraseña
            contraseña = self.solicitar_contraseña()
            if contraseña is None:
                return
                
            if not self.validar_contraseña(contraseña):
                QtWidgets.QMessageBox.warning(self, "Error", "Contraseña incorrecta")
                return
            
            # Si la contraseña es válida, procede con eliminar
            selected_row = self.result_table3.currentRow()
            
            if selected_row >= 0:
                # Obtiene el ID de la fase 1 de la primera columna
                id_fase1 = self.result_table3.item(selected_row, 0).text()
                
                reply = QtWidgets.QMessageBox.question(
                    self, 
                    'Confirmar eliminación', 
                    '¿Está seguro de que desea eliminar este registro de fase 1?',
                    QtWidgets.QMessageBox.Yes | QtWidgets.QMessageBox.No, 
                    QtWidgets.QMessageBox.No
                )
                
                if reply == QtWidgets.QMessageBox.Yes:
                    if self.db.eliminar_datos_3(id_fase1):
                        self.result_table3.removeRow(selected_row)
                        QtWidgets.QMessageBox.information(self, "Éxito", "Registro de fase 1 eliminado correctamente.")
                        self.limpiar_fase1()
                    else:
                        QtWidgets.QMessageBox.warning(self, "Error", "No se pudo eliminar el registro de fase 1.")
            else:
                QtWidgets.QMessageBox.warning(self, "Advertencia", "Por favor, seleccione un registro de fase 1 para eliminar.")
                
        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Error", f"Error al eliminar fase 1: {str(e)}")

    def calcular_imc(self):
        peso = self.peso.value()  # Peso en kg
        talla = self.talla.value()  # Talla en metros
        
        if talla == 0:
            QMessageBox.warning(self, "Error", "La talla no puede ser cero.")
            return
        
        # Calcula IMC con peso en kg y talla en metros
        imc = peso / (talla ** 2)
        self.valor_imc.setValue(round(imc, 2))
        
        # Clasifica el IMC
        if imc < 18.5:
            clasificacion = "Bajo peso"
        elif 18.5 <= imc < 25:
            clasificacion = "Normal"
        elif 25 <= imc < 30:
            clasificacion = "Sobrepeso"
        elif 30 <= imc < 40:
            clasificacion = "Obesidad"
        else:
            clasificacion = "Obesidad mórbida"
        
        # Establece la clasificación en el combobox
        index = self.est_imc.findText(clasificacion)
        if index >= 0:
            self.est_imc.setCurrentIndex(index)
        
        # Muestra mensaje con el resultado del IMC y su clasificación
        QMessageBox.information(self, "IMC Calculado", f"IMC: {imc:.2f}\nClasificación: {clasificacion}")

   ################## FASE 2A - Funciones básicas para interactuar con los datos (buscar,agregar,limpiar,eliminar datos)

    def agregar_fase2a(self):
        try:
            # Solicita y valida la contraseña
            contraseña = self.solicitar_contraseña()
            if contraseña is None:
                return
                
            if not self.validar_contraseña(contraseña):
                QtWidgets.QMessageBox.warning(self, "Error", "Contraseña incorrecta")
                return
            # Recoge datos de los campos de entrada para la fase 2a
            id_fase2a = self.id_fase2a.text()
            id_paciente_4 = self.id_paciente_4.text()
            fecha_registro_4 = self.fecha_registro_4.date().toString("yyyy-MM-dd")
            vac_cov19 = self.estado_vacuna_map.get(self.vac_cov19.currentText(), None)
            dos1_cov19 = self.dos1_cov19.date().toString("yyyy-MM-dd")
            tipo_dos1 = self.tipo_vacuna_map.get(self.tipo_dos1.currentText(), None)
            dos2_cov19 = self.dos2_cov19.date().toString("yyyy-MM-dd")
            tipo_dos2 = self.tipo_vacuna_map.get(self.tipo_dos2.currentText(), None)
            dos3_cov19 = self.dos3_cov19.date().toString("yyyy-MM-dd")
            tipo_dos3 = self.tipo_vacuna_map.get(self.tipo_dos3.currentText(), None)
            vac_influen = self.estado_vacuna_map.get(self.vac_influen.currentText(), None)
            dos_influen = self.dos_influen.date().toString("yyyy-MM-dd")
            vac_neum = self.estado_vacuna_map.get(self.vac_neum.currentText(), None)
            dos_neum = self.dos_neum.date().toString("yyyy-MM-dd")
            cert_dental = self.estado_eval_map.get(self.cert_dental.currentText(), None)
            obs_dental = self.obs_dental.toPlainText()
            eval_psico = self.estado_eval_map.get(self.eval_psico.currentText(), None)
            obs_eval1 = self.obs_eval1.toPlainText()
            eval_trab_social = self.estado_eval_map.get(self.eval_trab_social.currentText(), None)
            obs_eval2 = self.obs_eval2.toPlainText()
            eval_legal = self.estado_eval_map.get(self.eval_legal.currentText(), None)
            obs_eval3 = self.obs_eval3.toPlainText()
            eval_nutri = self.estado_eval_map.get(self.eval_nutri.currentText(), None)
            obs_eval4 = self.obs_eval4.toPlainText()
            masa_muscular = self.masa_muscular.value()
            grasa = self.grasa.value()
            agua = self.agua.value()
            # Valida que todos los campos obligatorios estén llenos
            if not all([id_fase2a, id_paciente_4, fecha_registro_4]):
                QMessageBox.warning(self, "Advertencia", "Los campos ID Fase 2a, ID Paciente y Fecha de Registro son obligatorios.")
                return

            # Verifica si el id_paciente existe en las tablas receptor o donante
            if not (self.db.existe_paciente_en_receptor_o_donante(id_paciente_4)):
                QMessageBox.warning(self, "Error", f"El ID del paciente {id_paciente_4} no se encuentra en las tablas receptor o donante.")
                return
            
            # Llama al método de la clase de conexión para agregar los datos de la fase 2a
            if self.db.agregar_datos_4(id_fase2a, id_paciente_4, fecha_registro_4, vac_cov19,
                                    dos1_cov19, tipo_dos1, dos2_cov19, tipo_dos2,
                                    dos3_cov19, tipo_dos3, vac_influen, dos_influen,
                                    vac_neum, dos_neum, cert_dental, obs_dental,
                                    eval_psico, obs_eval1, eval_trab_social,
                                    obs_eval2, eval_legal, obs_eval3,
                                    eval_nutri, obs_eval4,
                                    masa_muscular, grasa, agua):
                # Limpia la tabla antes de agregar el nuevo registro
                self.result_table4.setRowCount(0)  # Limpia la tabla

                # Actualizarsolo las primeras tres columnas en el QTableWidget para la fase 2a
                row_position = self.result_table4.rowCount()
                self.result_table4.insertRow(row_position)
                items_to_insert = [id_fase2a, id_paciente_4, fecha_registro_4]
                for column_index, item in enumerate(items_to_insert):
                    self.result_table4.setItem(row_position, column_index, QTableWidgetItem(str(item)))

                QMessageBox.information(self, "Éxito", "Información de Fase 2a agregada correctamente.")
                self.limpiar_fase2a()  # Limpia el formulario después de agregar
            else:
                QMessageBox.warning(self, "Error", "No se pudo agregar la información de Fase 2a.")

        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Error", f"Error al agregar fase 2a: {str(e)}")
        
 #Es un condicional para verificar que RadioButton esta seleccionado y así restringir widgets dependiendo del tipo de paciente (receptor o donante)
    def buscar_fase2a_condicional(self):
        if self.id_receptor_f2a.isChecked():
            self.mostrar_f2a_receptor()
        elif self.id_donante_f2a.isChecked():
            self.mostrar_f2a_donante()

    def mostrar_f2a_receptor(self):
        results = self.db.buscar_f2a_receptor()  # Llama al método del objeto de conexión
        self.limpiar_fase2a()

        # Limpia la tabla antes de agregar nuevos datos
        self.result_table4.setRowCount(0)

        for row in results:
            row_position = self.result_table4.rowCount()
            self.result_table4.insertRow(row_position)
            
            for column, value in enumerate(row.values()):
                self.result_table4.setItem(row_position, column, QTableWidgetItem(str(value)))

    def mostrar_f2a_donante(self):
        if self.id_donante_f2a.isChecked():
            try:
                # Conecta a la base de datos
                self.db.conecta_base_datos()
                self.limpiar_fase2a()

                query = """
                SELECT f2a.* FROM fase_2a f2a
                JOIN donante d ON f2a.id_paciente_4 = d.id_donante
                """
                self.db.cursor.execute(query)
                results = self.db.cursor.fetchall()

                self.result_table4.setRowCount(0)

                for row in results:
                    row_position = self.result_table4.rowCount()
                    self.result_table4.insertRow(row_position)
                    
                    for column, value in enumerate(row.values()):
                        self.result_table4.setItem(row_position, column, QTableWidgetItem(str(value)))

            except Exception as e:
                QMessageBox.critical(self, "Error", f"Error al mostrar datos de donantes: {str(e)}")
            finally:
                if self.db.cursor:
                    self.db.cursor.close()
                if self.db.con:
                    self.db.con.close()

    def limpiar_fase2a(self):
        # Limpia los campos de texto de la fase 2a
        self.id_fase2a.clear()
        self.id_paciente_4.clear()
        
        # Restablece las fechas a un valor predeterminado
        fecha_default = QDate(2000, 1, 1)
        self.fecha_registro_4.setDate(fecha_default)
        self.dos1_cov19.setDate(fecha_default)
        self.dos2_cov19.setDate(fecha_default)
        self.dos3_cov19.setDate(fecha_default)
        self.dos_influen.setDate(fecha_default)
        self.dos_neum.setDate(fecha_default)
        self.vac_cov19.setCurrentIndex(0)
        self.tipo_dos1.setCurrentIndex(0)
        self.tipo_dos2.setCurrentIndex(0)
        self.tipo_dos3.setCurrentIndex(0)
        self.vac_influen.setCurrentIndex(0)
        self.vac_neum.setCurrentIndex(0)
        self.cert_dental.setCurrentIndex(0)
        self.eval_psico.setCurrentIndex(0)
        self.eval_trab_social.setCurrentIndex(0)
        self.eval_legal.setCurrentIndex(0)
        self.eval_nutri.setCurrentIndex(0)
        self.obs_dental.clear()
        self.obs_eval1.clear()
        self.obs_eval2.clear()
        self.obs_eval3.clear()
        self.obs_eval4.clear()
        self.masa_muscular.setValue(0.0)
        self.grasa.setValue(0.0)
        self.agua.setValue(0.0)
    
    def seleccionar_fase2a(self):
        selected_row = self.result_table4.currentRow()
        if selected_row >= 0:
            id_fase2a = self.result_table4.item(selected_row, 0).text()
            
            # Obtiene todos los datos de la fase 2a seleccionada de la base de datos
            fase2a_data = self.db.obtener_datos_fase2a(id_fase2a)
            
            if fase2a_data:
                # Función auxiliar para obtener valores float o establecerlos en 0
                def get_float_value(value):
                    try:
                        return float(value) if value not in [None, '', 'None'] else 0.0
                    except ValueError:
                        return 0.0

                # Función auxiliar para obtener valores int o establecerlos en 0
                def get_int_value(value):
                    try:
                        return int(value) if value not in [None, '', 'None'] else 0
                    except ValueError:
                        return 0

                # Limpia todos los campos antes de llenarlos con nueva información
                self.limpiar_fase2a()

                # Llena los widgets con la información
                self.id_fase2a.setText(str(fase2a_data['id_fase2a']))
                self.id_paciente_4.setText(str(fase2a_data['id_paciente_4']))
                self.fecha_registro_4.setDate(QDate.fromString(str(fase2a_data['fecha_registro_4']), "yyyy-MM-dd"))
                
                vac_cov19_value = get_int_value(fase2a_data['vac_cov19'])
                if vac_cov19_value in self.estado_vacuna_map.values():
                    self.vac_cov19.setCurrentText(next(key for key, value in self.estado_vacuna_map.items() if value == vac_cov19_value))
                
                self.dos1_cov19.setDate(QDate.fromString(str(fase2a_data['dos1_cov19']), "yyyy-MM-dd"))
                
                tipo_dos1_value = get_int_value(fase2a_data['tipo_dos1'])
                if tipo_dos1_value in self.tipo_vacuna_map.values():
                    self.tipo_dos1.setCurrentText(next(key for key, value in self.tipo_vacuna_map.items() if value == tipo_dos1_value))
                
                self.dos2_cov19.setDate(QDate.fromString(str(fase2a_data['dos2_cov19']), "yyyy-MM-dd"))
                
                tipo_dos2_value = get_int_value(fase2a_data['tipo_dos2'])
                if tipo_dos2_value in self.tipo_vacuna_map.values():
                    self.tipo_dos2.setCurrentText(next(key for key, value in self.tipo_vacuna_map.items() if value == tipo_dos2_value))
                
                self.dos3_cov19.setDate(QDate.fromString(str(fase2a_data['dos3_cov19']), "yyyy-MM-dd"))
                
                tipo_dos3_value = get_int_value(fase2a_data['tipo_dos3'])
                if tipo_dos3_value in self.tipo_vacuna_map.values():
                    self.tipo_dos3.setCurrentText(next(key for key, value in self.tipo_vacuna_map.items() if value == tipo_dos3_value))
                
                vac_influen_value = get_int_value(fase2a_data['vac_influen'])
                if vac_influen_value in self.estado_vacuna_map.values():
                    self.vac_influen.setCurrentText(next(key for key, value in self.estado_vacuna_map.items() if value == vac_influen_value))
                
                self.dos_influen.setDate(QDate.fromString(str(fase2a_data['dos_influen']), "yyyy-MM-dd"))
                
                vac_neum_value = get_int_value(fase2a_data['vac_neum'])
                if vac_neum_value in self.estado_vacuna_map.values():
                    self.vac_neum.setCurrentText(next(key for key, value in self.estado_vacuna_map.items() if value == vac_neum_value))
                
                self.dos_neum.setDate(QDate.fromString(str(fase2a_data['dos_neum']), "yyyy-MM-dd"))
                
                cert_dental_value = get_int_value(fase2a_data['cert_dental'])
                if cert_dental_value in self.estado_eval_map.values():
                    self.cert_dental.setCurrentText(next(key for key, value in self.estado_eval_map.items() if value == cert_dental_value))
                
                self.obs_dental.setPlainText(str(fase2a_data['obs_dental']) if fase2a_data['obs_dental'] else "")
                
                eval_psico_value = get_int_value(fase2a_data['eval_psico'])
                if eval_psico_value in self.estado_eval_map.values():
                    self.eval_psico.setCurrentText(next(key for key, value in self.estado_eval_map.items() if value == eval_psico_value))
                
                self.obs_eval1.setPlainText(str(fase2a_data['obs_eval1']) if fase2a_data['obs_eval1'] else "")
                
                eval_trab_social_value = get_int_value(fase2a_data['eval_trab_social'])
                if eval_trab_social_value in self.estado_eval_map.values():
                    self.eval_trab_social.setCurrentText(next(key for key, value in self.estado_eval_map.items() if value == eval_trab_social_value))
                
                self.obs_eval2.setPlainText(str(fase2a_data['obs_eval2']) if fase2a_data['obs_eval2'] else "")
                
                eval_legal_value = get_int_value(fase2a_data['eval_legal'])
                if eval_legal_value in self.estado_eval_map.values():
                    self.eval_legal.setCurrentText(next(key for key, value in self.estado_eval_map.items() if value == eval_legal_value))
                
                self.obs_eval3.setPlainText(str(fase2a_data['obs_eval3']) if fase2a_data['obs_eval3'] else "")
                
                eval_nutri_value = get_int_value(fase2a_data['eval_nutri'])
                if eval_nutri_value in self.estado_eval_map.values():
                    self.eval_nutri.setCurrentText(next(key for key, value in self.estado_eval_map.items() if value == eval_nutri_value))
                
                self.obs_eval4.setPlainText(str(fase2a_data['obs_eval4']) if fase2a_data['obs_eval4'] else "")
                
                self.masa_muscular.setValue(get_float_value(fase2a_data['masa_muscular']))
                self.grasa.setValue(get_float_value(fase2a_data['grasa']))
                self.agua.setValue(get_float_value(fase2a_data['agua']))
            else:
                QMessageBox.warning(self, "Error", "No se pudo obtener la información de la Fase 2a.")

    def actualizar_fase2a(self):
        try:
            # Solicita y valida la contraseña
            contraseña = self.solicitar_contraseña()
            if contraseña is None:
                return
                
            if not self.validar_contraseña(contraseña):
                QtWidgets.QMessageBox.warning(self, "Error", "Contraseña incorrecta")
                return
            # Recoger datos de los campos de entrada para la fase 2a
            id_fase2a = self.id_fase2a.text()
            id_paciente_4 = self.id_paciente_4.text()
            fecha_registro_4 = self.fecha_registro_4.date().toString("yyyy-MM-dd")
            vac_cov19 = self.estado_vacuna_map[self.vac_cov19.currentText()]
            dos1_cov19 = self.dos1_cov19.date().toString("yyyy-MM-dd")
            tipo_dos1 = self.tipo_vacuna_map[self.tipo_dos1.currentText()]
            dos2_cov19 = self.dos2_cov19.date().toString("yyyy-MM-dd")
            tipo_dos2 = self.tipo_vacuna_map[self.tipo_dos2.currentText()]
            dos3_cov19 = self.dos3_cov19.date().toString("yyyy-MM-dd")
            tipo_dos3 = self.tipo_vacuna_map[self.tipo_dos3.currentText()]
            vac_influen = self.estado_vacuna_map[self.vac_influen.currentText()]
            dos_influen = self.dos_influen.date().toString("yyyy-MM-dd")
            vac_neum = self.estado_vacuna_map[self.vac_neum.currentText()]
            dos_neum = self.dos_neum.date().toString("yyyy-MM-dd")
            cert_dental = self.estado_eval_map[self.cert_dental.currentText()]
            obs_dental = self.obs_dental.toPlainText()
            eval_psico = self.estado_eval_map[self.eval_psico.currentText()]
            obs_eval1 = self.obs_eval1.toPlainText()
            eval_trab_social = self.estado_eval_map[self.eval_trab_social.currentText()]
            obs_eval2 = self.obs_eval2.toPlainText()
            eval_legal = self.estado_eval_map[self.eval_legal.currentText()]
            obs_eval3 = self.obs_eval3.toPlainText()
            eval_nutri = self.estado_eval_map[self.eval_nutri.currentText()]
            obs_eval4 = self.obs_eval4.toPlainText()
            masa_muscular = self.masa_muscular.value()
            grasa = self.grasa.value()
            agua = self.agua.value()

            # Llamar al método de la clase de conexión para actualizar los datos de la fase 2a
            if self.db.actualizar_datos_4(id_fase2a, id_paciente_4, fecha_registro_4, vac_cov19,
                                            dos1_cov19, tipo_dos1, dos2_cov19, tipo_dos2,
                                            dos3_cov19, tipo_dos3, vac_influen, dos_influen,
                                            vac_neum, dos_neum, cert_dental, obs_dental,
                                            eval_psico, obs_eval1, eval_trab_social,
                                            obs_eval2, eval_legal, obs_eval3,
                                            eval_nutri, obs_eval4,
                                            masa_muscular, grasa, agua):
                QMessageBox.information(self, "Éxito", "Información actualizada correctamente.")
                # Limpia el formulario después de actualizar
                self.limpiar_fase2a()
                # Vuelve a cargar los datos para mostrar todos
                self.buscar_fase2a_condicional()

        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Error", f"Error al actualizar fase 2a: {str(e)}")

    def eliminar_fase2a(self):
        #Elimina un registro de fase 2a después de validar la contraseña
        try:
            # Solicita y valida la contraseña
            contraseña = self.solicitar_contraseña()
            if contraseña is None:
                return
                
            if not self.validar_contraseña(contraseña):
                QtWidgets.QMessageBox.warning(self, "Error", "Contraseña incorrecta")
                return
            
            # Si la contraseña es válida, procede con eliminar
            selected_row = self.result_table4.currentRow()
            
            if selected_row >= 0:
                # Obtiene el ID de la fase 2a de la primera columna
                id_fase2a = self.result_table4.item(selected_row, 0).text()
                
                reply = QtWidgets.QMessageBox.question(
                    self, 
                    'Confirmar eliminación', 
                    '¿Está seguro de que desea eliminar este registro de fase 2a?',
                    QtWidgets.QMessageBox.Yes | QtWidgets.QMessageBox.No, 
                    QtWidgets.QMessageBox.No
                )
                
                if reply == QtWidgets.QMessageBox.Yes:
                    if self.db.eliminar_datos_4(id_fase2a):
                        self.result_table4.removeRow(selected_row)
                        QtWidgets.QMessageBox.information(self, "Éxito", "Registro de fase 2a eliminado correctamente.")
                        self.limpiar_fase2a()
                    else:
                        QtWidgets.QMessageBox.warning(self, "Error", "No se pudo eliminar el registro de fase 2a.")
            else:
                QtWidgets.QMessageBox.warning(self, "Advertencia", "Por favor, seleccione un registro de fase 2a para eliminar.")
                
        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Error", f"Error al eliminar fase 2a: {str(e)}")


################## FASE 2B - Funciones básicas para interactuar con los datos (buscar,agregar,limpiar,eliminar datos)
    
    def agregar_fase2b(self):
        try:
            # Solicita y valida la contraseña
            contraseña = self.solicitar_contraseña()
            if contraseña is None:
                return
                
            if not self.validar_contraseña(contraseña):
                QtWidgets.QMessageBox.warning(self, "Error", "Contraseña incorrecta")
                return
            
            # Recoge datos de los campos de entrada para la fase 2b
            id_fase2b = self.id_fase2b.text()
            id_paciente_5 = self.id_paciente_5.text()
            fecha_registro_5 = self.fecha_registro_5.date().toString("yyyy-MM-dd")

            # Recoge todos los demás campos necesarios
            vih = self.estado_pruebas.get(self.vih.currentText(), None)
            hepa_b = self.estado_pruebas.get(self.hepa_b.currentText(), None)
            hepa_c = self.estado_pruebas.get(self.hepa_c.currentText(), None)
            vdrl = self.estado_pruebas.get(self.vdrl.currentText(), None)
            toxo_igg = self.estado_pruebas1.get(self.toxo_igg.currentText(), None)
            toxo_igm = self.estado_pruebas1.get(self.toxo_igm.currentText(), None)
            rubeo_igg = self.estado_pruebas1.get(self.rubeo_igg.currentText(), None)
            rubeo_igm = self.estado_pruebas1.get(self.rubeo_igm.currentText(), None)
            cmv_igg = self.estado_pruebas1.get(self.cmv_igg.currentText(), None)
            cmv_igm = self.estado_pruebas1.get(self.cmv_igm.currentText(), None)
            herpes_igg = self.estado_pruebas1.get(self.herpes_igg.currentText(), None)
            herpes_igm = self.estado_pruebas1.get(self.herpes_igm.currentText(), None)
            veb_igg = self.estado_pruebas1.get(self.veb_igg.currentText(), None)
            veb_igm = self.estado_pruebas1.get(self.veb_igm.currentText(), None)
            cuantiferon = self.estado_pruebas.get(self.cuantiferon.currentText(), None)
            obs_cuantiferon = self.obs_cuantiferon.toPlainText()
            tgo = self.tgo.value()
            tgp = self.tgp.value()
            bt = self.bt.value()
            bd = self.bd.value()
            bi = self.bi.value()
            dhl = self.dhl.value()
            fa = self.fa.value()
            albumina = self.albumina.value()
            acido_urico = self.acido_urico.value()
            perfil_hdl = self.perfil_hdl.value()
            perfil_ldl = self.perfil_ldl.value()
            perfil_ct = self.perfil_ct.value()
            perfil_tg = self.perfil_tg.value()
            perfil_vldl = self.perfil_vldl.value()
            sodio = self.sodio.value()
            potasio = self.potasio.value()
            fosforo = self.fosforo.value()
            calcio = self.calcio.value()
            magnesio = self.magnesio.value()
            pth = self.pth.value()
            tsh = self.tsh.value()
            t4_libre = self.t4_libre.value()
            c3 = self.c3.value()
            c4 = self.c4.value()
            bnp = self.bnp.value()
            tp = self.tp.value()
            tpt = self.tpt.value()
            inr = self.inr.value()
            nivel_ecg = self.nivel_ecg_map.get(self.nivel_ecg.currentText(), None)
            obs_ecg = self.obs_ecg.toPlainText()
            anti_dna = self.estado_pruebas.get(self.anti_dna.currentText(), None)
            ana = self.estado_pruebas.get(self.ana.currentText(), None)
            b2_glicoprot = self.estado_pruebas.get(self.b2_glicoprot.currentText(), None)
            anticoag_lupico = self.estado_pruebas.get(self.anticoag_lupico.currentText(), None)
            anticardio_igg_igm = self.estado_pruebas.get(self.anticardio_igg_igm.currentText(), None)
            p_anca = self.estado_pruebas.get(self.p_anca.currentText(), None)
            c_anca = self.estado_pruebas.get(self.c_anca.currentText(), None)
            urocultivo = self.estado_pruebas.get(self.urocultivo.currentText(), None)
            obs_urocultivo = self.obs_urocultivo.toPlainText()
            orocultivo = self.estado_pruebas.get(self.orocultivo.currentText(), None)
            obs_orocultivo = self.obs_orocultivo.toPlainText()
            ex_nasal = self.estado_pruebas.get(self.ex_nasal.currentText(), None)
            obs_nasal = self.obs_nasal.toPlainText()

            # Valida que todos los campos obligatorios estén llenos
            if not all([id_fase2b, id_paciente_5, fecha_registro_5]):
                QMessageBox.warning(self, "Advertencia", "Los campos ID Fase 2b, ID Paciente y Fecha de Registro son obligatorios.")
                return
            
            # Verifica si el id_paciente existe en las tablas receptor o donante
            if not (self.db.existe_paciente_en_receptor_o_donante(id_paciente_5)):
                QMessageBox.warning(self, "Error", f"El ID del paciente {id_paciente_5} no se encuentra en las tablas receptor o donante.")
                return

            # Llama al método de la clase de conexión para agregar los datos de la fase 2b
            if self.db.agregar_datos_5(
                id_fase2b, id_paciente_5, fecha_registro_5,
                vih, hepa_b, hepa_c, vdrl,
                toxo_igg, toxo_igm, rubeo_igg, rubeo_igm,
                cmv_igg, cmv_igm, herpes_igg, herpes_igm,
                veb_igg, veb_igm, cuantiferon, obs_cuantiferon,
                tgo, tgp, bt, bd, bi,
                dhl, fa, albumina, acido_urico,
                perfil_hdl, perfil_ldl, perfil_ct,
                perfil_tg, perfil_vldl,
                sodio, potasio, fosforo,
                calcio, magnesio,
                pth, tsh, t4_libre,
                c3, c4, bnp,
                tp, tpt, inr,
                nivel_ecg, obs_ecg,
                anti_dna, ana,
                b2_glicoprot, anticoag_lupico,
                anticardio_igg_igm,
                p_anca, c_anca,
                urocultivo, obs_urocultivo,
                orocultivo, obs_orocultivo,
                ex_nasal, obs_nasal
            ):
                # Limpia la tabla antes de agregar el nuevo registro
                self.result_table5.setRowCount(0)  # Limpia la tabla

                # Actualiza solo las primeras tres columnas en el QTableWidget para la fase 2b
                row_position = self.result_table5.rowCount()
                self.result_table5.insertRow(row_position)
                items_to_insert = [id_fase2b, id_paciente_5, fecha_registro_5]
                for column_index, item in enumerate(items_to_insert):
                    self.result_table5.setItem(row_position, column_index, QTableWidgetItem(str(item)))

                QMessageBox.information(self, "Éxito", "Información de Fase 2b agregada correctamente.")
                self.limpiar_fase2b()  # Limpia el formulario después de agregar
            else:
                QMessageBox.warning(self, "Error", "No se pudo agregar la información de Fase 2b.")
        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Error", f"Error al agregar fase 2b: {str(e)}")

#Es un condicional para verificar que RadioButton esta seleccionado y así restringir widgets dependiendo del tipo de paciente (receptor o donante)
    def buscar_fase2b_condicional(self):
        if self.id_receptor_f2b.isChecked():
            self.mostrar_f2b_receptor()
        elif self.id_donante_f2b.isChecked():
            self.mostrar_f2b_donante()

    def mostrar_f2b_receptor(self):
        self.limpiar_fase2b()
        results = self.db.buscar_f2b_receptor()  # Llama al método del objeto de conexión

        # Limpia la tabla antes de agregar nuevos datos
        self.result_table5.setRowCount(0)

        for row in results:
            row_position = self.result_table5.rowCount()
            self.result_table5.insertRow(row_position)
            
            for column, value in enumerate(row.values()):
                self.result_table5.setItem(row_position, column, QTableWidgetItem(str(value)))

    def mostrar_f2b_donante(self):
        if self.id_donante_f2b.isChecked():
            try:
                # Conecta a la base de datos
                self.db.conecta_base_datos()
                self.limpiar_fase2b()

                query = """
                SELECT f2b.* FROM fase_2b f2b
                JOIN donante d ON f2b.id_paciente_5 = d.id_donante
                """
                self.db.cursor.execute(query)
                results = self.db.cursor.fetchall()


                self.result_table5.setRowCount(0)

                for row in results:
                    row_position = self.result_table5.rowCount()
                    self.result_table5.insertRow(row_position)
                    
                    for column, value in enumerate(row.values()):
                        self.result_table5.setItem(row_position, column, QTableWidgetItem(str(value)))

            except Exception as e:
                QMessageBox.critical(self, "Error", f"Error al mostrar datos de donantes: {str(e)}")
            finally:
                if self.db.cursor:
                    self.db.cursor.close()
                if self.db.con:
                    self.db.con.close()

    def seleccionar_fase2b(self):
        selected_row = self.result_table5.currentRow()
        if selected_row >= 0:
            id_fase2b = self.result_table5.item(selected_row, 0).text()
            # Obtiene todos los datos de la fase 2b seleccionada de la base de datos
            fase2b_data = self.db.obtener_datos_fase2b(id_fase2b)

            if fase2b_data:
                # Función auxiliar para obtener valores float o establecerlos en 0
                def get_float_value(value):
                    try:
                        return float(value) if value not in [None, '', 'None'] else 0.0
                    except ValueError:
                        return 0.0

                # Función auxiliar para obtener valores int o establecerlos en 0
                def get_int_value(value):
                    try:
                        return int(value) if value not in [None, '', 'None'] else 0
                    except ValueError:
                        return 0

                # Limpia todos los campos antes de llenarlos con nueva información
                self.limpiar_fase2b()

                # Llena los widgets con la información
                self.id_fase2b.setText(str(fase2b_data['id_fase2b']))
                self.id_paciente_5.setText(str(fase2b_data['id_paciente_5']))
                self.fecha_registro_5.setDate(QDate.fromString(str(fase2b_data['fecha_registro_5']), "yyyy-MM-dd"))

                vih_value = get_int_value(fase2b_data['vih'])
                if vih_value in self.estado_pruebas.values():
                    self.vih.setCurrentText(next(key for key, value in self.estado_pruebas.items() if value == vih_value))

                hepa_b_value = get_int_value(fase2b_data['hepa_b'])
                if hepa_b_value in self.estado_pruebas.values():
                    self.hepa_b.setCurrentText(next(key for key, value in self.estado_pruebas.items() if value == hepa_b_value))

                hepa_c_value = get_int_value(fase2b_data['hepa_c'])
                if hepa_c_value in self.estado_pruebas.values():
                    self.hepa_c.setCurrentText(next(key for key, value in self.estado_pruebas.items() if value == hepa_c_value))

                vdrl_value = get_int_value(fase2b_data['vdrl'])
                if vdrl_value in self.estado_pruebas.values():
                    self.vdrl.setCurrentText(next(key for key, value in self.estado_pruebas.items() if value == vdrl_value))

                toxo_igg_value = get_int_value(fase2b_data['toxo_igg'])
                if toxo_igg_value in self.estado_pruebas1.values():
                    self.toxo_igg.setCurrentText(next(key for key, value in self.estado_pruebas1.items() if value == toxo_igg_value))

                toxo_igm_value = get_int_value(fase2b_data['toxo_igm'])
                if toxo_igm_value in self.estado_pruebas1.values():
                    self.toxo_igm.setCurrentText(next(key for key, value in self.estado_pruebas1.items() if value == toxo_igm_value))

                rubeo_igg_value = get_int_value(fase2b_data['rubeo_igg'])
                if rubeo_igg_value in self.estado_pruebas1.values():
                    self.rubeo_igg.setCurrentText(next(key for key, value in self.estado_pruebas1.items() if value == rubeo_igg_value))

                rubeo_igm_value = get_int_value(fase2b_data['rubeo_igm'])
                if rubeo_igm_value in self.estado_pruebas1.values():
                    self.rubeo_igm.setCurrentText(next(key for key, value in self.estado_pruebas1.items() if value == rubeo_igm_value))

                cmv_igg_value = get_int_value(fase2b_data['cmv_igg'])
                if cmv_igg_value in self.estado_pruebas1.values():
                    self.cmv_igg.setCurrentText(next(key for key, value in self.estado_pruebas1.items() if value == cmv_igg_value))

                cmv_igm_value = get_int_value(fase2b_data['cmv_igm'])
                if cmv_igm_value in self.estado_pruebas1.values():
                    self.cmv_igm.setCurrentText(next(key for key, value in self.estado_pruebas1.items() if value == cmv_igm_value))

                herpes_igg_value = get_int_value(fase2b_data['herpes_igg'])
                if herpes_igg_value in self.estado_pruebas1.values():
                    self.herpes_igg.setCurrentText(next(key for key, value in self.estado_pruebas1.items() if value == herpes_igg_value))

                herpes_igm_value = get_int_value(fase2b_data['herpes_igm'])
                if herpes_igm_value in self.estado_pruebas1.values():
                    self.herpes_igm.setCurrentText(next(key for key, value in self.estado_pruebas1.items() if value == herpes_igm_value))

                veb_igg_value = get_int_value(fase2b_data['veb_igg'])
                if veb_igg_value in self.estado_pruebas1.values():
                    self.veb_igg.setCurrentText(next(key for key, value in self.estado_pruebas1.items() if value == veb_igg_value))

                veb_igm_value = get_int_value(fase2b_data['veb_igm'])
                if veb_igm_value in self.estado_pruebas1.values():
                    self.veb_igm.setCurrentText(next(key for key, value in self.estado_pruebas1.items() if value == veb_igm_value))

                cuantiferon_value = get_int_value(fase2b_data['cuantiferon'])
                if cuantiferon_value in self.estado_pruebas.values():
                    self.cuantiferon.setCurrentText(next(key for key, value in self.estado_pruebas.items() if value == cuantiferon_value))

                self.obs_cuantiferon.setPlainText(str(fase2b_data['obs_cuantiferon']) if fase2b_data['obs_cuantiferon'] else "")

                self.tgo.setValue(get_float_value(fase2b_data['tgo']))
                self.tgp.setValue(get_float_value(fase2b_data['tgp']))
                self.bt.setValue(get_float_value(fase2b_data['bt']))
                self.bd.setValue(get_float_value(fase2b_data['bd']))
                self.bi.setValue(get_float_value(fase2b_data['bi']))
                self.dhl.setValue(get_float_value(fase2b_data['dhl']))
                self.fa.setValue(get_float_value(fase2b_data['fa']))
                self.albumina.setValue(get_float_value(fase2b_data['albumina']))
                self.acido_urico.setValue(get_float_value(fase2b_data['acido_urico']))
                self.perfil_hdl.setValue(get_float_value(fase2b_data['perfil_hdl']))
                self.perfil_ldl.setValue(get_float_value(fase2b_data['perfil_ldl']))
                self.perfil_ct.setValue(get_float_value(fase2b_data['perfil_ct']))
                self.perfil_tg.setValue(get_float_value(fase2b_data['perfil_tg']))
                self.perfil_vldl.setValue(get_float_value(fase2b_data['perfil_vldl']))
                self.sodio.setValue(get_float_value(fase2b_data['sodio']))
                self.potasio.setValue(get_float_value(fase2b_data['potasio']))
                self.fosforo.setValue(get_float_value(fase2b_data['fosforo']))
                self.calcio.setValue(get_float_value(fase2b_data['calcio']))
                self.magnesio.setValue(get_float_value(fase2b_data['magnesio']))
                self.pth.setValue(get_float_value(fase2b_data['pth']))
                self.tsh.setValue(get_float_value(fase2b_data['tsh']))
                self.t4_libre.setValue(get_float_value(fase2b_data['t4_libre']))
                self.c3.setValue(get_float_value(fase2b_data['c3']))
                self.c4.setValue(get_float_value(fase2b_data['c4']))
                self.bnp.setValue(get_float_value(fase2b_data['bnp']))
                self.tp.setValue(get_float_value(fase2b_data['tp']))
                self.tpt.setValue(get_float_value(fase2b_data['tpt']))
                self.inr.setValue(get_float_value(fase2b_data['inr']))

                nivel_ecg_value = get_int_value(fase2b_data['nivel_ecg'])
                if nivel_ecg_value in self.nivel_ecg_map.values():
                    self.nivel_ecg.setCurrentText(next(key for key, value in self.nivel_ecg_map.items() if value == nivel_ecg_value))

                self.obs_ecg.setPlainText(str(fase2b_data['obs_ecg']) if fase2b_data['obs_ecg'] else "")

                anti_dna_value = get_int_value(fase2b_data['anti_dna'])
                if anti_dna_value in self.estado_pruebas.values():
                    self.anti_dna.setCurrentText(next(key for key, value in self.estado_pruebas.items() if value == anti_dna_value))

                ana_value = get_int_value(fase2b_data['ana'])
                if ana_value in self.estado_pruebas.values():
                    self.ana.setCurrentText(next(key for key, value in self.estado_pruebas.items() if value == ana_value))

                b2_glicoprot_value = get_int_value(fase2b_data['b2_glicoprot'])
                if b2_glicoprot_value in self.estado_pruebas.values():
                    self.b2_glicoprot.setCurrentText(next(key for key, value in self.estado_pruebas.items() if value == b2_glicoprot_value))

                anticoag_lupico_value = get_int_value(fase2b_data['anticoag_lupico'])
                if anticoag_lupico_value in self.estado_pruebas.values():
                    self.anticoag_lupico.setCurrentText(next(key for key, value in self.estado_pruebas.items() if value == anticoag_lupico_value))

                anticardio_igg_igm_value = get_int_value(fase2b_data['anticardio_igg_igm'])
                if anticardio_igg_igm_value in self.estado_pruebas.values():
                    self.anticardio_igg_igm.setCurrentText(next(key for key, value in self.estado_pruebas.items() if value == anticardio_igg_igm_value))

                p_anca_value = get_int_value(fase2b_data['p_anca'])
                if p_anca_value in self.estado_pruebas.values():
                    self.p_anca.setCurrentText(next(key for key, value in self.estado_pruebas.items() if value == p_anca_value))

                c_anca_value = get_int_value(fase2b_data['c_anca'])
                if c_anca_value in self.estado_pruebas.values():
                    self.c_anca.setCurrentText(next(key for key, value in self.estado_pruebas.items() if value == c_anca_value))

                urocultivo_value = get_int_value(fase2b_data['urocultivo'])
                if urocultivo_value in self.estado_pruebas.values():
                    self.urocultivo.setCurrentText(next(key for key, value in self.estado_pruebas.items() if value == urocultivo_value))

                self.obs_urocultivo.setPlainText(str(fase2b_data['obs_urocultivo']) if fase2b_data['obs_urocultivo'] else "")

                orocultivo_value = get_int_value(fase2b_data['orocultivo'])
                if orocultivo_value in self.estado_pruebas.values():
                    self.orocultivo.setCurrentText(next(key for key, value in self.estado_pruebas.items() if value == orocultivo_value))

                self.obs_orocultivo.setPlainText(str(fase2b_data['obs_orocultivo']) if fase2b_data['obs_orocultivo'] else "")

                ex_nasal_value = get_int_value(fase2b_data['ex_nasal'])
                if ex_nasal_value in self.estado_pruebas.values():
                    self.ex_nasal.setCurrentText(next(key for key, value in self.estado_pruebas.items() if value == ex_nasal_value))

                self.obs_nasal.setPlainText(str(fase2b_data['obs_nasal']) if fase2b_data['obs_nasal'] else "")

            else:
                QMessageBox.warning(self, "Error", "No se pudo obtener la información de la Fase 2b.")


    def limpiar_fase2b(self):
        # Limpia los campos de texto y restablecer valores predeterminados para la fase 2b
        self.id_fase2b.clear()
        self.id_paciente_5.clear() 
        fecha_default = QDate(2000, 1, 1)
        self.fecha_registro_5.setDate(fecha_default)
        self.vih.setCurrentIndex(0) 
        self.hepa_b.setCurrentIndex(0) 
        self.hepa_c.setCurrentIndex(0)
        self.vdrl.setCurrentIndex(0) 
        self.toxo_igg.setCurrentIndex(0) 
        self.toxo_igm.setCurrentIndex(0) 
        self.rubeo_igg.setCurrentIndex(0) 
        self.rubeo_igm.setCurrentIndex(0) 
        self.cmv_igg.setCurrentIndex(0) 
        self.cmv_igm.setCurrentIndex(0) 
        self.herpes_igg.setCurrentIndex(0) 
        self.herpes_igm.setCurrentIndex(0) 
        self.veb_igg.setCurrentIndex(0) 
        self.veb_igm.setCurrentIndex(0) 
        self.cuantiferon.setCurrentIndex(0) 
        self.obs_cuantiferon.clear() 
        self.tgo.setValue(0.0) 
        self.tgp.setValue(0.0) 
        self.bt.setValue(0.0) 
        self.bd.setValue(0.0) 
        self.bi.setValue(0.0) 
        self.dhl.setValue(0.0) 
        self.fa.setValue(0.0) 
        self.albumina.setValue(0.0) 
        self.acido_urico.setValue(0.0) 
        self.perfil_hdl.setValue(0.0) 
        self.perfil_ldl.setValue(0.0) 
        self.perfil_ct.setValue(0.0) 
        self.perfil_tg.setValue(0.0) 
        self.perfil_vldl.setValue(0.0) 
        self.sodio.setValue(0.0) 
        self.potasio.setValue(0.0) 
        self.fosforo.setValue(0.0) 
        self.calcio.setValue(0.0) 
        self.magnesio.setValue(0.0) 
        self.pth.setValue(0.0) 
        self.tsh.setValue(0.0) 
        self.t4_libre.setValue(0.0) 
        self.c3.setValue(0.0) 
        self.c4.setValue(0.0) 
        self.bnp.setValue(0.0) 
        self.tp.setValue(0.0) 
        self.tpt.setValue(0.0)
        self.inr.setValue(0.0) 
        self.nivel_ecg.setCurrentIndex(0) 
        self.obs_ecg.clear() 
        self.anti_dna.setCurrentIndex(0) 
        self.ana.setCurrentIndex(0) 
        self.b2_glicoprot.setCurrentIndex(0) 
        self.anticoag_lupico.setCurrentIndex(0) 
        self.anticardio_igg_igm.setCurrentIndex(0) 
        self.p_anca.setCurrentIndex(0) 
        self.c_anca.setCurrentIndex(0)
        self.urocultivo.setCurrentIndex(0) 
        self.obs_urocultivo.clear() 
        self.orocultivo.setCurrentIndex(0)  
        self.obs_orocultivo.clear() 
        self.ex_nasal .setCurrentIndex (0)  
        self.obs_nasal.clear() 


    def actualizar_fase2b(self):
        try:
            # Solicita y valida la contraseña
            contraseña = self.solicitar_contraseña()
            if contraseña is None:
                return
                
            if not self.validar_contraseña(contraseña):
                QtWidgets.QMessageBox.warning(self, "Error", "Contraseña incorrecta")
                return
            # Recoge datos de los campos de entrada para la fase 2b
            id_fase2b = self.id_fase2b.text()
            id_paciente_5 = self.id_paciente_5.text()
            fecha_registro_5 = self.fecha_registro_5.date().toString("yyyy-MM-dd")

            vih = self.estado_pruebas.get(self.vih.currentText(), None)
            hepa_b = self.estado_pruebas.get(self.hepa_b.currentText(), None)
            hepa_c = self.estado_pruebas.get(self.hepa_c.currentText(), None)
            vdrl = self.estado_pruebas.get(self.vdrl.currentText(), None)
            toxo_igg = self.estado_pruebas1.get(self.toxo_igg.currentText(), None)
            toxo_igm = self.estado_pruebas1.get(self.toxo_igm.currentText(), None)
            rubeo_igg = self.estado_pruebas1.get(self.rubeo_igg.currentText(), None)
            rubeo_igm = self.estado_pruebas1.get(self.rubeo_igm.currentText(), None)
            cmv_igg = self.estado_pruebas1.get(self.cmv_igg.currentText(), None)
            cmv_igm = self.estado_pruebas1.get(self.cmv_igm.currentText(), None)
            herpes_igg = self.estado_pruebas1.get(self.herpes_igg.currentText(), None)
            herpes_igm = self.estado_pruebas1.get(self.herpes_igm.currentText(), None)
            veb_igg = self.estado_pruebas1.get(self.veb_igg.currentText(), None)
            veb_igm = self.estado_pruebas1.get(self.veb_igm.currentText(), None)
            cuantiferon = self.estado_pruebas.get(self.cuantiferon.currentText(), None)
            obs_cuantiferon = self.obs_cuantiferon.toPlainText()
            tgo = self.tgo.value()
            tgp = self.tgp.value()
            bt = self.bt.value()
            bd = self.bd.value()
            bi = self.bi.value()
            dhl = self.dhl.value()
            fa = self.fa.value()
            albumina = self.albumina.value()
            acido_urico = self.acido_urico.value()
            perfil_hdl= self.perfil_hdl.value()
            perfil_ldl= self.perfil_ldl.value()
            perfil_ct= self.perfil_ct.value()
            perfil_tg= self.perfil_tg.value()
            perfil_vldl= self.perfil_vldl.value()
            sodio= self.sodio.value()
            potasio= self.potasio.value()
            fosforo= self.fosforo.value()
            calcio= self.calcio.value()
            magnesio= self.magnesio.value()
            pth= self.pth.value()
            tsh= self.tsh.value()
            t4_libre= self.t4_libre.value()
            c3= self.c3.value()
            c4= self.c4.value()
            bnp= self.bnp.value()
            tp= self.tp.value()
            tpt= self.tpt.value()
            inr= self.inr.value()

            nivel_ecg= self.nivel_ecg_map.get(self.nivel_ecg.currentText(), None)
            obs_ecg= self.obs_ecg.toPlainText()
            anti_dna= self.estado_pruebas.get(self.anti_dna.currentText(), None)
            ana= self.estado_pruebas.get(self.ana.currentText(), None)
            b2_glicoprot= self.estado_pruebas.get(self.b2_glicoprot.currentText(), None)
            anticoag_lupico= self.estado_pruebas.get(self.anticoag_lupico.currentText(), None)
            anticardio_igg_igm= self.estado_pruebas.get(self.anticardio_igg_igm.currentText(), None)
            p_anca= self.estado_pruebas.get(self.p_anca.currentText(), None)
            c_anca= self.estado_pruebas.get(self.c_anca.currentText(), None)

            urocultivo= self.estado_pruebas.get(self.urocultivo.currentText(), None)
            obs_urocultivo= self.obs_urocultivo.toPlainText()

            orocultivo= self.estado_pruebas.get(self.orocultivo.currentText(), None)
            obs_orocultivo= self.obs_orocultivo.toPlainText()

            ex_nasal= self.estado_pruebas.get(self.ex_nasal.currentText(), None)
            obs_nasal=	self.obs_nasal.toPlainText()

            # Llama al método de la clase de conexión para actualizar los datos de la fase 2b
            if (self.db.actualizar_datos_5(
                    id_fase2b,id_paciente_5 ,fecha_registro_5,
                    vih ,hepa_b ,hepa_c ,vdrl,
                    toxo_igg ,toxo_igm ,rubeo_igg ,rubeo_igm,
                    cmv_igg ,cmv_igm ,herpes_igg ,herpes_igm,
                    veb_igg ,veb_igm ,cuantiferon,
                    obs_cuantiferon,
                    tgo,tgp ,bt ,bd ,bi,
                    dhl ,fa ,albumina ,acido_urico,
                    perfil_hdl ,perfil_ldl ,perfil_ct,
                    perfil_tg ,perfil_vldl,
                    sodio,potasio,fosforo,
                    calcio ,magnesio,
                    pth,tsh,t4_libre,
                    c3,c4,bnp,
                    tp,tpt,inr,
                    nivel_ecg ,obs_ecg ,
                    anti_dna ,ana ,
                    b2_glicoprot ,anticoag_lupico ,
                    anticardio_igg_igm ,
                    p_anca,c_anca,
                    urocultivo, obs_urocultivo,
                    orocultivo, obs_orocultivo ,
                    ex_nasal ,obs_nasal 
                )):
                QMessageBox.information(self,"Éxito","Información actualizada correctamente.")
                # Limpia el formulario después de actualizar
                self.limpiar_fase2b()
                # Vuelve a cargar los datos para mostrar todos
                self.buscar_fase2b_condicional()

        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Error", f"Error al actualizar fase 2b: {str(e)}")
    
    def eliminar_fase2b(self):
        #Elimina un registro de fase 2b después de validar la contraseña
        try:
            # Solicita y valida la contraseña
            contraseña = self.solicitar_contraseña()
            if contraseña is None:
                return
                
            if not self.validar_contraseña(contraseña):
                QtWidgets.QMessageBox.warning(self, "Error", "Contraseña incorrecta")
                return
            
            # Si la contraseña es válida, procede con eliminar
            selected_row = self.result_table5.currentRow()
            
            if selected_row >= 0:
                # Obtiene el ID de la fase 2b de la primera columna
                id_fase2b = self.result_table5.item(selected_row, 0).text()
                
                reply = QtWidgets.QMessageBox.question(
                    self, 
                    'Confirmar eliminación', 
                    '¿Está seguro de que desea eliminar este registro de fase 2b?',
                    QtWidgets.QMessageBox.Yes | QtWidgets.QMessageBox.No, 
                    QtWidgets.QMessageBox.No
                )
                
                if reply == QtWidgets.QMessageBox.Yes:
                    if self.db.eliminar_datos_5(id_fase2b):
                        self.result_table5.removeRow(selected_row)
                        QtWidgets.QMessageBox.information(self, "Éxito", "Registro de fase 2b eliminado correctamente.")
                        self.limpiar_fase2b()
                    else:
                        QtWidgets.QMessageBox.warning(self, "Error", "No se pudo eliminar el registro de fase 2b.")
            else:
                QtWidgets.QMessageBox.warning(self, "Advertencia", "Por favor, seleccione un registro de fase 2b para eliminar.")
                
        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Error", f"Error al eliminar fase 2b: {str(e)}")

################## FASE 3 - Funciones básicas para interactuar con los datos (buscar,agregar,limpiar,eliminar datos)

    def agregar_fase3(self):
        try:
            # Solicita y valida la contraseña
            contraseña = self.solicitar_contraseña()
            if contraseña is None:
                return
                
            if not self.validar_contraseña(contraseña):
                QtWidgets.QMessageBox.warning(self, "Error", "Contraseña incorrecta")
                return
            # Recoge datos de los campos de entrada para la fase 3
            id_fase3 = self.id_fase3.text()
            id_paciente_6 = self.id_paciente_6.text()
            fecha_registro_6 = self.fecha_registro_6.date().toString("yyyy-MM-dd")
            est_antig_prost = self.estado_pruebas2.get(self.est_antig_prost.currentText(), None)
            obs_antig_prost = self.obs_antig_prost.toPlainText()
            est_hgc_sbeta = self.estado_pruebas2.get(self.est_hgc_sbeta.currentText(), None)
            obs_hgc_sbeta = self.obs_hgc_sbeta.toPlainText()
            est_pap = self.estado_pruebas2.get(self.est_pap.currentText(), None)
            obs_pap = self.obs_pap.toPlainText()
            est_mamo = self.estado_pruebas2.get(self.est_mamo.currentText())
            obs_mamo = self.obs_mamo.toPlainText()
            est_guay_hec = self.estado_pruebas2.get(self.est_guay_hec.currentText(), None)
            obs_guay_hec = self.obs_guay_hec.toPlainText()
            est_endo_colon = self.estado_pruebas2.get(self.est_endo_colon.currentText(), None)
            obs_endo_colon = self.obs_endo_colon.toPlainText()
            est_rxt = self.estado_pruebas2.get(self.est_rxt.currentText(), None)
            obs_rxt = self.obs_rxt.toPlainText()
            est_rx_spn = self.estado_pruebas2.get(self.est_rx_spn.currentText(), None)
            obs_rx_spn = self.obs_rx_spn.toPlainText()
            est_cisto = self.estado_pruebas2.get(self.est_cisto.currentText(), None)
            obs_cisto = self.obs_cisto.toPlainText()
            est_usg_vesi = self.estado_pruebas2.get(self.est_usg_vesi.currentText(), None)
            obs_usg_vesi = self.obs_usg_vesi.toPlainText()
            est_eco_trans = self.estado_pruebas2.get(self.est_eco_trans.currentText(), None)
            obs_eco_trans = self.obs_eco_trans.toPlainText()
            est_eco_trans_dm = self.estado_pruebas2.get(self.est_eco_trans_dm.currentText(), None)
            obs_eco_trans_dm = self.obs_eco_trans_dm.toPlainText()
            est_dop_iliac = self.estado_pruebas2.get(self.est_dop_iliac.currentText(), None)
            obs_dop_iliac = self.obs_dop_iliac.toPlainText()
            est_dop_art = self.estado_pruebas2.get(self.est_dop_art.currentText(), None)
            obs_dop_art = self.obs_dop_art.toPlainText()
            est_2donantes = self.estado_pruebas2.get(self.est_2donantes.currentText(), None)
            obs_2donantes = self.obs_2donantes.toPlainText()
            est_pielograma = self.estado_pruebas2.get(self.est_pielograma.currentText(), None)
            obs_pielograma = self.obs_pielograma.toPlainText()

            # Valida que todos los campos obligatorios estén llenos
            if not all([id_fase3, id_paciente_6, fecha_registro_6]):
                QMessageBox.warning(self, "Advertencia", "Los campos ID Fase 3, ID Paciente y Fecha de Registro son obligatorios.")
                return
            
            # Verificar si el id_paciente_3 existe en las tablas receptor o donante
            if not (self.db.existe_paciente_en_receptor_o_donante(id_paciente_6)):
                QMessageBox.warning(self, "Error", f"El ID del paciente {id_paciente_6} no se encuentra en las tablas receptor o donante.")
                return

            # Llama al método de la clase de conexión para agregar los datos de la fase 3
            if self.db.agregar_datos_6(
                id_fase3, id_paciente_6, fecha_registro_6,
                est_antig_prost, obs_antig_prost,
                est_hgc_sbeta, obs_hgc_sbeta,
                est_pap, obs_pap,
                est_mamo, obs_mamo,
                est_guay_hec, obs_guay_hec,
                est_endo_colon, obs_endo_colon,
                est_rxt, obs_rxt,
                est_rx_spn, obs_rx_spn,
                est_cisto, obs_cisto,
                est_usg_vesi, obs_usg_vesi,
                est_eco_trans, obs_eco_trans,
                est_eco_trans_dm, obs_eco_trans_dm,
                est_dop_iliac, obs_dop_iliac,
                est_dop_art, obs_dop_art,
                est_2donantes, obs_2donantes,
                est_pielograma, obs_pielograma
            ):
                # Limpia la tabla antes de agregar el nuevo registro
                self.result_table6.setRowCount(0)  # Limpia la tabla

                # Actualiza solo las primeras tres columnas en el QTableWidget para la fase 3
                row_position = self.result_table6.rowCount()
                self.result_table6.insertRow(row_position)
                items_to_insert = [id_fase3, id_paciente_6, fecha_registro_6]
                for column_index, item in enumerate(items_to_insert):
                    self.result_table6.setItem(row_position, column_index, QTableWidgetItem(str(item)))

                QMessageBox.information(self, "Éxito", "Información de Fase 3 agregada correctamente.")
                self.limpiar_fase3()  # Limpia el formulario después de agregar
            else:
                QMessageBox.warning(self, "Error", "No se pudo agregar la información de Fase 3.")

        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Error", f"Error al agregar fase 3: {str(e)}")
        
#Es un condicional para verificar que RadioButton esta seleccionado y así restringir widgets dependiendo del tipo de paciente (receptor o donante)
    def buscar_fase3_condicional(self):
        if self.id_receptor_f3.isChecked():
            self.mostrar_f3_receptor()
        elif self.id_donante_f2b.isChecked():
            self.mostrar_f3_donante()

    def mostrar_f3_receptor(self):
        self.limpiar_fase3()
        results = self.db.buscar_f3_receptor()  # Llama al método del objeto de conexión

        # Limpia la tabla antes de agregar nuevos datos
        self.result_table6.setRowCount(0)

        for row in results:
            row_position = self.result_table6.rowCount()
            self.result_table6.insertRow(row_position)
            
            for column, value in enumerate(row.values()):
                self.result_table6.setItem(row_position, column, QTableWidgetItem(str(value)))

    def mostrar_f3_donante(self):
        if self.id_donante_f3.isChecked():
            try:
                # Conecta a la base de datos
                self.limpiar_fase3()
                self.db.conecta_base_datos()

                query = """
                SELECT f3.* FROM fase_3 f3
                JOIN donante d ON f3.id_paciente_6 = d.id_donante
                """
                self.db.cursor.execute(query)
                results = self.db.cursor.fetchall()

                self.result_table6.setRowCount(0)

                for row in results:
                    row_position = self.result_table6.rowCount()
                    self.result_table6.insertRow(row_position)
                    
                    for column, value in enumerate(row.values()):
                        self.result_table6.setItem(row_position, column, QTableWidgetItem(str(value)))

            except Exception as e:
                QMessageBox.critical(self, "Error", f"Error al mostrar datos de donantes: {str(e)}")
            finally:
                if self.db.cursor:
                    self.db.cursor.close()
                if self.db.con:
                    self.db.con.close()

    def limpiar_fase3(self):
        # Limpia los campos de texto de la fase 3
        self.id_fase3.clear()
        self.id_paciente_6.clear()

        fecha_default = QDate(2000, 1, 1)
        self.fecha_registro_6.setDate(fecha_default)

        # Restablece los ComboBox a su índice inicial
        self.est_antig_prost.setCurrentIndex(0)
        self.est_hgc_sbeta.setCurrentIndex(0)
        self.est_pap.setCurrentIndex(0)
        self.est_mamo.setCurrentIndex(0)
        self.est_guay_hec.setCurrentIndex(0)
        self.est_endo_colon.setCurrentIndex(0)
        self.est_rxt.setCurrentIndex(0)
        self.est_rx_spn.setCurrentIndex(0)
        self.est_cisto.setCurrentIndex(0)
        self.est_usg_vesi.setCurrentIndex(0)
        self.est_eco_trans.setCurrentIndex(0)
        self.est_eco_trans_dm.setCurrentIndex(0)
        self.est_dop_iliac.setCurrentIndex(0)
        self.est_dop_art.setCurrentIndex(0)
        self.est_2donantes.setCurrentIndex(0)
        self.est_pielograma.setCurrentIndex(0)
        self.obs_antig_prost.clear()
        self.obs_hgc_sbeta.clear()
        self.obs_pap.clear()
        self.obs_mamo.clear()
        self.obs_guay_hec.clear()
        self.obs_endo_colon.clear()
        self.obs_rxt.clear()
        self.obs_rx_spn.clear()
        self.obs_cisto.clear()
        self.obs_usg_vesi.clear()
        self.obs_eco_trans.clear()
        self.obs_eco_trans_dm.clear()
        self.obs_dop_iliac.clear()
        self.obs_dop_art.clear()
        self.obs_2donantes.clear()
        self.obs_pielograma.clear()

    def seleccionar_fase3(self):
        selected_row = self.result_table6.currentRow()
        if selected_row >= 0:
            id_fase3 = self.result_table6.item(selected_row, 0).text()
            # Obtiene todos los datos de la fase 3 seleccionada de la base de datos
            fase3_data = self.db.obtener_datos_fase3(id_fase3)

            if fase3_data:

                # Función auxiliar para obtener valores int o establecerlos en 0
                def get_int_value(value):
                    try:
                        return int(value) if value not in [None, '', 'None'] else 0
                    except ValueError:
                        return 0

                # Limpia todos los campos antes de llenarlos con nueva información
                self.limpiar_fase3()

                # Llena los widgets con la información
                self.id_fase3.setText(str(fase3_data['id_fase3']))
                self.id_paciente_6.setText(str(fase3_data['id_paciente_6']))
                self.fecha_registro_6.setDate(QDate.fromString(str(fase3_data['fecha_registro_6']), "yyyy-MM-dd"))

                est_antig_prost_value = get_int_value(fase3_data['est_antig_prost'])
                if est_antig_prost_value in self.estado_pruebas2.values():
                    self.est_antig_prost.setCurrentText(next(key for key, value in self.estado_pruebas2.items() if value == est_antig_prost_value))
                self.obs_antig_prost.setPlainText(str(fase3_data['obs_antig_prost']) if fase3_data['obs_antig_prost'] else "")

                est_hgc_sbeta_value = get_int_value(fase3_data['est_hgc_sbeta'])
                if est_hgc_sbeta_value in self.estado_pruebas2.values():
                    self.est_hgc_sbeta.setCurrentText(next(key for key, value in self.estado_pruebas2.items() if value == est_hgc_sbeta_value))
                self.obs_hgc_sbeta.setPlainText(str(fase3_data['obs_hgc_sbeta']) if fase3_data['obs_hgc_sbeta'] else "")

                est_pap_value = get_int_value(fase3_data['est_pap'])
                if est_pap_value in self.estado_pruebas2.values():
                    self.est_pap.setCurrentText(next(key for key, value in self.estado_pruebas2.items() if value == est_pap_value))
                self.obs_pap.setPlainText(str(fase3_data['obs_pap']) if fase3_data['obs_pap'] else "")

                est_mamo_value = get_int_value(fase3_data['est_mamo'])
                if est_mamo_value in self.estado_pruebas2.values():
                    self.est_mamo.setCurrentText(next(key for key, value in self.estado_pruebas2.items() if value == est_mamo_value))
                self.obs_mamo.setPlainText(str(fase3_data['obs_mamo']) if fase3_data['obs_mamo'] else "")

                est_guay_hec_value = get_int_value(fase3_data['est_guay_hec'])
                if est_guay_hec_value in self.estado_pruebas2.values():
                    self.est_guay_hec.setCurrentText(next(key for key, value in self.estado_pruebas2.items() if value == est_guay_hec_value))
                self.obs_guay_hec.setPlainText(str(fase3_data['obs_guay_hec']) if fase3_data['obs_guay_hec'] else "")

                est_endo_colon_value = get_int_value(fase3_data['endo_colon'])
                if est_endo_colon_value in self.estado_pruebas2.values():
                    self.est_endo_colon.setCurrentText(next(key for key, value in self.estado_pruebas2.items() if value == est_endo_colon_value))
                self.obs_endo_colon.setPlainText(str(fase3_data['obs_endo_colon']) if fase3_data['obs_endo_colon'] else "")

                est_rxt_value = get_int_value(fase3_data['est_rxt'])
                if est_rxt_value in self.estado_pruebas2.values():
                    self.est_rxt.setCurrentText(next(key for key, value in self.estado_pruebas2.items() if value == est_rxt_value))
                self.obs_rxt.setPlainText(str(fase3_data['obs_rxt']) if fase3_data['obs_rxt'] else "")

                est_rx_spn_value = get_int_value(fase3_data['est_rx_spn'])
                if est_rx_spn_value in self.estado_pruebas2.values():
                    self.est_rx_spn.setCurrentText(next(key for key, value in self.estado_pruebas2.items() if value == est_rx_spn_value))
                self.obs_rx_spn.setPlainText(str(fase3_data['obs_rx_spn']) if fase3_data['obs_rx_spn'] else "")

                est_cisto_value = get_int_value(fase3_data['est_cisto'])
                if est_cisto_value in self.estado_pruebas2.values():
                    self.est_cisto.setCurrentText(next(key for key, value in self.estado_pruebas2.items() if value == est_cisto_value))
                self.obs_cisto.setPlainText(str(fase3_data['obs_cisto']) if fase3_data['obs_cisto'] else "")

                est_usg_vesi_value = get_int_value(fase3_data['est_usg_vesi'])
                if est_usg_vesi_value in self.estado_pruebas2.values():
                    self.est_usg_vesi.setCurrentText(next(key for key, value in self.estado_pruebas2.items() if value == est_usg_vesi_value))
                self.obs_usg_vesi.setPlainText(str(fase3_data['obs_usg_vesi']) if fase3_data['obs_usg_vesi'] else "")

                est_eco_trans_value = get_int_value(fase3_data['est_eco_trans'])
                if est_eco_trans_value in self.estado_pruebas2.values():
                    self.est_eco_trans.setCurrentText(next(key for key, value in self.estado_pruebas2.items() if value == est_eco_trans_value))
                self.obs_eco_trans.setPlainText(str(fase3_data['obs_eco_trans']) if fase3_data['obs_eco_trans'] else "")

                est_eco_trans_dm_value = get_int_value(fase3_data['est_eco_trans_dm'])
                if est_eco_trans_dm_value in self.estado_pruebas2.values():
                    self.est_eco_trans_dm.setCurrentText(next(key for key, value in self.estado_pruebas2.items() if value == est_eco_trans_dm_value))
                self.obs_eco_trans_dm.setPlainText(str(fase3_data['obs_eco_trans_dm']) if fase3_data['obs_eco_trans_dm'] else "")

                est_dop_iliac_value = get_int_value(fase3_data['est_dop_iliac'])
                if est_dop_iliac_value in self.estado_pruebas2.values():
                    self.est_dop_iliac.setCurrentText(next(key for key, value in self.estado_pruebas2.items() if value == est_dop_iliac_value))
                self.obs_dop_iliac.setPlainText(str(fase3_data['obs_dop_iliac']) if fase3_data['obs_dop_iliac'] else "")

                est_dop_art_value = get_int_value(fase3_data['est_dop_art'])
                if est_dop_art_value in self.estado_pruebas2.values():
                    self.est_dop_art.setCurrentText(next(key for key, value in self.estado_pruebas2.items() if value == est_dop_art_value))
                self.obs_dop_art.setPlainText(str(fase3_data['obs_dop_art']) if fase3_data['obs_dop_art'] else "")

                est_2donantes_value = get_int_value(fase3_data['est_2donantes'])
                if est_2donantes_value in self.estado_pruebas2.values():
                    self.est_2donantes.setCurrentText(next(key for key, value in self.estado_pruebas2.items() if value == est_2donantes_value))
                self.obs_2donantes.setPlainText(str(fase3_data['obs_2donantes']) if fase3_data['obs_2donantes'] else "")

                est_pielograma_value = get_int_value(fase3_data['est_pielograma'])
                if est_pielograma_value in self.estado_pruebas2.values():
                    self.est_pielograma.setCurrentText(next(key for key, value in self.estado_pruebas2.items() if value == est_pielograma_value))
                self.obs_pielograma.setPlainText(str(fase3_data['obs_pielograma']) if fase3_data['obs_pielograma'] else "")

            else:
                QMessageBox.warning(self, "Error", "No se pudo obtener la información de la Fase 3.")

    def actualizar_fase3(self):
        try:
            # Solicita y valida la contraseña
            contraseña = self.solicitar_contraseña()
            if contraseña is None:
                return
                
            if not self.validar_contraseña(contraseña):
                QtWidgets.QMessageBox.warning(self, "Error", "Contraseña incorrecta")
                return
            # Recoge datos de los campos de entrada para la fase 3
            id_fase3 = self.id_fase3.text()
            id_paciente_6 = self.id_paciente_6.text()
            fecha_registro_6 = self.fecha_registro_6.date().toString("yyyy-MM-dd")

            # Obtiene los valores enteros correspondientes a los estados seleccionados
            est_antig_prost = self.estado_pruebas2.get(self.est_antig_prost.currentText(), None)
            est_hgc_sbeta = self.estado_pruebas2.get(self.est_hgc_sbeta.currentText(), None)
            est_pap = self.estado_pruebas2.get(self.est_pap.currentText(), None)
            est_mamo = self.estado_pruebas2.get(self.est_mamo.currentText(), None)
            est_guay_hec = self.estado_pruebas2.get(self.est_guay_hec.currentText(), None)
            est_endo_colon = self.estado_pruebas2.get(self.endo_colon.currentText(), None)
            est_rxt = self.estado_pruebas2.get(self.est_rxt.currentText(), None)
            est_rx_spn = self.estado_pruebas2.get(self.est_rx_spn.currentText(), None)
            est_cisto = self.estado_pruebas2.get(self.est_cisto.currentText(), None)
            est_usg_vesi = self.estado_pruebas2.get(self.est_usg_vesi.currentText(), None)
            est_eco_trans = self.estado_pruebas2.get(self.est_eco_trans.currentText(), None)
            est_eco_trans_dm = self.estado_pruebas2.get(self.est_eco_trans_dm.currentText(), None)
            est_dop_iliac = self.estado_pruebas2.get(self.est_dop_iliac.currentText(), None)
            est_dop_art = self.estado_pruebas2.get(self.est_dop_art.currentText(), None)
            est_2donantes = self.estado_pruebas2.get(self.est_2donantes.currentText(), None)
            est_pielograma = self.estado_pruebas2.get(self.est_pielograma.currentText(), None)

            # Observaciones
            obs_antig_prost = self.obs_antig_prost.toPlainText() or None
            obs_hgc_sbeta = self.obs_hgc_sbeta.toPlainText() or None
            obs_pap = self.obs_pap.toPlainText() or None
            obs_mamo = self.obs_mamo.toPlainText() or None
            obs_guay_hec = self.obs_guay_hec.toPlainText() or None
            obs_endo_colon = self.obs_endo_colon.toPlainText() or None
            obs_rxt = self.obs_rxt.toPlainText() or None
            obs_rx_spn = self.obs_rx_spn.toPlainText() or None
            obs_cisto = self.obs_cisto.toPlainText() or None
            obs_usg_vesi = self.obs_usg_vesi.toPlainText() or None
            obs_eco_trans = self.obs_eco_trans.toPlainText() or None
            obs_eco_trans_dm = self.obs_eco_trans_dm.toPlainText() or None
            obs_dop_iliac = self.obs_dop_iliac.toPlainText() or None
            obs_dop_art = self.obs_dop_art.toPlainText() or None
            obs_2donantes = self.obs_2donantes.toPlainText() or None
            obs_pielograma = self.obs_pielograma.toPlainText() or None

            # Llama al método de la clase de conexión para actualizar los datos de la fase 3
            if self.db.actualizar_datos_6(
                id_fase3, id_paciente_6, fecha_registro_6,
                est_antig_prost, est_hgc_sbeta, est_pap, est_mamo,
                est_guay_hec, est_endo_colon, est_rxt, est_rx_spn,
                est_cisto, est_usg_vesi, est_eco_trans, est_eco_trans_dm,
                est_dop_iliac, est_dop_art, est_2donantes, est_pielograma,
                obs_antig_prost, obs_hgc_sbeta, obs_pap, obs_mamo,
                obs_guay_hec, obs_endo_colon, obs_rxt, obs_rx_spn,
                obs_cisto, obs_usg_vesi, obs_eco_trans, obs_eco_trans_dm,
                obs_dop_iliac, obs_dop_art, obs_2donantes, obs_pielograma
            ):
                QMessageBox.information(self, "Éxito", "Información actualizada correctamente.")
                self.limpiar_fase3()  # Limpia el formulario después de actualizar
                self.buscar_fase3_condicional()  # Vuelve a cargar los datos para mostrar todos

        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Error", f"Error al actualizar fase 3: {str(e)}")

    def eliminar_fase3(self):
        #Elimina un registro de fase 3 después de validar la contraseña
        try:
            # Solicita y valida la contraseña
            contraseña = self.solicitar_contraseña()
            if contraseña is None:
                return
                
            if not self.validar_contraseña(contraseña):
                QtWidgets.QMessageBox.warning(self, "Error", "Contraseña incorrecta")
                return
            
            # Si la contraseña es válida, procede con eliminar
            selected_row = self.result_table6.currentRow()
            
            if selected_row >= 0:
                # Obtiene el ID de la fase 3 de la primera columna
                id_fase3 = self.result_table6.item(selected_row, 0).text()
                
                reply = QtWidgets.QMessageBox.question(
                    self, 
                    'Confirmar eliminación', 
                    '¿Está seguro de que desea eliminar este registro de fase 3?',
                    QtWidgets.QMessageBox.Yes | QtWidgets.QMessageBox.No, 
                    QtWidgets.QMessageBox.No
                )
                
                if reply == QtWidgets.QMessageBox.Yes:
                    if self.db.eliminar_datos_6(id_fase3):
                        self.result_table6.removeRow(selected_row)
                        QtWidgets.QMessageBox.information(self, "Éxito", "Registro de fase 3 eliminado correctamente.")
                        self.limpiar_fase3()
                    else:
                        QtWidgets.QMessageBox.warning(self, "Error", "No se pudo eliminar el registro de fase 3.")
            else:
                QtWidgets.QMessageBox.warning(self, "Advertencia", "Por favor, seleccione un registro de fase 3 para eliminar.")
                
        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Error", f"Error al eliminar registro de fase 3: {str(e)}")

#####################FASE 4 - Funciones básicas para interactuar con los datos (buscar,agregar,limpiar,eliminar datos)

    def agregar_fase4(self):
        try:
            # Solicita y valida la contraseña
            contraseña = self.solicitar_contraseña()
            if contraseña is None:
                return
                
            if not self.validar_contraseña(contraseña):
                QtWidgets.QMessageBox.warning(self, "Error", "Contraseña incorrecta")
                return
            # Recoge datos de los campos de entrada para la fase 4
            id_fase4 = self.id_fase4.text()
            id_paciente_7 = self.id_paciente_7.text()
            fecha_registro_7 = self.fecha_registro_7.date().toString("yyyy-MM-dd")
            eval_urologia = self.eval_urologia.toPlainText()
            eval_cardiologia = self.eval_cardiologia.toPlainText()
            angiotac_miem_inf = self.angiotac_miem_inf.toPlainText()
            angiotac_ven_art = self.angiotac_ven_art.toPlainText()
            a1 = self.a1.value()
            a2 = self.a2.value()
            b1 = self.b1.value()
            b2 = self.b2.value()
            dr1 = self.dr1.value()
            dr2 = self.dr2.value()
            dq1 = self.dq1.value()
            dq2 = self.dq2.value()

            est_protocolo = self.est_protocolo_map.get(self.est_protocolo.currentText(), None)

            # Valida que todos los campos obligatorios estén llenos
            if not all([id_fase4, id_paciente_7, fecha_registro_7]):
                QMessageBox.warning(self, "Advertencia", "Los campos ID Fase 4, ID Paciente y Fecha de Registro son obligatorios.")
                return

            # Verifica si el id_paciente existe en las tablas receptor o donante
            if not (self.db.existe_paciente_en_receptor_o_donante(id_paciente_7)):
                QMessageBox.warning(self, "Error", f"El ID del paciente {id_paciente_7} no se encuentra en las tablas receptor o donante.")
                return
            
            # Llama al método de la clase de conexión para agregar los datos de la fase 4
            if self.db.agregar_datos_7(
                id_fase4, id_paciente_7, fecha_registro_7, eval_urologia,
                eval_cardiologia, angiotac_miem_inf, angiotac_ven_art,
                a1, a2, b1, b2, dr1, dr2, dq1, dq2,est_protocolo
            ):
                # Limpia la tabla antes de agregar el nuevo registro
                self.result_table7.setRowCount(0)  # Limpia la tabla

                # Actualiza solo las primeras tres columnas en el QTableWidget para la fase 4
                row_position = self.result_table7.rowCount()
                self.result_table7.insertRow(row_position)
                items_to_insert = [id_fase4, id_paciente_7, fecha_registro_7]
                for column_index, item in enumerate(items_to_insert):
                    self.result_table7.setItem(row_position, column_index, QTableWidgetItem(str(item)))

                QMessageBox.information(self, "Éxito", "Información de Fase 4 agregada correctamente.")
                self.limpiar_fase4()  # Limpia el formulario después de agregar
            else:
                QMessageBox.warning(self, "Error", "No se pudo agregar la información de Fase 4.")
        except Exception as e:
                QtWidgets.QMessageBox.critical(self, "Error", f"Error al agregar fase 4: {str(e)}")

#Es un condicional para verificar que RadioButton esta seleccionado y así restringir widgets dependiendo del tipo de paciente (receptor o donante)
    def buscar_fase4_condicional(self):
        if self.id_receptor_f4.isChecked():
            self.mostrar_f4_receptor()
        elif self.id_donante_f4.isChecked():
            self.mostrar_f4_donante()

    def mostrar_f4_receptor(self):
        self.limpiar_fase4()
        results = self.db.buscar_f4_receptor()  # Llama al método del objeto de conexión

        # Limpia la tabla antes de agregar nuevos datos
        self.result_table7.setRowCount(0)

        for row in results:
            row_position = self.result_table7.rowCount()
            self.result_table7.insertRow(row_position)
            
            for column, value in enumerate(row.values()):
                self.result_table7.setItem(row_position, column, QTableWidgetItem(str(value)))

    def mostrar_f4_donante(self):
        if self.id_donante_f4.isChecked():
            try:
                self.limpiar_fase4()
                #Conecta la base datos
                self.db.conecta_base_datos()

                query = """
                SELECT f4.* FROM fase_4 f4
                JOIN donante d ON f4.id_paciente_7 = d.id_donante
                """
                self.db.cursor.execute(query)
                results = self.db.cursor.fetchall()

                self.result_table7.setRowCount(0)

                for row in results:
                    row_position = self.result_table7.rowCount()
                    self.result_table7.insertRow(row_position)
                    
                    for column, value in enumerate(row.values()):
                        self.result_table7.setItem(row_position, column, QTableWidgetItem(str(value)))

            except Exception as e:
                QMessageBox.critical(self, "Error", f"Error al mostrar datos de donantes: {str(e)}")
            finally:
                if self.db.cursor:
                    self.db.cursor.close()
                if self.db.con:
                    self.db.con.close()

    def limpiar_fase4(self):
        # Limpia los campos de texto y restablece valores predeterminados para la fase 4

        # Campos de texto
        self.id_fase4.clear()
        self.id_paciente_7.clear()
        fecha_default = QDate(2000, 1, 1)
        self.fecha_registro_7.setDate(fecha_default)
        self.eval_urologia.clear()
        self.eval_cardiologia.clear()
        self.angiotac_miem_inf.clear()
        self.angiotac_ven_art.clear()
        self.a1.setValue(0)
        self.a2.setValue(0)
        self.b1.setValue(0)
        self.b2.setValue(0)
        self.dr1.setValue(0)
        self.dr2.setValue(0)
        self.dq1.setValue(0)
        self.dq2.setValue(0)
        self.est_protocolo.setCurrentIndex(0)

    def seleccionar_fase4(self):
        selected_row = self.result_table7.currentRow()
        if selected_row >= 0:
            id_fase4 = self.result_table7.item(selected_row, 0).text()
            
            # Obtiene todos los datos de la fase 4 seleccionada de la base de datos
            fase4_data = self.db.obtener_datos_fase4(id_fase4)
            
            if fase4_data:
                # Función auxiliar para obtener valores int o establecerlos en 0
                def get_int_value(value):
                    try:
                        return int(value) if value not in [None, '', 'None'] else 0
                    except ValueError:
                        return 0

                self.limpiar_fase4()

                # Llena los widgets con la información
                self.id_fase4.setText(str(fase4_data['id_fase4']))
                self.id_paciente_7.setText(str(fase4_data['id_paciente_7']))
                self.fecha_registro_7.setDate(QDate.fromString(str(fase4_data['fecha_registro_7']), "yyyy-MM-dd"))
                
                self.eval_urologia.setPlainText(str(fase4_data['eval_urologia']) if fase4_data['eval_urologia'] else "")
                self.eval_cardiologia.setPlainText(str(fase4_data['eval_cardiologia']) if fase4_data['eval_cardiologia'] else "")
                self.angiotac_miem_inf.setPlainText(str(fase4_data['angiotac_miem_inf']) if fase4_data['angiotac_miem_inf'] else "")
                self.angiotac_ven_art.setPlainText(str(fase4_data['angiotac_ven_art']) if fase4_data['angiotac_ven_art'] else "")
                
                self.a1.setValue(get_int_value(fase4_data['a1']))
                self.a2.setValue(get_int_value(fase4_data['a2']))
                self.b1.setValue(get_int_value(fase4_data['b1']))
                self.b2.setValue(get_int_value(fase4_data['b2']))
                self.dr1.setValue(get_int_value(fase4_data['dr1']))
                self.dr2.setValue(get_int_value(fase4_data['dr2']))
                self.dq1.setValue(get_int_value(fase4_data['dq1']))
                self.dq2.setValue(get_int_value(fase4_data['dq2']))
                
                est_protocolo_value = get_int_value(fase4_data['est_protocolo'])
                if est_protocolo_value in self.est_protocolo_map.values():
                    self.est_protocolo.setCurrentText(next(key for key, value in self.est_protocolo_map.items() if value == est_protocolo_value))
                
            else:
                QMessageBox.warning(self, "Error", "No se pudo obtener la información de la Fase 4.")

    def actualizar_fase4(self):
        try:
            # Solicita y valida la contraseña
            contraseña = self.solicitar_contraseña()
            if contraseña is None:
                return
                
            if not self.validar_contraseña(contraseña):
                QtWidgets.QMessageBox.warning(self, "Error", "Contraseña incorrecta")
                return

            # Recoge datos de los campos de entrada para la fase 4
            id_fase4 = self.id_fase4.text()
            id_paciente_7 = self.id_paciente_7.text()
            fecha_registro_7 = self.fecha_registro_7.date().toString("yyyy-MM-dd")
            
            eval_urologia = self.eval_urologia.toPlainText()
            eval_cardiologia = self.eval_cardiologia.toPlainText()
            angiotac_miem_inf = self.angiotac_miem_inf.toPlainText()
            angiotac_ven_art = self.angiotac_ven_art.toPlainText()
            
            a1 = self.a1.value()
            a2 = self.a2.value()
            b1 = self.b1.value()
            b2 = self.b2.value()
            dr1 = self.dr1.value()
            dr2 = self.dr2.value()
            dq1 = self.dq1.value()
            dq2 = self.dq2.value()
            est_protocolo = self.est_protocolo_map.get(self.est_protocolo.currentText(), None)

            # Valida que todos los campos obligatorios estén llenos
            if not all([id_fase4, id_paciente_7, fecha_registro_7, eval_urologia, eval_cardiologia,
                        angiotac_miem_inf, angiotac_ven_art, a1, a2, b1, b2, dr1, dr2, dq1, dq2,est_protocolo]):
                QMessageBox.warning(self, "Advertencia", "Todos los campos son obligatorios.")
                return

            # Llama al método de la clase de conexión para actualizar los datos de la fase 4
            if self.db.actualizar_datos_7(id_fase4, id_paciente_7, fecha_registro_7, eval_urologia,
                                        eval_cardiologia, angiotac_miem_inf, angiotac_ven_art,
                                        a1, a2, b1, b2, dr1, dr2, dq1, dq2,est_protocolo):
                QMessageBox.information(self, "Éxito", "Información actualizada correctamente.")
                # Limpia el formulario después de actualizar
                self.limpiar_fase4()
                # Vuelve a cargar los datos para mostrar todos
                self.buscar_fase4_condicional()
        except Exception as e:
                QtWidgets.QMessageBox.critical(self, "Error", f"Error al agregar fase 4: {str(e)}")
        

    def eliminar_fase4(self):
        #Elimina un registro de fase 4 después de validar la contraseña"""
        try:
            # Solicita y valida la contraseña
            contraseña = self.solicitar_contraseña()
            if contraseña is None:
                return
                
            if not self.validar_contraseña(contraseña):
                QtWidgets.QMessageBox.warning(self, "Error", "Contraseña incorrecta")
                return
            
            # Si la contraseña es válida, procede con eliminar
            selected_row = self.result_table7.currentRow()
            
            if selected_row >= 0:
                # Obtiene el ID de la fase 4 de la primera columna
                id_fase4 = self.result_table7.item(selected_row, 0).text()
                
                reply = QtWidgets.QMessageBox.question(
                    self, 
                    'Confirmar eliminación', 
                    '¿Está seguro de que desea eliminar este registro de fase 4?',
                    QtWidgets.QMessageBox.Yes | QtWidgets.QMessageBox.No, 
                    QtWidgets.QMessageBox.No
                )
                
                if reply == QtWidgets.QMessageBox.Yes:
                    if self.db.eliminar_datos_7(id_fase4):
                        self.result_table7.removeRow(selected_row)
                        QtWidgets.QMessageBox.information(self, "Éxito", "Registro de fase 4 eliminado correctamente.")
                        self.limpiar_fase4()
                    else:
                        QtWidgets.QMessageBox.warning(self, "Error", "No se pudo eliminar el registro de fase 4.")
            else:
                QtWidgets.QMessageBox.warning(self, "Advertencia", "Por favor, seleccione un registro de fase 4 para eliminar.")
                
        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Error", f"Error al eliminar fase 4: {str(e)}")


######################IMAGENES###############################}

    def cargar_imagen(self):
        try:
            # Se abre un diálogo para seleccionar la imagen
            opciones = QFileDialog.Options()
            nombre_archivo, _ = QFileDialog.getOpenFileName(self, "Seleccionar Imagen", "", "Images (*.png *.jpg *.jpeg)", options=opciones)

            if nombre_archivo:
                # Carga la imagen en un QPixmap
                pixmap = QPixmap(nombre_archivo)

                # Muestra el pixmap en el QLabel
                self.ui.labelimagen.setPixmap(pixmap)
                self.ui.labelimagen.setScaledContents(True)  # Ajusta la imagen al tamaño del QLabel si es necesario

                # Guarda temporalmente la ruta de la imagen seleccionada
                self.ruta_imagen = nombre_archivo

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error al seleccionar la imagen: {str(e)}")

    def guardar_imagen(self):
        #Guarda una imagen después de validar la contraseña
        try:
            # Solicita y valida la contraseña
            contraseña = self.solicitar_contraseña()
            if contraseña is None:
                return
                
            if not self.validar_contraseña(contraseña):
                QtWidgets.QMessageBox.warning(self, "Error", "Contraseña incorrecta")
                return
            
            # Si la contraseña es válida, se procede con guardar
            # Obtiene valores de los widgets, asegurando que sean numéricos
            id_imagen_text = self.id_imagen.text().strip()
            id_paciente_8_text = self.id_paciente_8.text().strip()
            fecha_registro_8 = self.fecha_registro_8.date().toString("yyyy-MM-dd")
            
            # Obtiene el texto seleccionado del ComboBox y se mapea a su valor numérico
            tipo_examen_text = self.tipo_examen.currentText().strip()
            tipo_examen = self.tipo_imagen.get(tipo_examen_text)

            # Valida que los campos no estén vacíos y contengan valores numéricos donde sea necesario
            if not id_imagen_text.isdigit():
                QtWidgets.QMessageBox.warning(self, "Advertencia", "ID de imagen debe ser un número.")
                return

            if not id_paciente_8_text.isdigit():
                QtWidgets.QMessageBox.warning(self, "Advertencia", "ID del paciente debe ser un número.")
                return

            if tipo_examen is None:
                QtWidgets.QMessageBox.warning(self, "Advertencia", "Seleccione un tipo de examen válido.")
                return

            # Verifica si el id_paciente_3 existe en las tablas receptor o donante
            if not (self.db.existe_paciente_en_receptor_o_donante(id_paciente_8_text)):
                QtWidgets.QMessageBox.warning(self, "Error", f"El ID del paciente {id_paciente_8_text} no se encuentra en las tablas receptor o donante.")
                return
            
            # Convierte a enteros después de la validación
            id_imagen = int(id_imagen_text)
            id_paciente_8 = int(id_paciente_8_text)

            if hasattr(self, 'ruta_imagen'):
                with open(self.ruta_imagen, 'rb') as file:
                    binary_data = file.read()

                # Conecta a la base de datos y ejecuta la consulta
                self.db.conecta_base_datos()
                query = "INSERT INTO imagenes (id_imagen, id_paciente_8, fecha_registro_8, tipo_imagen, data_imagen) VALUES (%s, %s, %s, %s, %s)"
                self.db.cursor.execute(query, (id_imagen, id_paciente_8, fecha_registro_8, tipo_examen, binary_data))
                self.db.con.commit()

                QtWidgets.QMessageBox.information(self, "Éxito", "Imagen guardada correctamente.")
                self.limpiar_imagen()  # Limpia el formulario después de actualizar
                self.buscar_imagen_condicional()  # Vuelve a cargar los datos para mostrar todos
            else:
                QtWidgets.QMessageBox.warning(self, "Advertencia", "No se ha seleccionado ninguna imagen para guardar.")
                
        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Error", f"Error al guardar la imagen: {str(e)}")
            
        finally:
            if hasattr(self.db, 'cursor') and self.db.cursor:
                self.db.cursor.close()
            if hasattr(self.db, 'con') and self.db.con:
                self.db.con.close()

#Es un condicional para verificar que RadioButton esta seleccionado y así restringir widgets dependiendo del tipo de paciente (receptor o donante)
    def buscar_imagen_condicional(self):
        if self.id_receptor_imagen.isChecked():
            self.mostrar_imagen_receptor()
        elif self.id_donante_imagen.isChecked():
            self.mostrar_imagen_donante()

    def mostrar_imagen_receptor(self):
        try:
            #Conecta a la base de datos
            self.db.conecta_base_datos()

            query = """
            SELECT img.id_imagen, img.id_paciente_8, img.fecha_registro_8, img.tipo_imagen, img.data_imagen
            FROM imagenes img
            JOIN receptor r ON img.id_paciente_8 = r.id_receptor
            """
            self.db.cursor.execute(query)
            results = self.db.cursor.fetchall()

            self.ui.tableWidget_34.setRowCount(0)

            for row in results:
                row_position = self.ui.tableWidget_34.rowCount()
                self.ui.tableWidget_34.insertRow(row_position)

                # Llena la data de imagen
                self.ui.tableWidget_34.setItem(row_position, 0, QTableWidgetItem(str(row['id_imagen'])))
                self.ui.tableWidget_34.setItem(row_position, 1, QTableWidgetItem(str(row['id_paciente_8'])))
                fecha_str = row['fecha_registro_8'].strftime("%Y-%m-%d")
                self.ui.tableWidget_34.setItem(row_position, 2, QTableWidgetItem(fecha_str))
                self.ui.tableWidget_34.setItem(row_position, 3, QTableWidgetItem(str(row['tipo_imagen'])))

                # Indica la aviabilidad de la imagen
                if row['data_imagen']:
                    self.ui.tableWidget_34.setItem(row_position, 4, QTableWidgetItem("Imagen disponible"))
                else:
                    self.ui.tableWidget_34.setItem(row_position, 4, QTableWidgetItem("Sin imagen"))

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error al mostrar datos de receptores: {str(e)}")
        finally:
            if hasattr(self.db, 'cursor') and self.db.cursor:
                self.db.cursor.close()
            if hasattr(self.db, 'con') and self.db.con:
                self.db.con.close()

    def mostrar_imagen_donante(self):
        try:
            # Conecta a la base de datos
            self.db.conecta_base_datos()

            query = """
            SELECT img.id_imagen, img.id_paciente_8, img.fecha_registro_8, img.tipo_imagen, img.data_imagen
            FROM imagenes img
            JOIN donante d ON img.id_paciente_8 = d.id_donante
            """
            self.db.cursor.execute(query)
            results = self.db.cursor.fetchall()

            self.ui.tableWidget_34.setRowCount(0)

            # Llena la tabla con la data filtrada incluida la aviabilidad de la imagen
            for row in results:
                row_position = self.ui.tableWidget_34.rowCount()
                self.ui.tableWidget_34.insertRow(row_position)

                self.ui.tableWidget_34.setItem(row_position, 0, QTableWidgetItem(str(row['id_imagen'])))
                self.ui.tableWidget_34.setItem(row_position, 1, QTableWidgetItem(str(row['id_paciente_8'])))
                fecha_str = row['fecha_registro_8'].strftime("%Y-%m-%d")
                self.ui.tableWidget_34.setItem(row_position, 2, QTableWidgetItem(fecha_str))
                self.ui.tableWidget_34.setItem(row_position, 3, QTableWidgetItem(str(row['tipo_imagen'])))

                if row['data_imagen']:
                    self.ui.tableWidget_34.setItem(row_position, 4, QTableWidgetItem("Imagen disponible"))
                else:
                    self.ui.tableWidget_34.setItem(row_position, 4, QTableWidgetItem("Sin imagen"))

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error al mostrar datos de donantes: {str(e)}")
        finally:
            if hasattr(self.db, 'cursor') and self.db.cursor:
                self.db.cursor.close()
            if hasattr(self.db, 'con') and self.db.con:
                self.db.con.close()
        

    def seleccionar_imagen(self):
        selected_row = self.ui.tableWidget_34.currentRow()

        if selected_row >= 0:
            # Obtiene datos de la fila seleccionada
            id_imagen = self.ui.tableWidget_34.item(selected_row, 0).text()
            id_paciente_8 = self.ui.tableWidget_34.item(selected_row, 1).text()
            fecha_registro_8 = self.ui.tableWidget_34.item(selected_row, 2).text()
            tipo_imagen_num = self.ui.tableWidget_34.item(selected_row, 3).text()

            # Convierte el número de tipo de examen a texto usando el mapeo
            tipo_examen_text = next((key for key, value in self.tipo_imagen.items() if str(value) == tipo_imagen_num), "")

            # Establece datos en los widgets
            self.id_imagen.setText(id_imagen if id_imagen else "")
            self.id_paciente_8.setText(id_paciente_8 if id_paciente_8 else "")
            
            if fecha_registro_8:
                self.fecha_registro_8.setDate(QDate.fromString(fecha_registro_8, "yyyy-MM-dd"))
            else:
                self.fecha_registro_8.setDate(QDate.currentDate())
            
            if tipo_examen_text:
                self.tipo_examen.setCurrentText(tipo_examen_text)
            else:
                self.tipo_examen.setCurrentIndex(0)

            # Obtiene la imagen de la base de datos
            try:
                self.db.conecta_base_datos()
                query = "SELECT data_imagen FROM imagenes WHERE id_imagen = %s"
                self.db.cursor.execute(query, (id_imagen,))
                result = self.db.cursor.fetchone()

                if result and result['data_imagen']:
                    data_imagen = result['data_imagen']

                    # Convierte BLOB a QPixmap y muestra en QLabel
                    pixmap = QPixmap()
                    pixmap.loadFromData(data_imagen)
                    self.ui.labelimagen.setPixmap(pixmap)
                    self.ui.labelimagen.setScaledContents(True)
                else:
                    self.ui.labelimagen.clear()

            except Exception as e:
                QMessageBox.critical(self, "Error", f"Error al cargar la imagen: {str(e)}")
            finally:
                if hasattr(self.db, 'cursor') and self.db.cursor:
                    self.db.cursor.close()
                if hasattr(self.db, 'con') and self.db.con:
                    self.db.con.close()

    def limpiar_imagen(self): #Quita la imagen del widget
        self.id_imagen.clear()
        self.id_paciente_8.clear()
        self.fecha_registro_8.setDate(QDate.currentDate())
        self.tipo_examen.setCurrentIndex(0)
        self.ui.labelimagen.clear()

    def actualizar_imagen(self):
        try:
            # Solicita y valida la contraseña
            contraseña = self.solicitar_contraseña()
            if contraseña is None:
                return
                
            if not self.validar_contraseña(contraseña):
                QtWidgets.QMessageBox.warning(self, "Error", "Contraseña incorrecta")
                return
            id_imagen = self.id_imagen.text().strip()
            id_paciente_8 = self.id_paciente_8.text().strip()
            tipo_examen_text = self.tipo_examen.currentText().strip()
            tipo_examen = self.tipo_imagen.get(tipo_examen_text)
            fecha_registro_8 = self.fecha_registro_8.date().toString("yyyy-MM-dd")

            if not id_imagen.isdigit() or not id_paciente_8.isdigit() or tipo_examen is None:
                QMessageBox.warning(self, "Advertencia", "Datos inválidos o incompletos.")
                return

            id_imagen = int(id_imagen)
            id_paciente_8 = int(id_paciente_8)

            # Conecta a la base de datos
            self.db.conecta_base_datos()

            if hasattr(self, 'ruta_imagen'):
                with open(self.ruta_imagen, 'rb') as file:
                    data_binaria = file.read()

                # Actualiza el registro con nueva imagen
                query = """
                    UPDATE imagenes 
                    SET id_paciente_8 = %s, fecha_registro_8 = %s, tipo_imagen = %s, data_imagen = %s 
                    WHERE id_imagen = %s
                """
                self.db.cursor.execute(query, (id_paciente_8, fecha_registro_8, tipo_examen, data_binaria, id_imagen))
            else:
                # Actualiza el registro sin cambiar la imagen
                query = """
                    UPDATE imagenes 
                    SET id_paciente_8 = %s, fecha_registro_8 = %s, tipo_imagen = %s 
                    WHERE id_imagen = %s
                """
                self.db.cursor.execute(query, (id_paciente_8, fecha_registro_8, tipo_examen, id_imagen))

            self.db.con.commit()
            QMessageBox.information(self, "Éxito", "Información actualizada correctamente.")
            self.limpiar_imagen()  # Limpia el formulario después de actualizar
            self.buscar_imagen_condicional()  # Vuelve a cargar los datos para mostrar todos

        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Error", f"Error al eliminar fase 4: {str(e)}")

        finally:
            if hasattr(self.db, 'cursor') and self.db.cursor:
                self.db.cursor.close()
            if hasattr(self.db, 'con') and self.db.con:
                self.db.con.close()

    def eliminar_imagen(self):
        #Elimina una imagen después de validar la contraseña
        try:
            # Solicita y valida la contraseña
            contraseña = self.solicitar_contraseña()
            if contraseña is None:
                return
                
            if not self.validar_contraseña(contraseña):
                QtWidgets.QMessageBox.warning(self, "Error", "Contraseña incorrecta")
                return
            
            # Si la contraseña es válida, procede con eliminar
            id_imagen = self.id_imagen.text().strip()

            if not id_imagen.isdigit():
                QtWidgets.QMessageBox.warning(self, "Advertencia", "ID de imagen inválido.")
                return

            id_imagen = int(id_imagen)

            reply = QtWidgets.QMessageBox.question(
                self,
                'Confirmar eliminación',
                '¿Está seguro de que desea eliminar este registro?',
                QtWidgets.QMessageBox.Yes | QtWidgets.QMessageBox.No,
                QtWidgets.QMessageBox.No
            )

            if reply == QtWidgets.QMessageBox.Yes:
                # Elimina el registro de la base de datos
                self.db.conecta_base_datos()
                query = "DELETE FROM imagenes WHERE id_imagen = %s"
                self.db.cursor.execute(query, (id_imagen,))
                self.db.con.commit()

                QtWidgets.QMessageBox.information(self, "Éxito", "Registro eliminado correctamente.")
                self.limpiar_imagen() 
                self.buscar_imagen_condicional() 
                
        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Error", f"Error al eliminar la imagen: {str(e)}")
            
        finally:
            if hasattr(self.db, 'cursor') and self.db.cursor:
                self.db.cursor.close()
            if hasattr(self.db, 'con') and self.db.con:
                self.db.con.close()

#Gráficos - Sección para evaluación periódica

    def habilitar_graficos(self):
        id_paciente = self.id_paciente_grafico.text().strip()
        
        if id_paciente:
            # Habilita los radio buttons
            self.habilitar_radio_buttons(True)
            self.id_paciente_actual = id_paciente  # Guarda el ID del paciente actual
            QMessageBox.information(self, "Éxito", f"ID de paciente {id_paciente} válido. Puede seleccionar los gráficos.")
        else:
            # Deshabilita los radio buttons
            self.habilitar_radio_buttons(False)

            # Limpia el ID del paciente actual
            self.id_paciente_actual = None
            
            # Limpia los gráficos
            self.limpiar_graficos()
            
            QMessageBox.warning(self, "Error", "Por favor, ingrese un ID de paciente antes de graficar.")

    def habilitar_radio_buttons(self, habilitar):
        radio_buttons = [
            self.ui.radioButton_29, self.ui.radioButton_30, self.ui.radioButton,
            self.ui.radioButton_14, self.ui.radioButton_17, self.ui.radioButton_18,
            self.ui.radioButton_21, self.ui.radioButton_2, self.ui.radioButton_9,
            self.ui.radioButton_27, self.ui.radioButton_22,
            self.ui.radioButton_44, self.ui.radioButton_43, self.ui.radioButton_23,
            self.ui.radioButton_24, self.ui.radioButton_25, self.ui.radioButton_26,
            self.ui.radioButton_31, self.ui.radioButton_32, self.ui.radioButton_33,
            self.ui.radioButton_34, self.ui.radioButton_35, self.ui.radioButton_36,
            self.ui.radioButton_37, self.ui.radioButton_38, self.ui.radioButton_10,
            self.ui.radioButton_11, self.ui.radioButton_12, self.ui.radioButton_28,
            self.ui.radioButton_39, self.ui.radioButton_40, self.ui.radioButton_45,
            self.ui.radioButton_46, self.ui.radioButton_53, self.ui.radioButton_47,
            self.ui.radioButton_48, self.ui.radioButton_51, self.ui.radioButton_52
        ]
        
        for radio_button in radio_buttons:
            radio_button.setEnabled(habilitar)
            if not habilitar:
                radio_button.setChecked(False)


    def limpiar_graficos(self):
        # Limpia el gráfico de la Fase 1
        self.figure.clear()
        self.canvas.draw()
        
        # Limpia el gráfico de la Fase 2b
        self.figure2.clear()
        self.canvas2.draw()

    #Funciones para obtener los indicadores bioquímicos de la fase 1
    def obtener_datos(self, variable, id_paciente):
        try:
            self.db.conecta_base_datos()
            query = f"""
            SELECT secuencial.fecha_registro, fase_1.{variable}
            FROM fase_1
            JOIN secuencial ON fase_1.id_fase1 = secuencial.id_fase
            WHERE fase_1.id_paciente_3 = %s
            """
            self.db.cursor.execute(query, (id_paciente,))
            results = self.db.cursor.fetchall()
            
            # Convierte resultados a DataFrame
            df = pd.DataFrame(results, columns=['fecha_registro', variable])
            df['fecha_registro'] = pd.to_datetime(df['fecha_registro'])
            
            return df
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error al obtener datos: {str(e)}")
            return pd.DataFrame()
        finally:
            if self.db.cursor is not None:
                self.db.cursor.close()
            if self.db.con is not None:
                self.db.con.close()

    #Se procede a graficar algun indicador seleccionado en fase 1
    def graficar_datos(self, variable):
        if not self.id_paciente_actual:
            QMessageBox.warning(self, "Error", "Por favor, ingrese un ID de paciente antes de graficar.")
            return
        
        df = self.obtener_datos(variable, self.id_paciente_actual)
        
        if df.empty:
            QMessageBox.warning(self, "Advertencia", "No hay datos disponibles para graficar.")
            return

        self.figure.clear()
        ax = self.figure.add_subplot(111)
        
        ax.set_facecolor((255/255, 255/255, 255/255))  # Establecer color de fondo

        # Se grafica la línea
        line, = ax.plot(df['fecha_registro'], df[variable], label=variable)

        # Añade marcadores para todos los puntos
        ax.scatter(df['fecha_registro'], 
                    df[variable], 
                    color='red', s=50, label='Valores Registrados')

        # Configura el formato de las fechas
        ax.xaxis.set_major_formatter(mdates.DateFormatter('%Y-%m-%d'))
        ax.xaxis.set_major_locator(mdates.AutoDateLocator())

        ax.set_title(f'Evolución de {variable} en el tiempo')
        ax.set_xlabel('Fecha')
        ax.set_ylabel('Valor')
        ax.legend()
        
        plt.setp(ax.xaxis.get_majorticklabels(), rotation=15)
    
        cursor = mplcursors.cursor(line, hover=True)
        
        cursor.connect("add", lambda sel: sel.annotation.set_text(
            f'{df["fecha_registro"].iloc[int(sel.index)]}: {sel.target[1]:.2f}'))

        self.canvas.draw()
        
#Funciones para obtener los indicadores bioquímicos de la fase 2b
    def obtener_datos2(self, variable, id_paciente):
        try:
            self.db.conecta_base_datos()
            query = f"""
            SELECT secuencial.fecha_registro, fase_2b.{variable}
            FROM fase_2b
            JOIN secuencial ON fase_2b.id_fase2b = secuencial.id_fase
            WHERE fase_2b.id_paciente_5 = %s
            """
            self.db.cursor.execute(query, (id_paciente,))
            results = self.db.cursor.fetchall()
            
            # Convierte resultados a DataFrame
            df = pd.DataFrame(results, columns=['fecha_registro', variable])
            df['fecha_registro'] = pd.to_datetime(df['fecha_registro'])
            
            return df
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error al obtener datos: {str(e)}")
            return pd.DataFrame()
        finally:
            if self.db.cursor is not None:
                self.db.cursor.close()
            if self.db.con is not None:
                self.db.con.close()

#Se procede a graficar algun indicador seleccionado antes para fase 2b
    def graficar_datos2(self, variable):
        if not self.id_paciente_actual:
            QMessageBox.warning(self, "Error", "Por favor, ingrese un ID de paciente antes de graficar.")
            return
        
        df = self.obtener_datos2(variable, self.id_paciente_actual)
        
        if df.empty:
            QMessageBox.warning(self, "Advertencia", "No hay datos disponibles para graficar.")
            return

        self.figure2.clear()
        ax = self.figure2.add_subplot(111)
        
        ax.set_facecolor((255/255, 255/255, 255/255))  # Establecer color de fondo

        # Se grafica la línea
        line, = ax.plot(df['fecha_registro'], df[variable], label=variable)

        # Añade marcadores para todos los puntos
        ax.scatter(df['fecha_registro'], 
                    df[variable], 
                    color='red', s=50, label='Valores Registrados')

        # Configura el formato de las fechas
        ax.xaxis.set_major_formatter(mdates.DateFormatter('%Y-%m-%d'))
        ax.xaxis.set_major_locator(mdates.AutoDateLocator())

        ax.set_title(f'Evolución de {variable} en el tiempo')
        ax.set_xlabel('Fecha')
        ax.set_ylabel('Valor')
        ax.legend()
        
        plt.setp(ax.xaxis.get_majorticklabels(), rotation=15)
        
        cursor = mplcursors.cursor(line, hover=True)
        
        cursor.connect("add", lambda sel: sel.annotation.set_text(
            f'{df["fecha_registro"].iloc[int(sel.index)]}: {sel.target[1]:.2f}'))

        self.canvas2.draw()


############################FASE DE COMPATIBILIDAD -Selección de donantes potenciales

    def buscar_compatibilidad(self):
        # Recoge valores del formulario
        id_receptor_2 = self.id_receptor_2.text()
        rango_edad = self.rango_edad.currentText()

        # Obtiene los datos del receptor desde la base de datos
        receptor_data = self.db.buscar_datos_0()
        grupo_sanguineo_receptor_num = None
        for receptor in receptor_data:
            if receptor['id_receptor'] == int(id_receptor_2):
                grupo_sanguineo_receptor_num = receptor['grupo_sanguineo']
                break

        if grupo_sanguineo_receptor_num is None:
            QMessageBox.critical(self, "Error", "Receptor no encontrado.")
            return

        # Determina grupos sanguíneos compatibles basados en números
        compatible_grupos_num = self.obtener_grupos_sanguineos_compatibles(grupo_sanguineo_receptor_num)

        # Define los límites de edad basados en el rango seleccionado
        if rango_edad == "Adulto joven (20 - 39) años":
            edad_min, edad_max = 20, 39
        elif rango_edad == "Adulto medio (40-59) años":
            edad_min, edad_max = 40, 59
        elif rango_edad == "Adulto mayor (60+) años":
            edad_min, edad_max = 60, 120

        try:
            # Conecta a la base de datos
            self.db.conecta_base_datos()

            # Construiye la consulta SQL para obtener datos del donante y fase 4
            placeholders = ','.join(['%s'] * len(compatible_grupos_num))
            query = f"""
            SELECT d.id_donante, d.edad, d.grupo_sanguineo, f.est_protocolo, f.a1, f.a2, f.b1, f.b2, f.dr1, f.dr2, f.dq1, f.dq2
            FROM donante d
            JOIN fase_4 f ON d.id_donante = f.id_paciente_7
            WHERE d.edad BETWEEN %s AND %s AND d.grupo_sanguineo IN ({placeholders})
            """
            
            # Ejecuta la consulta con los argumentos correctos
            self.db.cursor.execute(query, (edad_min, edad_max, *compatible_grupos_num))
            
            # Obtiene resultados
            results = self.db.cursor.fetchall()

            # Limpia la tabla antes de mostrar nuevos resultados
            self.result_table8.setRowCount(0)

            # Inserta resultados en la tabla
            for row in results:
                row_position = self.result_table8.rowCount()
                self.result_table8.insertRow(row_position)
                self.result_table8.setItem(row_position, 0, QTableWidgetItem(str(row['id_donante'])))
                self.result_table8.setItem(row_position, 1, QTableWidgetItem(str(row['edad'])))
                
                # Convierte los combobox de numérico a texto usando el mapeo
                grupo_sanguineo_texto = next(key for key, value in self.grupo_sanguineo_map.items() if value == row['grupo_sanguineo'])
                self.result_table8.setItem(row_position, 2, QTableWidgetItem(grupo_sanguineo_texto))
                est_protocolo_texto = next(key for key, value in self.est_protocolo_map.items() if value == row['est_protocolo'])
                self.result_table8.setItem(row_position, 3, QTableWidgetItem(est_protocolo_texto))
                self.result_table8.setItem(row_position, 4, QTableWidgetItem(str(row['a1'])))
                self.result_table8.setItem(row_position, 5, QTableWidgetItem(str(row['a2'])))
                self.result_table8.setItem(row_position, 6, QTableWidgetItem(str(row['b1'])))
                self.result_table8.setItem(row_position, 7, QTableWidgetItem(str(row['b2'])))
                self.result_table8.setItem(row_position, 8, QTableWidgetItem(str(row['dr1'])))
                self.result_table8.setItem(row_position, 9, QTableWidgetItem(str(row['dr2'])))
                self.result_table8.setItem(row_position, 10,QTableWidgetItem(str(row['dq1'])))
                self.result_table8.setItem(row_position, 11,QTableWidgetItem(str(row['dq2'])))

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error al buscar compatibilidad: {str(e)}")
        finally:
            if self.db.cursor:
                self.db.cursor.close()
            if self.db.con:
                self.db.con.close()

    def obtener_grupos_sanguineos_compatibles(self, grupo_num):
        # Define compatibilidad de grupos sanguíneos usando números
        compatibilidad_numerica = {
            1: [1, 5, 3, 7],   # A+
            2: [2, 6, 3, 7],   # B+
            3: [3, 7],         # O+
            4: [4, 8, 1, 5, 2, 6, 3, 7],   # AB+
            5: [5, 7],         # A-
            6: [6, 7],         # B-
            7: [7],            # O-
            8: [8, 5, 6, 7]    # AB-
        }
        return compatibilidad_numerica.get(grupo_num, [])

    def mostrar_datos_receptor(self):
        # Recoge el ID del receptor ingresado
        id_receptor_2 = self.id_receptor_2.text()

        try:
            # Conecta a la base de datos
            self.db.conecta_base_datos()
            query_receptor = "SELECT grupo_sanguineo, est_acc_vasc FROM receptor WHERE id_receptor = %s"
            self.db.cursor.execute(query_receptor, (id_receptor_2,))
            receptor_data = self.db.cursor.fetchone()

            if not receptor_data:
                QMessageBox.critical(self, "Error", "Receptor no encontrado en la tabla receptor.")
                return

            grupo_sanguineo_num = receptor_data['grupo_sanguineo']
            est_acc_vascul_num = receptor_data['est_acc_vasc']

            # Convierte valores numéricos a texto usando mapeos
            grupo_sanguineo_texto = next(key for key, value in self.grupo_sanguineo_map.items() if value == grupo_sanguineo_num)
            est_acc_vascul_texto = next(key for key, value in self.est_acc_vascul_map.items() if value == est_acc_vascul_num)

            # Verifica si el id_paciente está en fase_4 para obtener est_protocolo
            query_fase4 = "SELECT est_protocolo FROM fase_4 WHERE id_paciente_7 = %s"
            self.db.cursor.execute(query_fase4, (id_receptor_2,))
            fase4_data = self.db.cursor.fetchone()

            if not fase4_data:
                QMessageBox.critical(self, "Error", "Datos de fase 4 no encontrados para este receptor.")
                return

            # Obtiene y convierte el estado del protocolo
            est_protocolo_num = fase4_data['est_protocolo']
            est_protocolo_texto = next(key for key, value in self.est_protocolo_map.items() if value == est_protocolo_num)

            # Muestra los datos en los widgets de solo vista
            self.grupo_sanguineo_r.setText(grupo_sanguineo_texto)
            self.est_acc_vascul.setText(est_acc_vascul_texto)
            self.estado_protocolo.setText(est_protocolo_texto)

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error al mostrar datos del receptor: {str(e)}")
        finally:
            if self.db.cursor:
                self.db.cursor.close()
            if self.db.con:
                self.db.con.close()

    def seleccionar_compatibilidad(self):
        selected_row = self.result_table8.currentRow()
        if selected_row < 0:
            QMessageBox.warning(self, "Advertencia", "Por favor, seleccione un registro para calcular compatibilidad.")
            return

        try:
            # Obtiene el ID del donante seleccionado
            id_donante = self.result_table8.item(selected_row, 0).text()
            
            # Obtiene el ID del receptor ingresado
            id_receptor_2 = self.id_receptor_2.text()

            # Muestra los IDs en los widgets correspondientes
            self.id_receptor_3.setText(id_receptor_2)
            self.id_donante_2.setText(id_donante)

            self.db.conecta_base_datos()

            query_donante = """
            SELECT a1, a2, b1, b2, dr1, dr2, dq1, dq2 
            FROM fase_4 
            WHERE id_paciente_7 = %s
            """
            self.db.cursor.execute(query_donante, (id_donante,))
            donante_data = self.db.cursor.fetchone()

            if not donante_data:
                QMessageBox.critical(self, "Error", "Datos de fase 4 no encontrados para este donante.")
                return
            
            #Datos del receptor
            query_receptor = """
            SELECT a1, a2, b1, b2, dr1, dr2, dq1, dq2 
            FROM fase_4 
            WHERE id_paciente_7 = %s
            """
            self.db.cursor.execute(query_receptor, (id_receptor_2,))
            receptor_data = self.db.cursor.fetchone()

            if not receptor_data:
                QMessageBox.critical(self, "Error", "Datos de fase 4 no encontrados para este receptor.")
                return

            # Calcula HLA tipo 1 (a1, a2, b1, b2)
            hla_tipo_1_matches = sum(
                1 for key in ['a1', 'a2', 'b1', 'b2']
                if donante_data[key] == receptor_data[key]
            )
            
            hla_tipo_1_result = f"{hla_tipo_1_matches}/4"
            
            # Calcula HLA tipo 2 (dr1, dr2, dq1, dq2)
            hla_tipo_2_matches = sum(
                1 for key in ['dr1', 'dr2', 'dq1', 'dq2']
                if donante_data[key] == receptor_data[key]
            )
            
            hla_tipo_2_result = f"{hla_tipo_2_matches}/4"

            # Muestra resultados en los widgets correspondientes
            self.hla_1.setText(hla_tipo_1_result)
            self.hla_2.setText(hla_tipo_2_result)

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error al seleccionar compatibilidad: {str(e)}")
        finally:
            if self.db.cursor:
                self.db.cursor.close()
            if self.db.con:
                self.db.con.close()

    def agregar_compatibilidad(self):
        #Toma información de los widgets
        id_fase_final = self.id_fase_final.text()
        id_receptor_3 = self.id_receptor_3.text()
        id_donante_2 = self.id_donante_2.text()
        fecha_registro_9 = self.fecha_registro_9.date().toString('yyyy-MM-dd')
        parentesco_value = self.parentesco.currentText() 
        tipo_donante_value = self.tipo_donante.currentText()
        hla_1_value = self.hla_1.text()
        hla_2_value = self.hla_2.text()
        pra_hla1_2_value = self.pra_hla1_2.text()
        prueba_cruzada_value = self.prueba_cruzada.currentText()

        if not all([id_fase_final, id_receptor_3, id_donante_2, fecha_registro_9,
                    hla_1_value, hla_2_value, pra_hla1_2_value]):
            QMessageBox.warning(self, "Advertencia", "Todos los campos son obligatorios.")
            return

        # Limpia la tabla antes de insertar datos
        self.result_table_9.setRowCount(0)

        row_position = self.result_table_9.rowCount()
        items_to_insert = [
            id_fase_final,
            id_receptor_3,
            id_donante_2,
            fecha_registro_9,
            parentesco_value,
            tipo_donante_value,
            hla_1_value,
            hla_2_value,
            pra_hla1_2_value,
            prueba_cruzada_value
        ]
        
        self.result_table_9.insertRow(row_position)
        for column_index, item in enumerate(items_to_insert):
            self.result_table_9.setItem(row_position, column_index, QTableWidgetItem(item))

        # Llama al método para guardar datos en la base de datos
        if self.db.agregar_datos_9(
            int(id_fase_final), id_receptor_3, id_donante_2, fecha_registro_9,
            self.parentesco.currentIndex(), self.tipo_donante.currentIndex(),
            hla_1_value, hla_2_value, pra_hla1_2_value,
            self.prueba_cruzada.currentIndex()
        ):
            QMessageBox.information(self, "Éxito", "Información agregada correctamente.")
        self.buscar_registros()
        self.limpiar_compatibilidad()  # Limpia el formulario después de agregar

    def limpiar_compatibilidad(self):
        # Limpia los campos de texto
        self.id_fase_final.clear()
        self.id_receptor_3.clear()
        self.id_donante_2.clear()
        fecha_default = QDate(2000, 1, 1)
        self.fecha_registro_9.setDate(fecha_default)
        self.parentesco.setCurrentIndex(0)
        self.tipo_donante.setCurrentIndex(0)
        self.prueba_cruzada.setCurrentIndex(0)
        self.hla_1.clear()
        self.hla_2.clear()
        self.pra_hla1_2.clear()

    def buscar_registros(self):
        # Llama al método de la clase de conexión para buscar datos de la fase 9
        results = self.db.buscar_datos_9()

        # Limpia la tabla antes de agregar nuevos datos
        self.result_table_9.setRowCount(0)

        # Invierte los mapeos para transformar índices numéricos a texto
        parentesco_map_inv = {v: k for k, v in self.parentesco_map.items()}
        tipo_donante_map_inv = {v: k for k, v in self.tipo_donante_map.items()}
        prueba_cruzada_map_inv = {v: k for k, v in self.prueba_cruzada_map.items()}

        for row in results:
            row_position = self.result_table_9.rowCount()
            self.result_table_9.insertRow(row_position)

            # Transforma los valores numéricos a texto
            row['parentesco'] = parentesco_map_inv.get(row['parentesco'], 'Desconocido')
            row['tipo_donante'] = tipo_donante_map_inv.get(row['tipo_donante'], 'Desconocido')
            row['prueba_cruzada'] = prueba_cruzada_map_inv.get(row['prueba_cruzada'], 'Desconocido')

            # Inserta los valores en la tabla
            for column, value in enumerate(row.values()):
                self.result_table_9.setItem(row_position, column, QTableWidgetItem(str(value)))

    def eliminar_registro(self):
                selected_row = self.result_table_9.currentRow()
                if selected_row >= 0:
                    id_fase_final = self.result_table_9.item(selected_row, 0).text()
                    
                    # Se confirma la eliminación
                    reply = QMessageBox.question(self, 'Confirmar eliminación',
                                                '¿Está seguro de que desea eliminar este registro de fase final?',
                                                QMessageBox.Yes | QMessageBox.No, QMessageBox.No)
                    if reply == QMessageBox.Yes:
                        # Llama al método de la clase de conexión para eliminar el registro
                        if self.db.eliminar_datos_9(id_fase_final):
                            self.result_table_9.removeRow(selected_row)
                            QMessageBox.information(self, "Éxito", "Registro de fase 3 eliminado correctamente.")
                            self.limpiar_compatibilidad()
                        else:
                            QMessageBox.warning(self, "Error", "No se pudo eliminar el registro de fase final.")
                    else:
                        QMessageBox.warning(self, "Advertencia", "Por favor, seleccione un registro de fase final para eliminar.")

    def seleccionar_registro(self):
        selected_row = self.result_table_9.currentRow()
        
        if selected_row >= 0:
            # Obtiene y establece los datos desde la tabla para la fase 4
            self.id_fase_final.setText(self.result_table_9.item(selected_row, 0).text())
            self.id_receptor_3.setText(self.result_table_9.item(selected_row, 1).text())
            self.id_donante_2.setText(self.result_table_9.item(selected_row, 2).text())
            self.fecha_registro_9.setDate(QDate.fromString(self.result_table_9.item(selected_row, 3).text(), "yyyy-MM-dd"))

            parentesco_text = self.result_table_9.item(selected_row, 4).text()
            if parentesco_text:
                self.parentesco.setCurrentText(parentesco_text)
            
            tipo_donante_text = self.result_table_9.item(selected_row, 5).text()
            if tipo_donante_text:
                self.tipo_donante.setCurrentText(tipo_donante_text)

            self.hla_1.setText(self.result_table_9.item(selected_row, 6).text())
            self.hla_2.setText(self.result_table_9.item(selected_row, 7).text())
            self.pra_hla1_2.setText(self.result_table_9.item(selected_row, 8).text())
            
            prueba_cruzada_text = self.result_table_9.item(selected_row, 9).text()
            if prueba_cruzada_text:
                self.prueba_cruzada.setCurrentText(prueba_cruzada_text)

    def actualizar_registro(self):
        try:
            # Recoge datos de los campos de entrada que son editables
            id_fase_final = self.id_fase_final.text()
            fecha_registro_9 = self.fecha_registro_9.date().toString("yyyy-MM-dd")
            
            # Obtiene valores de combo boxes
            parentesco_text = self.parentesco.currentText()
            tipo_donante_text = self.tipo_donante.currentText()
            prueba_cruzada_text = self.prueba_cruzada.currentText()

            # Valida que todos los campos obligatorios estén llenos
            if not all([id_fase_final, fecha_registro_9, parentesco_text, tipo_donante_text, prueba_cruzada_text]):
                QMessageBox.warning(self, "Advertencia", "Todos los campos editables son obligatorios.")
                return

            # Obtiene el valor entero correspondiente a los textos seleccionados en los combo boxes
            parentesco = self.parentesco_map.get(parentesco_text)
            tipo_donante = self.tipo_donante_map.get(tipo_donante_text)
            prueba_cruzada = self.prueba_cruzada_map.get(prueba_cruzada_text)

            # Llama al método de la clase de conexión para actualizar el registro
            if self.db.actualizar_datos_9(id_fase_final, fecha_registro_9, parentesco, tipo_donante, prueba_cruzada):
                QMessageBox.information(self, "Éxito", "Información actualizada correctamente.")
                # Vuelve a cargar los datos para mostrar todos los registros
                self.buscar_registros()

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error al actualizar el registro: {str(e)}")

#Radiobuttons - Configuración para detectar que opción del radio button está seleccionada en cada fase

    def toggle_antecedentes_widgets(self):
        if self.id_donante_ant.isChecked():
            self.mostrar_antecedentes_donante()
            #Al tener seleeccionado donante, se setea ant_traumaticos y ant_transfusionales para modo lectura y se cambia el color del widget
            self.ant_traumaticos.setReadOnly(True)
            self.ant_transfusionales.setReadOnly(True)
            
            # Aplica el color gris al fondo  para indicar que no es editable 
            self.ant_traumaticos.setStyleSheet("background-color: lightgray;")
            self.ant_transfusionales.setStyleSheet("background-color: lightgray;")
        else:
            self.mostrar_antecedentes_receptor()
            # Permite editar si el receptor es seleccionado y resetea el color
            self.ant_traumaticos.setReadOnly(False)
            self.ant_transfusionales.setReadOnly(False)
            
            # Resetea el color de fondo al predeterminado
            self.ant_traumaticos.setStyleSheet("")
            self.ant_transfusionales.setStyleSheet("")

    def toggle_f1_widgets(self):
        if self.id_donante_f1.isChecked():
            self.mostrar_f1_donante()
            self.usg_hep_bil.setEditable(False)
            self.usg_hep_bil.setEnabled(False)
            self.obs_usg_hep_bil.setReadOnly(True)
            self.pra.setReadOnly(True)
            self.usg_hep_bil.setStyleSheet("background-color: lightgray;")
            self.obs_usg_hep_bil.setStyleSheet("background-color: lightgray;")
            self.pra.setStyleSheet("background-color: lightgray;")
        else:
            self.mostrar_f1_receptor()
            self.usg_hep_bil.setEditable(True)
            self.usg_hep_bil.setEnabled(True)
            self.obs_usg_hep_bil.setReadOnly(False)
            self.pra.setReadOnly(False)
            self.usg_hep_bil.setStyleSheet("")
            self.obs_usg_hep_bil.setStyleSheet("")
            self.pra.setStyleSheet("")

    def toggle_f2a_widgets(self):
        if self.id_donante_f2a.isChecked():
            self.mostrar_f2a_donante()
            self.vac_influen.setEditable(False)
            self.vac_influen.setEnabled(False)
            self.dos_influen.setReadOnly(True)
            self.vac_neum.setEditable(False)
            self.vac_neum.setEnabled(False)
            self.dos_neum.setReadOnly(True)
            self.cert_dental.setEditable(False)
            self.cert_dental.setEnabled(False)
            self.obs_dental.setReadOnly(True)
            self.vac_influen.setStyleSheet("background-color: lightgray;")
            self.dos_influen.setStyleSheet("background-color: lightgray;")
            self.vac_neum.setStyleSheet("background-color: lightgray;")
            self.dos_neum.setStyleSheet("background-color: lightgray;")
            self.cert_dental.setStyleSheet("background-color: lightgray;")
            self.obs_dental.setStyleSheet("background-color: lightgray;")

        else:
            self.mostrar_f2a_receptor()
            self.vac_influen.setEditable(True)
            self.vac_influen.setEnabled(True)
            self.dos_influen.setReadOnly(False)
            self.vac_neum.setEditable(True)
            self.vac_neum.setEnabled(True)
            self.dos_neum.setReadOnly(False)
            self.cert_dental.setEditable(True)
            self.cert_dental.setEnabled(True)
            self.obs_dental.setReadOnly(False)
            self.vac_influen.setStyleSheet("")
            self.dos_influen.setStyleSheet("")
            self.vac_neum.setStyleSheet("")
            self.dos_neum.setStyleSheet("")
            self.cert_dental.setStyleSheet("")
            self.obs_dental.setStyleSheet("")

    def toggle_f2b_widgets(self):
            if self.id_donante_f2b.isChecked():
                self.mostrar_f2b_donante()
                self.cuantiferon.setEditable(False)
                self.cuantiferon.setEnabled(False)
                self.obs_cuantiferon.setReadOnly(True)
                self.albumina.setReadOnly(True)
                self.sodio.setReadOnly(True)
                self.potasio.setReadOnly(True)
                self.fosforo.setReadOnly(True)
                self.calcio.setReadOnly(True)
                self.magnesio.setReadOnly(True)
                self.pth.setReadOnly(True)
                self.c3.setReadOnly(True)
                self.c4.setReadOnly(True)
                self.anti_dna.setEditable(False)
                self.anti_dna.setEnabled(False)
                self.ana.setEditable(False)
                self.ana.setEnabled(False)
                self.b2_glicoprot.setEditable(False)
                self.b2_glicoprot.setEnabled(False)
                self.anticoag_lupico.setEditable(False)
                self.anticoag_lupico.setEnabled(False)
                self.anticardio_igg_igm.setEditable(False)
                self.anticardio_igg_igm.setEnabled(False)
                self.p_anca.setEditable(False)
                self.p_anca.setEnabled(False)
                self.c_anca.setEditable(False)
                self.c_anca.setEnabled(False)
                self.bnp.setReadOnly(True)
                self.orocultivo.setEditable(False)
                self.orocultivo.setEnabled(False)
                self.obs_orocultivo.setReadOnly(True)
                self.ex_nasal.setEditable(False)
                self.ex_nasal.setEnabled(False)
                self.obs_nasal.setReadOnly(True)
                self.cuantiferon.setStyleSheet("background-color: lightgray;")
                self.obs_cuantiferon.setStyleSheet("background-color: lightgray;")
                self.albumina.setStyleSheet("background-color: lightgray;")
                self.sodio.setStyleSheet("background-color: lightgray;")
                self.potasio.setStyleSheet("background-color: lightgray;")
                self.fosforo.setStyleSheet("background-color: lightgray;")
                self.calcio.setStyleSheet("background-color: lightgray;")
                self.magnesio.setStyleSheet("background-color: lightgray;")
                self.pth.setStyleSheet("background-color: lightgray;")
                self.c3.setStyleSheet("background-color: lightgray;")
                self.c4.setStyleSheet("background-color: lightgray;")
                self.anti_dna.setStyleSheet("background-color: lightgray;")
                self.ana.setStyleSheet("background-color: lightgray;")
                self.b2_glicoprot.setStyleSheet("background-color: lightgray;")
                self.anticoag_lupico.setStyleSheet("background-color: lightgray;")
                self.anticardio_igg_igm.setStyleSheet("background-color: lightgray;")
                self.p_anca.setStyleSheet("background-color: lightgray;")
                self.c_anca.setStyleSheet("background-color: lightgray;")
                self.bnp.setStyleSheet("background-color: lightgray;")
                self.orocultivo.setStyleSheet("background-color: lightgray;")
                self.obs_orocultivo.setStyleSheet("background-color: lightgray;")
                self.ex_nasal.setStyleSheet("background-color: lightgray;")
                self.obs_nasal.setStyleSheet("background-color: lightgray;")

            else:
                self.mostrar_f2b_receptor()
                self.cuantiferon.setEditable(True)
                self.cuantiferon.setEnabled(True)
                self.obs_cuantiferon.setReadOnly(False)
                self.albumina.setReadOnly(False)
                self.sodio.setReadOnly(False)
                self.potasio.setReadOnly(False)
                self.fosforo.setReadOnly(False)
                self.calcio.setReadOnly(False)
                self.magnesio.setReadOnly(False)
                self.pth.setReadOnly(False)
                self.c3.setReadOnly(False)
                self.c4.setReadOnly(False)
                self.anti_dna.setEditable(True)
                self.anti_dna.setEnabled(True)
                self.ana.setEditable(True)
                self.ana.setEnabled(True)
                self.b2_glicoprot.setEditable(True)
                self.b2_glicoprot.setEnabled(True)
                self.anticoag_lupico.setEditable(True)
                self.anticoag_lupico.setEnabled(True)
                self.anticardio_igg_igm.setEditable(True)
                self.anticardio_igg_igm.setEnabled(True)
                self.p_anca.setEditable(True)
                self.p_anca.setEnabled(True)
                self.c_anca.setEditable(True)
                self.c_anca.setEnabled(True)
                self.bnp.setReadOnly(False)
                self.orocultivo.setEditable(True)
                self.orocultivo.setEnabled(True)
                self.obs_orocultivo.setReadOnly(False)
                self.ex_nasal.setEditable(True)
                self.ex_nasal.setEnabled(True)
                self.obs_nasal.setReadOnly(False)
                self.cuantiferon.setStyleSheet("")
                self.obs_cuantiferon.setStyleSheet("")
                self.albumina.setStyleSheet("")
                self.sodio.setStyleSheet("")
                self.potasio.setStyleSheet("")
                self.fosforo.setStyleSheet("")
                self.calcio.setStyleSheet("")
                self.magnesio.setStyleSheet("")
                self.pth.setStyleSheet("")
                self.c3.setStyleSheet("")
                self.c4.setStyleSheet("")
                self.anti_dna.setStyleSheet("")
                self.ana.setStyleSheet("")
                self.b2_glicoprot.setStyleSheet("")
                self.anticoag_lupico.setStyleSheet("")
                self.anticardio_igg_igm.setStyleSheet("")
                self.p_anca.setStyleSheet("")
                self.c_anca.setStyleSheet("")
                self.bnp.setStyleSheet("")
                self.orocultivo.setStyleSheet("")
                self.obs_orocultivo.setStyleSheet("")
                self.ex_nasal.setStyleSheet("")
                self.obs_nasal.setStyleSheet("")


    def toggle_f3_widgets(self):
        if self.id_donante_f3.isChecked():
            self.mostrar_f3_donante()
            self.est_guay_hec.setEditable(False)
            self.est_guay_hec.setEnabled(False) 
            self.obs_guay_hec.setReadOnly(True)
            self.est_endo_colon.setEditable(False)
            self.est_endo_colon.setEnabled(False) 
            self.obs_endo_colon.setReadOnly(True)
            self.est_rx_spn.setEditable(False)
            self.est_rx_spn.setEnabled(False)
            self.obs_rx_spn.setReadOnly(True)
            self.est_usg_vesi.setEditable(False)
            self.est_usg_vesi.setEnabled(False)
            self.obs_usg_vesi.setReadOnly(True)
            self.est_eco_trans.setEditable(False)
            self.est_eco_trans.setEnabled(False)
            self.obs_eco_trans.setReadOnly(True)
            self.est_eco_trans_dm.setEditable(False)
            self.est_eco_trans_dm.setEnabled(False)
            self.obs_eco_trans_dm.setReadOnly(True)
            self.est_dop_iliac.setEditable(False)
            self.est_dop_iliac.setEnabled(False)
            self.obs_dop_iliac.setReadOnly(True)
            self.est_dop_art.setEditable(False)
            self.est_dop_art.setEnabled(False)
            self.obs_dop_art.setReadOnly(True)
            self.est_pielograma.setEditable(True)
            self.obs_pielograma.setReadOnly(False)
            self.est_cisto.setEditable(False)
            self.obs_cisto.setReadOnly(True)
            self.est_guay_hec.setStyleSheet("background-color: lightgray;")
            self.obs_guay_hec.setStyleSheet("background-color: lightgray;")
            self.est_endo_colon.setStyleSheet("background-color: lightgray;")
            self.obs_endo_colon.setStyleSheet("background-color: lightgray;")
            self.est_rx_spn.setStyleSheet("background-color: lightgray;")
            self.obs_rx_spn.setStyleSheet("background-color: lightgray;")
            self.est_usg_vesi.setStyleSheet("background-color: lightgray;")
            self.obs_usg_vesi.setStyleSheet("background-color: lightgray;")
            self.est_eco_trans.setStyleSheet("background-color: lightgray;")
            self.obs_eco_trans.setStyleSheet("background-color: lightgray;")
            self.est_eco_trans_dm.setStyleSheet("background-color: lightgray;")
            self.obs_eco_trans_dm.setStyleSheet("background-color: lightgray;")
            self.est_dop_iliac.setStyleSheet("background-color: lightgray;")
            self.obs_dop_iliac.setStyleSheet("background-color: lightgray;")
            self.est_dop_art.setStyleSheet("background-color: lightgray;")
            self.obs_dop_art.setStyleSheet("background-color: lightgray;")
            self.est_pielograma.setStyleSheet("")
            self.obs_pielograma.setStyleSheet("")
            self.est_cisto.setStyleSheet("background-color: lightgray;")
            self.obs_cisto.setStyleSheet("background-color: lightgray;")
            
        else:
            self.mostrar_f3_receptor()
            self.est_guay_hec.setEditable(True)
            self.est_guay_hec.setEnabled(True) 
            self.obs_guay_hec.setReadOnly(False)
            self.est_endo_colon.setEditable(True)
            self.est_endo_colon.setEnabled(True) 
            self.obs_endo_colon.setReadOnly(False)
            self.est_rx_spn.setEditable(True)
            self.est_rx_spn.setEnabled(True) 
            self.obs_rx_spn.setReadOnly(False)
            self.est_usg_vesi.setEditable(True)
            self.est_usg_vesi.setEnabled(True) 
            self.obs_usg_vesi.setReadOnly(False)
            self.est_eco_trans.setEditable(True)
            self.est_eco_trans.setEnabled(True) 
            self.obs_eco_trans.setReadOnly(False)
            self.est_eco_trans_dm.setEditable(True)
            self.est_eco_trans_dm.setEnabled(True) 
            self.obs_eco_trans_dm.setReadOnly(False)
            self.est_dop_iliac.setEditable(True)
            self.est_dop_iliac.setEnabled(True) 
            self.obs_dop_iliac.setReadOnly(False)
            self.est_dop_art.setEditable(True)
            self.est_dop_art.setEnabled(True) 
            self.obs_dop_art.setReadOnly(False)
            self.est_pielograma.setEditable(False)
            self.est_pielograma.setEnabled(False) 
            self.obs_pielograma.setReadOnly(True)
            self.est_cisto.setEditable(True)
            self.est_cisto.setEnabled(True) 
            self.obs_cisto.setReadOnly(False)
            self.est_guay_hec.setStyleSheet("")
            self.obs_guay_hec.setStyleSheet("")
            self.est_endo_colon.setStyleSheet("")
            self.obs_endo_colon.setStyleSheet("")
            self.est_rx_spn.setStyleSheet("")
            self.obs_rx_spn.setStyleSheet("")
            self.est_usg_vesi.setStyleSheet("")
            self.obs_usg_vesi.setStyleSheet("")
            self.est_eco_trans.setStyleSheet("")
            self.obs_eco_trans.setStyleSheet("")
            self.est_eco_trans_dm.setStyleSheet("")
            self.obs_eco_trans_dm.setStyleSheet("")
            self.est_dop_iliac.setStyleSheet("")
            self.obs_dop_iliac.setStyleSheet("")
            self.est_dop_art.setStyleSheet("")
            self.obs_dop_art.setStyleSheet("")
            self.est_pielograma.setStyleSheet("background-color: lightgray;")
            self.obs_pielograma.setStyleSheet("background-color: lightgray;")
            self.est_cisto.setStyleSheet("")
            self.obs_cisto.setStyleSheet("")

    def toggle_f4_widgets(self):
            if self.id_donante_f4.isChecked():
                self.mostrar_f1_donante()

                self.eval_urologia.setReadOnly(True)
                self.eval_cardiologia.setReadOnly(True)
                self.angiotac_miem_inf.setReadOnly(True)
                self.angiotac_ven_art.setReadOnly(False)
                self.eval_urologia.setStyleSheet("background-color: lightgray;")
                self.eval_cardiologia.setStyleSheet("background-color: lightgray;")
                self.angiotac_miem_inf.setStyleSheet("background-color: lightgray;")
                self.angiotac_ven_art.setStyleSheet("")

            else:
                self.mostrar_f4_receptor()
                self.eval_urologia.setReadOnly(False)
                self.eval_cardiologia.setReadOnly(False)
                self.angiotac_miem_inf.setReadOnly(False)
                self.angiotac_ven_art.setReadOnly(True)
                self.eval_urologia.setStyleSheet("")
                self.eval_cardiologia.setStyleSheet("")
                self.angiotac_miem_inf.setStyleSheet("")
                self.angiotac_ven_art.setStyleSheet("background-color: lightgray;")

    def exportar_docx(self):
        # Obtiene el ID del receptor de la fase final seleccionada
        selected_row = self.result_table_9.currentRow()
        if selected_row < 0:
            QMessageBox.warning(self, "Error", "Por favor, seleccione una fase final.")
            return

        id_receptor = self.result_table_9.item(selected_row, 1).text()
        id_donante = self.result_table_9.item(selected_row, 2).text()

        # Obtiene datos de fase_final
        fase_final_data = self.db.obtener_datos_fase_final_por_receptor(id_receptor)
        if not fase_final_data:
            QMessageBox.warning(self, "Error", "No se encontraron datos de fase final para el receptor especificado.")
            return

        # Obtiene datos del receptor
        receptor_data = self.db.obtener_datos_receptor_por_id(id_receptor)
        if not receptor_data:
            QMessageBox.warning(self, "Error", "No se encontraron datos del receptor.")
            return

        # Función para mapear valores numéricos a texto
        def map_value(value, mapping):
            if isinstance(value, (int, float)):
                return mapping.get(int(value), str(value))
            return str(value)

        # Crear un diccionario con los datos
        data = {

            #Receptor

            'FECHA_REGISTRO1': receptor_data['fecha_registro'].strftime('%d/%m/%Y'),
            'NOMBRE1': receptor_data['nombre'],
            'EDAD1': str(receptor_data['edad']),
            'ETNIA': map_value(receptor_data['etnia'], {v: k for k, v in self.etnia_map.items()}),
            'SEXO': map_value(receptor_data['sexo'], {v: k for k, v in self.sexo_map.items()}),
            'FECHA_ERC': receptor_data['fecha_dg_erc'].strftime('%d/%m/%Y'),
            'TER_SUST': map_value(receptor_data['ter_sust_act'], {v: k for k, v in self.ter_sust_act_map.items()}),
            'INST_HD': map_value(receptor_data['inst_provee_hd'], {v: k for k, v in self.inst_provee_hd_map.items()}),
            'TIPO_DON': map_value(fase_final_data['tipo_donante'], {v: k for k, v in self.tipo_donante_map.items()}),
            'VOL_RESI': str(receptor_data['vol_residual']),
            'T_ANURIA': str(receptor_data['tiempo_anuria']),
            'EXP_CLIN': str(receptor_data['id_receptor']),
            'TIPO_SANG_': map_value(receptor_data['grupo_sanguineo'], {v: k for k, v in self.grupo_sanguineo_map.items()}),
            'PROCEDENCIA': map_value(receptor_data['procedencia'], {v: k for k, v in self.procedencia_map.items()}),
            'RESIDENCIA': map_value(receptor_data['residencia'], {v: k for k, v in self.residencia_map.items()}),
            'OCUPACION': map_value(receptor_data['ocupacion'], {v: k for k, v in self.ocupacion_map.items()}),
            'ETIOLOGIA_ERC': map_value(receptor_data['etiologia_erc'], {v: k for k, v in self.etiologia_erc_map.items()}),
            'T_TER_SUST': receptor_data['tiempo_ini_st_renal'].strftime('%d/%m/%Y'),
            'T_DIALISIS': str(receptor_data['tiempo_dialisis']),
            'PARENTESCO': map_value(fase_final_data['parentesco'], {v: k for k, v in self.parentesco_map.items()}),
            'RIESGO_CMV': map_value(receptor_data['riesgo_cmv'], {v: k for k, v in self.riesgo_cmv_map.items()}),
            'YD': str(receptor_data['yd']),
            'YI': str(receptor_data['yi']),
            'FD': str(receptor_data['fd']),
            'FI': str(receptor_data['fi']),
            'OBS_ACC_VASC': receptor_data['obs_acc_vasc']

            ##Donante


            ##Antecedentes

            ##Fase 1

            ##Fase 2

            ##Fase 3

            ##Fase 4


        }

        # Imprime el diccionario data para depuración
        print("Contenido del diccionario data:")
        pprint.pprint(data)

        # Carga la plantilla
        doc = Document('plantilla.docx')

        # Reemplaza los marcadores en la plantilla
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                for key, value in data.items():
                    if '{{' + key + '}}' in run.text:
                        run.text = run.text.replace('{{' + key + '}}', str(value))

        # También busca en las tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            for key, value in data.items():
                                if '{{' + key + '}}' in run.text:
                                    run.text = run.text.replace('{{' + key + '}}', str(value))

        # Genera el nombre del archivo
        desktop = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop')
        base_filename = f'Informe_Final_{id_receptor}_{id_donante}'
        counter = 1
        output_path = os.path.join(desktop, f'{base_filename}.docx')

        # Verifica si el archivo ya existe y añade un número si es necesario
        while os.path.exists(output_path):
            output_path = os.path.join(desktop, f'{base_filename}_{counter}.docx')
            counter += 1

        # Guarda el documento
        doc.save(output_path)

        QMessageBox.information(self, "Éxito", f"Informe generado y guardado en: {output_path}")


class Application(QApplication):
    def __init__(self, argv):
        super().__init__(argv)
        self.login_window = LoginWindow()
        self.main_window = None
        self.login_window.sesion_exitosa.connect(self.mostrar_main_window)
        self.login_window.show()

    def mostrar_main_window(self):
        if self.main_window is None:
            self.main_window = MainWindow()
        self.main_window.show()
        self.login_window.close()


class contraseña_dialogo(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Verificación de administrador")
        self.setFixedSize(350, 150)

        # Aplica un estilo de fondo celeste pastel
        self.setStyleSheet("""
            QDialog {
                background-color: #B3E5FC;  /* Color celeste pastel */
                border-radius: 10px;
            }
            QLabel {
                color: #333;  /* Color del texto */
            }
            QLineEdit {
                background-color: #FFFFFF;  /* Fondo blanco para el campo de texto */
                color: #333;
                border: 1px solid #90CAF9;
                border-radius: 5px;
                padding: 5px;
            }
            QPushButton {
                background-color: #64B5F6;  /* Azul claro */
                color: white;
                border-radius: 5px;
                padding: 5px 15px;
            }
            QPushButton:hover {
                background-color: #42A5F5;  /* Azul más oscuro al pasar el mouse */
            }
        """)

        layout = QVBoxLayout(self)
        label = QLabel("Ingresa tu contraseña:")
        layout.addWidget(label)

        self.entrada_constraseña = QLineEdit()
        self.entrada_constraseña.setEchoMode(QLineEdit.Password)
        layout.addWidget(self.entrada_constraseña)

        ok_button = QPushButton("Aceptar")
        ok_button.clicked.connect(self.accept)
        layout.addWidget(ok_button)

    def obtener_contraseña(self):
        if self.exec_() == QDialog.Accepted:
            return self.entrada_constraseña.text()
        return None
    
if __name__ == "__main__":
    app = QApplication(sys.argv)
    
    # Crea la ventana principal
    window = MainWindow()
    
    # La ventana principal se muestra automáticamente después del login exitoso
    sys.exit(app.exec_())
    